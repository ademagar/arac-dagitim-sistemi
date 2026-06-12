"""
Thesis writing guide compliance fix script.
Applies all required changes to END4000_TEZ_FINAL.docx per the thesis writing guide rules.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import copy
from lxml import etree


def set_cell_text(cell, text, bold=False, font_size=12):
    """Clear cell and write text with formatting."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def get_para_text(p):
    return ''.join(r.text or '' for r in p.findall(f'.//{qn("w:t")}'))


def fix_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run containing new_text."""
    p_el = para._p
    # Remove all runs
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)
    # Add new run
    r_el = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = new_text
    if ' ' in new_text and (new_text.startswith(' ') or new_text.endswith(' ')):
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_el.append(t_el)
    p_el.append(r_el)


def copy_rpr(src_para, dst_run_el):
    """Copy rPr (run properties) from first run of src_para to dst_run_el."""
    first_r = src_para._p.find(qn('w:r'))
    if first_r is not None:
        rpr = first_r.find(qn('w:rPr'))
        if rpr is not None:
            dst_run_el.insert(0, copy.deepcopy(rpr))


def make_caption_paragraph(doc, text):
    """Create a new Caption-style paragraph with given text."""
    para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Caption')
    pPr.append(pStyle)
    para.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    para.append(r)
    return para


def make_normal_paragraph(doc, text):
    """Create a Normal paragraph with given text."""
    para = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    para.append(r)
    return para


def build_literature_table(doc):
    """Build a literature review summary table and return its XML element."""
    # Create table: 15 rows (1 header + 14 papers) x 5 columns
    tbl = doc.add_table(rows=15, cols=5)
    tbl.style = 'Table Grid'

    headers = ['Writer(s)', 'Year', 'Objective of Study', 'Method', 'Key Findings']
    data = [
        ('Cleveland et al.', '1990',
         'Decompose time series into trend, seasonal, and remainder components',
         'STL (Loess smoothing)',
         'Robust, iterative decomposition handles outliers and varying seasonal strength'),
        ('Crowder', '1989',
         'Design exponentially weighted moving average control charts',
         'Statistical optimization',
         'Optimal smoothing factor λ derived analytically for various process shift sizes'),
        ('Dantzig', '1963',
         'Formalize linear programming as a general optimization framework',
         'Simplex method',
         'Foundation of LP theory enabling large-scale real-world optimization'),
        ('Hwang & Yoon', '1981',
         'Rank alternatives under multiple conflicting criteria',
         'TOPSIS, MADM',
         'Systematic framework for multi-attribute ranking with weighted criteria'),
        ('Hyndman et al.', '2011',
         'Optimal combination of forecasts for hierarchical time series',
         'MinT reconciliation',
         'Trace minimization guarantees coherent, unbiased hierarchical forecasts'),
        ('Kok et al.', '2008',
         'Review assortment planning literature and industry practice',
         'Stochastic demand models',
         'Consumer demand substitution fundamentally affects optimal stocking decisions'),
        ('Koren et al.', '2009',
         'Improve collaborative filtering accuracy for recommender systems',
         'Matrix factorization (SVD++)',
         'Latent factor models outperform neighbourhood methods on sparse rating data'),
        ('Mahajan & van Ryzin', '2001',
         'Stock retail assortments under dynamic consumer substitution',
         'Stochastic programming',
         'Substitution raises optimal stock levels above independent-demand solutions'),
        ('Powell', '1996',
         'Stochastic vehicle assignment for truckload motor carriers',
         'Stochastic dynamic programming',
         'Dynamic assignment framework reduces empty miles in freight networks'),
        ('Rendle', '2010',
         'Model feature interactions for large-scale recommendation tasks',
         'Factorization machines',
         'Polynomial feature interactions captured in linear time with factored parameters'),
        ('Şahin & Kılıç', '2022',
         'Develop alternative MILP formulations for vehicle allocation',
         'MILP, Branch & Price',
         'Alternative formulations reduce computation time significantly for large VAP instances'),
        ('Taylor & Letham', '2018',
         'Scalable automated time series forecasting for business applications',
         'Prophet (additive decomposition)',
         'Analyst-adjustable model outperforms ARIMA on business series with holidays'),
        ('Toth & Vigo', '2002',
         'Survey exact and heuristic algorithms for the vehicle routing problem',
         'VRP algorithms (exact & heuristic)',
         'Comprehensive taxonomy of routing variants and solution methodologies'),
        ('Wickramasuriya et al.', '2019',
         'Reconcile hierarchical forecasts to satisfy aggregation constraints',
         'MinT (Minimum Trace)',
         'Reconciled forecasts are unbiased and minimize total mean squared error'),
    ]

    # Set header row
    for j, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[j], h, bold=True, font_size=11)

    # Set data rows
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            set_cell_text(tbl.rows[i + 1].cells[j], val, font_size=11)

    # Return the table's XML element (will be removed from doc.tables)
    return tbl._tbl


def main():
    doc = Document('END4000_TEZ_FINAL.docx')
    body = doc.element.body
    paras = doc.paragraphs  # live list

    print(f"Total paragraphs: {len(paras)}")
    print(f"Total tables: {len(doc.tables)}")

    # =========================================================
    # FIX 1: Heading 4.2 — remove "-- BURASI HATALI"
    # =========================================================
    for i, p in enumerate(paras):
        text = get_para_text(p._p)
        if '-- BURASI HATALI' in text:
            clean = text.replace('  -- BURASI HATALI', '').replace(' -- BURASI HATALI', '').replace('-- BURASI HATALI', '')
            fix_para_text(p, clean)
            print(f"FIX 1: Cleaned heading at P{i}: {clean}")

    # =========================================================
    # FIX 2: Heading 5.5 — remove trailing ")"
    # =========================================================
    for i, p in enumerate(paras):
        text = get_para_text(p._p)
        if 'Composite Score Computation)' in text:
            clean = text.replace('Composite Score Computation)', 'Composite Score Computation')
            fix_para_text(p, clean)
            print(f"FIX 2: Fixed heading at P{i}: {clean}")

    # =========================================================
    # FIX 3: Figure 1 caption — replace template text
    # =========================================================
    for i, p in enumerate(paras):
        text = get_para_text(p._p)
        if 'Equations into an MS Word Document' in text:
            fix_para_text(p, 'Figure 1. Three-Layer Decision-Support System Architecture')
            print(f"FIX 3: Fixed Figure 1 caption at P{i}")

    # =========================================================
    # FIX 4: Make figure captions use Caption style
    # =========================================================
    figure_caption_texts = [
        'Figure 2. STL Decomposition',
        'Figure 3. Seasonal Index',
        'Figure 4. EWMA Window Selection',
        'Figure 5. MCDM Composite Score',
        'Figure 6. January 2026 Allocation Results',
    ]
    for i, p in enumerate(paras):
        text = get_para_text(p._p)
        for fc in figure_caption_texts:
            if fc in text:
                # Set style to Caption
                pPr = p._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p._p.insert(0, pPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is None:
                    pStyle = OxmlElement('w:pStyle')
                    pPr.insert(0, pStyle)
                pStyle.set(qn('w:val'), 'Caption')
                print(f"FIX 4: Set Caption style at P{i}: {text[:60]}")

    # Reload paragraphs after modifications
    paras = doc.paragraphs

    # =========================================================
    # FIX 5: Table data corrections
    # =========================================================

    # Table 8 (doc index) = Vehicle Portfolio Table
    t8 = doc.tables[8]
    corrections_t8 = {1: (3, '35 units'), 2: (3, '291 units'), 3: (3, '97 units'), 4: (3, '180 units')}
    for row_idx, (col_idx, val) in corrections_t8.items():
        cell = t8.rows[row_idx].cells[col_idx]
        set_cell_text(cell, val)
    print("FIX 5a: Fixed Table 8 (vehicle portfolio) inventory numbers")

    # Table 13 = Depot Inventory Breakdown
    t13 = doc.tables[13]
    corrections_t13 = {
        1: [(1, '35'), (2, 'Black2, Grey3, Black1, Grey2'), (3, '5.80%')],
        2: [(1, '291'), (2, 'Black2, Grey3, Black1, Grey2'), (3, '48.26%')],
        3: [(1, '97'), (2, 'Black2, Grey3, Black1, Grey2'), (3, '16.09%')],
        4: [(1, '180'), (2, 'Black2, Grey3, Black1, Grey2'), (3, '29.85%')],
        5: [(1, '603'), (2, '4 colors per model'), (3, '100%')],
    }
    for row_idx, col_vals in corrections_t13.items():
        for col_idx, val in col_vals:
            set_cell_text(t13.rows[row_idx].cells[col_idx], val)
    print("FIX 5b: Fixed Table 13 (depot inventory) numbers")

    # Table 14 = Allocation Summary
    t14 = doc.tables[14]
    corrections_t14 = {
        1: [(1, '28'), (2, '35'), (3, '80.0%')],
        2: [(1, '227'), (2, '291'), (3, '78.0%')],
        3: [(1, '75'), (2, '97'), (3, '77.3%')],
        4: [(1, '171'), (2, '180'), (3, '95.0%')],
        5: [(1, '501'), (2, '603'), (3, '83.1%')],
    }
    for row_idx, col_vals in corrections_t14.items():
        for col_idx, val in col_vals:
            set_cell_text(t14.rows[row_idx].cells[col_idx], val)
    print("FIX 5c: Fixed Table 14 (allocation summary) numbers")

    # =========================================================
    # FIX 6: Fix para 257 — wrong A-group inventory sum
    # =========================================================
    for i, p in enumerate(paras):
        text = get_para_text(p._p)
        if '42+331+176=549' in text or '42+331+176 = 549' in text:
            new_text = text.replace('42+331+176=549', '35+291+97=423').replace('42+331+176 = 549', '35+291+97=423')
            fix_para_text(p, new_text)
            print(f"FIX 6: Fixed A-group inventory sum at P{i}")
        elif 'A-group inventory (' in text and ('42' in text or '549' in text):
            # Broader match for the sentence
            import re
            new_text = re.sub(r'\(42\+331\+176\s*=\s*549\s*units?\)', '(35+291+97=423 units)', text)
            if new_text != text:
                fix_para_text(p, new_text)
                print(f"FIX 6b: Fixed A-group inventory sum at P{i}")

    # =========================================================
    # FIX 7: Rename table caption texts (Table A→3, B→4, C→5, E→7, F→8)
    # New numbering with literature table as Table 1:
    # Table 1 = Literature Review (new)
    # Table 2 = Vehicle Portfolio (was Table 1)
    # Table 3 = Monthly Allocation Plan (was Table A)
    # Table 4 = EWMA Window Analysis (was Table B)
    # Table 5 = MCDM Criteria (was Table C)
    # Table 6 = MILP Notation (was Table D, no caption existed)
    # Table 7 = Depot Inventory (was Table E)
    # Table 8 = Allocation Summary (was Table F)
    # =========================================================
    table_renames = {
        'Table 1:': 'Table 2.',
        'Table 1.': 'Table 2.',
        'Table A.': 'Table 3.',
        'Table B.': 'Table 4.',
        'Table C.': 'Table 5.',
        'Table E.': 'Table 7.',
        'Table F.': 'Table 8.',
    }
    for i, p in enumerate(paras):
        text = get_para_text(p._p)
        for old, new in table_renames.items():
            if text.startswith(old):
                new_text = text.replace(old, new, 1)
                fix_para_text(p, new_text)
                # Ensure Caption style
                pPr = p._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p._p.insert(0, pPr)
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is None:
                    pStyle = OxmlElement('w:pStyle')
                    pPr.insert(0, pStyle)
                pStyle.set(qn('w:val'), 'Caption')
                print(f"FIX 7: Renamed '{old}' → '{new}' at P{i}: {new_text[:60]}")

    # =========================================================
    # FIX 8: Add "Table 6. MILP Model Notation and Parameters" caption
    # before docx Table 12 (MILP notation table)
    # and fix caption for current vehicle portfolio Table 1 → Table 2
    # =========================================================
    # Find Table 12 in body children and insert caption before it
    children = list(body)
    milp_tbl_elem = doc.tables[12]._tbl
    milp_idx = None
    for idx, child in enumerate(children):
        if child is milp_tbl_elem:
            milp_idx = idx
            break

    if milp_idx is not None:
        # Check if previous sibling is already a caption for Table 6
        prev_child = children[milp_idx - 1]
        runs = prev_child.findall(f'.//{qn("w:t")}')
        prev_text = ''.join(r.text or '' for r in runs)
        if 'Table 6' not in prev_text:
            # Insert caption paragraph before the MILP table
            cap_para = make_caption_paragraph(doc, 'Table 6. MILP Model Notation and Parameters')
            body.insert(milp_idx, cap_para)
            print(f"FIX 8: Inserted 'Table 6. MILP Model Notation and Parameters' before MILP table")

    # =========================================================
    # FIX 9: Insert Literature Review Table (Table 1) in Chapter 2
    # Place it after the last paragraph of section 2.5 (just before METHODOLOGY heading)
    # =========================================================
    # Reload paras after edits
    paras = doc.paragraphs
    children = list(body)

    # Find the METHODOLOGY (Heading 1) position
    methodology_elem = None
    for p in paras:
        text = get_para_text(p._p)
        if text.strip() == 'METHODOLOGY':
            methodology_elem = p._p
            break

    if methodology_elem is not None:
        meth_idx = None
        for idx, child in enumerate(list(body)):
            if child is methodology_elem:
                meth_idx = idx
                break

        if meth_idx is not None:
            # Build literature table
            lit_tbl_elem = build_literature_table(doc)

            # Insert caption + empty para + table just before METHODOLOGY heading
            # Order: ... [last Ch2 content] [cap_para] [lit_table] [empty] [METHODOLOGY]
            cap_para_el = make_caption_paragraph(doc, 'Table 1. Literature Review Summary')
            empty_para_el = make_normal_paragraph(doc, '')

            body.insert(meth_idx, empty_para_el)   # after table, before METHODOLOGY
            body.insert(meth_idx, lit_tbl_elem)     # table
            body.insert(meth_idx, cap_para_el)      # caption before table

            print(f"FIX 9: Inserted Literature Review Table (Table 1) before METHODOLOGY at body index {meth_idx}")

    doc.save('END4000_TEZ_FINAL.docx')
    print("\n✓ Saved END4000_TEZ_FINAL.docx with all writing guide fixes applied.")
    print("\nSummary of changes:")
    print("  1. Removed '-- BURASI HATALI' from heading 4.2")
    print("  2. Fixed '5.5. Composite Score Computation)' → removed trailing ')'")
    print("  3. Fixed Figure 1 caption (was template text, now system architecture)")
    print("  4. Applied Caption style to Figure 2-6 captions")
    print("  5a. Fixed Table 2 (vehicle portfolio): A1V01=35, A2V02=291, A3V02=97, B1V01=180")
    print("  5b. Fixed Table 7 (depot inventory): corrected all units and percentages")
    print("  5c. Fixed Table 8 (allocation summary): corrected all allocation/available numbers")
    print("  6. Fixed para text: A-group inventory sum (42+331+176=549 → 35+291+97=423)")
    print("  7. Renamed table captions: A→3, B→4, C→5, E→7, F→8")
    print("  8. Added 'Table 6. MILP Model Notation and Parameters' caption")
    print("  9. Added Literature Review Summary Table (Table 1) before METHODOLOGY chapter")


if __name__ == '__main__':
    main()
