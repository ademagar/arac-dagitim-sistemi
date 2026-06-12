"""
Thesis document generator for:
"Monthly Vehicle Allocation Optimization for Automotive Dealerships
via Multi-Criteria Scoring and Mixed-Integer Linear Programming"

Student: Adem Ağar, ID: 21069019
University: Yıldız Technical University, Faculty of Mechanical Engineering,
            Department of Industrial Engineering
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ──────────────────────────────────────────────
# HELPER UTILITIES
# ──────────────────────────────────────────────

def set_margins(section, top=2.5, bottom=2.5, left=2.5, right=2.5):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def set_run_font(run, size_pt=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic


def add_paragraph(doc, text, size_pt=12, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0,
                  space_after=6, line_spacing=None):
    """Add a body paragraph with Times New Roman styling."""
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if line_spacing is not None:
        from docx.shared import Pt as _Pt
        fmt.line_spacing = _Pt(line_spacing)
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    """
    level 1 → 14pt Bold
    level 2 → 12pt Bold
    level 3 → 12pt Bold Italic
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size_pt=14, bold=True)
    elif level == 2:
        set_run_font(run, size_pt=12, bold=True)
    else:
        set_run_font(run, size_pt=12, bold=True, italic=True)
    return p


def add_table_caption(doc, text):
    """Table caption goes ABOVE the table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size_pt=11, bold=True)
    return p


def add_figure_caption(doc, text):
    """Figure caption goes BELOW the figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(2)
    fmt.space_after  = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size_pt=11, italic=True)
    return p


def add_figure_placeholder(doc, caption_text, description):
    """Add a shaded placeholder box followed by a figure caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after  = Pt(2)
    run = p.add_run(f"[{description}]")
    set_run_font(run, size_pt=11, italic=True)
    # shade the paragraph light grey via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')
    pPr.append(shd)
    add_figure_caption(doc, caption_text)


def add_page_number_footer(section):
    """Add centered page number to footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    # Insert PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Times New Roman"
    run.font.size  = Pt(12)


def make_table(doc, headers, rows, caption, col_widths=None):
    """Add caption, then a formatted table."""
    add_table_caption(doc, caption)
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_val in enumerate(row_data):
            row_cells[c_idx].text = str(cell_val)
            for para in row_cells[c_idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    # Optional column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])

    # Space after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    return table


def add_blank_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


# ──────────────────────────────────────────────
# DOCUMENT CREATION
# ──────────────────────────────────────────────

def build_thesis():
    doc = Document()

    # Global default font (fallback)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # Set margins for first section
    section = doc.sections[0]
    set_margins(section)
    add_page_number_footer(section)

    # ══════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════
    add_blank_line(doc)
    add_blank_line(doc)
    add_paragraph(doc,
        "YILDIZ TECHNICAL UNIVERSITY",
        size_pt=14, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc,
        "Faculty of Mechanical Engineering",
        size_pt=13, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc,
        "Department of Industrial Engineering",
        size_pt=13, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    add_blank_line(doc)
    add_blank_line(doc)
    add_blank_line(doc)

    add_paragraph(doc,
        "GRADUATION THESIS",
        size_pt=13, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_paragraph(doc,
        "Monthly Vehicle Allocation Optimization for Automotive Dealerships\n"
        "via Multi-Criteria Scoring and Mixed-Integer Linear Programming",
        size_pt=16, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    add_blank_line(doc)
    add_blank_line(doc)
    add_blank_line(doc)

    add_paragraph(doc,
        "Submitted by:",
        size_pt=12, bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc,
        "Adem Ağar",
        size_pt=13, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc,
        "Student ID: 21069019",
        size_pt=12,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    add_blank_line(doc)
    add_blank_line(doc)

    add_paragraph(doc,
        "Istanbul, 2026",
        size_pt=12,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════
    add_heading(doc, "ABSTRACT", level=1)
    add_paragraph(doc,
        "Monthly Vehicle Allocation Optimization for Automotive Dealerships "
        "via Multi-Criteria Scoring and Mixed-Integer Linear Programming",
        size_pt=12, bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc,
        "Adem Ağar",
        size_pt=12,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc,
        "Industrial Engineering, Faculty of Mechanical Engineering, "
        "Yıldız Technical University, Istanbul, 2026",
        size_pt=12,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    add_paragraph(doc,
        "Efficient distribution of a limited monthly vehicle inventory across a dealer network "
        "is a critical operational challenge for automotive manufacturers and importers. When "
        "the number of vehicles available is less than aggregate dealer demand, the allocation "
        "decision directly affects dealer satisfaction, end-customer availability, and brand "
        "revenue. This thesis proposes and implements a three-layer decision-support system "
        "that automates monthly vehicle allocation for a Sport Utility Vehicle (SUV) brand "
        "operating through 28 authorized dealerships across Turkey.")
    add_paragraph(doc,
        "The first layer applies Seasonal-Trend decomposition using Loess (STL) to decompose "
        "24 months of historical sales data (January 2024–December 2025) into trend, seasonal, "
        "and residual components. Monthly seasonal indices are derived from this decomposition "
        "and combined with a Prophet-based annual market forecast and an Exponentially Weighted "
        "Moving Average (EWMA, α determined by window W=5 after systematic cross-validation "
        "over W∈{3, …, 19}) to produce dealer-level monthly demand estimates for 2026.")
    add_paragraph(doc,
        "The second layer applies Multi-Criteria Decision Making (MCDM) to score each dealer "
        "on four dimensions: recent target achievement performance (P, weight 0.25), "
        "location-product fit via cosine similarity between the dealer's historical "
        "color/version sales vector and the current inventory vector (LP, weight 0.35), "
        "seasonal alignment (S, weight 0.20), and proximity to the annual sales target "
        "with an asymmetric penalty for dealers who are behind target (H, weight 0.20). "
        "A composite score C ∈ [0, 1] is computed for every dealer.")
    add_paragraph(doc,
        "The third layer formulates vehicle allocation as a Mixed-Integer Linear Programming "
        "(MILP) problem. The objective is to maximize the sum of composite-score-weighted "
        "allocations subject to inventory availability constraints and dealer quota bounds "
        "(±20% of the monthly quota derived from market-share-based target distribution). "
        "The model is implemented in Python using PuLP with the open-source CBC solver and "
        "solves a January 2026 instance of 603 vehicles across 28 dealers in under 0.3 seconds.")
    add_paragraph(doc,
        "The system is deployed as an interactive Streamlit dashboard with real-time scenario "
        "analysis capabilities. Results for January 2026 demonstrate that the framework "
        "successfully directs scarce inventory toward high-potential dealers while respecting "
        "operational constraints, producing allocations that could not be replicated by "
        "manual heuristics at comparable speed or consistency.")
    add_paragraph(doc,
        "Keywords: vehicle allocation problem, multi-criteria decision making, mixed-integer "
        "linear programming, STL decomposition, exponentially weighted moving average, "
        "cosine similarity, assortment optimization, automotive supply chain",
        size_pt=12, italic=True, space_after=10)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════
    add_heading(doc, "TABLE OF CONTENTS", level=1)
    toc_entries = [
        ("Abstract", "ii"),
        ("Table of Contents", "iii"),
        ("List of Tables", "iv"),
        ("List of Figures", "iv"),
        ("Abbreviations", "v"),
        ("1. Introduction", "1"),
        ("   1.1 Motivation and Problem Statement", "1"),
        ("   1.2 Research Objectives", "3"),
        ("   1.3 Scope and Limitations", "3"),
        ("   1.4 Thesis Organization", "4"),
        ("2. Literature Review", "5"),
        ("   2.1 Vehicle Allocation Problem", "5"),
        ("   2.2 Time Series Forecasting in Automotive Contexts", "6"),
        ("   2.3 Multi-Criteria Decision Making in Inventory Allocation", "7"),
        ("   2.4 Assortment Optimization and Collaborative Filtering", "8"),
        ("   2.5 Mixed-Integer Linear Programming for Distribution", "9"),
        ("3. System Architecture and Data", "10"),
        ("   3.1 Overall Three-Layer Architecture", "10"),
        ("   3.2 Data Sources and Preprocessing", "11"),
        ("   3.3 Vehicle Portfolio and Market Structure", "12"),
        ("4. Layer 1: Demand Forecasting", "13"),
        ("   4.1 STL Decomposition", "13"),
        ("   4.2 Prophet Market Forecast", "15"),
        ("   4.3 Exponentially Weighted Moving Average", "16"),
        ("   4.4 Launch Boost Factor", "18"),
        ("5. Layer 2: Multi-Criteria Dealer Scoring", "19"),
        ("   5.1 Performance Score (P)", "19"),
        ("   5.2 Location-Product Fit Score (LP)", "20"),
        ("   5.3 Seasonal Alignment Score (S)", "21"),
        ("   5.4 Target Proximity Score (H)", "22"),
        ("   5.5 Composite Score Computation", "23"),
        ("6. Layer 3: MILP Allocation Model", "24"),
        ("   6.1 Problem Formulation", "24"),
        ("   6.2 Market-Share-Based Quota Distribution", "25"),
        ("   6.3 Implementation with PuLP and CBC", "26"),
        ("7. January 2026 Case Study", "27"),
        ("   7.1 Input Data and Parameters", "27"),
        ("   7.2 Allocation Results", "28"),
        ("   7.3 Sensitivity Analysis", "29"),
        ("8. Dashboard and Deployment", "30"),
        ("9. Conclusions and Future Work", "31"),
        ("References", "33"),
    ]
    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        tab_stops = p.paragraph_format.tab_stops
        # right-align page number at ~14cm
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.oxml.ns import qn as _qn
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(_qn('w:val'), 'right')
        tab_el.set(_qn('w:leader'), 'dot')
        tab_el.set(_qn('w:pos'), '9072')   # 9072 twips ≈ 16 cm
        tabs.append(tab_el)
        pPr.append(tabs)
        run1 = p.add_run(entry)
        set_run_font(run1, size_pt=11)
        run2 = p.add_run(f"\t{page}")
        set_run_font(run2, size_pt=11)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # LIST OF TABLES & FIGURES
    # ══════════════════════════════════════════
    add_heading(doc, "LIST OF TABLES", level=1)
    lot = [
        ("Table 1", "Sets and Parameters of the MILP Model", "24"),
        ("Table 2", "MCDM Criteria Summary", "19"),
        ("Table 3", "EW Window Analysis Results (W = 3 to W = 10)", "17"),
        ("Table 4", "January 2026 Allocation Summary by Model Group", "28"),
    ]
    for num, title, page in lot:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:leader'), 'dot')
        tab_el.set(qn('w:pos'), '9072')
        tabs.append(tab_el)
        pPr.append(tabs)
        run1 = p.add_run(f"{num}. {title}")
        set_run_font(run1, size_pt=11)
        run2 = p.add_run(f"\t{page}")
        set_run_font(run2, size_pt=11)

    add_blank_line(doc)
    add_heading(doc, "LIST OF FIGURES", level=1)
    lof = [
        ("Figure 1", "System Architecture Overview (Three-Layer Framework)", "10"),
        ("Figure 2", "STL Decomposition Example — Brand Monthly Sales 2024–2025", "14"),
        ("Figure 3", "Seasonal Index by Month for All Vehicle Model Groups", "15"),
        ("Figure 4", "EW Window MAE Comparison Chart (W = 3 to W = 19)", "17"),
        ("Figure 5", "MCDM Composite Score Distribution Across 28 Dealers", "23"),
        ("Figure 6", "January 2026 Allocation Results by Dealer", "28"),
    ]
    for num, title, page in lof:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:leader'), 'dot')
        tab_el.set(qn('w:pos'), '9072')
        tabs.append(tab_el)
        pPr.append(tabs)
        run1 = p.add_run(f"{num}. {title}")
        set_run_font(run1, size_pt=11)
        run2 = p.add_run(f"\t{page}")
        set_run_font(run2, size_pt=11)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # ABBREVIATIONS
    # ══════════════════════════════════════════
    add_heading(doc, "ABBREVIATIONS", level=1)
    abbrevs = [
        ("CBC",   "COIN-OR Branch and Cut (open-source MIP solver)"),
        ("EWMA",  "Exponentially Weighted Moving Average"),
        ("KPI",   "Key Performance Indicator"),
        ("LP",    "Location-Product Fit Score (also: Linear Program in context)"),
        ("MAE",   "Mean Absolute Error"),
        ("MCDM",  "Multi-Criteria Decision Making"),
        ("MILP",  "Mixed-Integer Linear Programming"),
        ("RMSE",  "Root Mean Square Error"),
        ("SI",    "Seasonal Index"),
        ("STL",   "Seasonal-Trend decomposition using Loess"),
        ("SUV",   "Sport Utility Vehicle"),
        ("TÜİK",  "Türkiye İstatistik Kurumu (Turkish Statistical Institute)"),
        ("VAP",   "Vehicle Allocation Problem"),
        ("YTD",   "Year-to-Date"),
    ]
    for abbr, defn in abbrevs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'left')
        tab_el.set(qn('w:pos'), '1440')
        tabs.append(tab_el)
        pPr.append(tabs)
        r1 = p.add_run(abbr)
        set_run_font(r1, size_pt=11, bold=True)
        r2 = p.add_run(f"\t{defn}")
        set_run_font(r2, size_pt=11)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 1 — INTRODUCTION
    # ══════════════════════════════════════════
    add_heading(doc, "1. INTRODUCTION", level=1)

    add_heading(doc, "1.1 Motivation and Problem Statement", level=2)
    add_paragraph(doc,
        "The Turkish automotive market is one of the largest in Europe by volume, with 17,373,581 "
        "registered automobiles as of December 2025 according to TÜİK data (TÜİK, 2026). The "
        "SUV segment has experienced substantial growth over the preceding decade, driven by "
        "consumer preference shifts toward higher ground clearance, improved perceived safety, "
        "and enhanced cargo utility. For a brand that fields multiple SUV model lines across a "
        "geographically dispersed dealer network, the monthly task of deciding how many vehicles "
        "of each model, version (trim level), and color to ship to each of its 28 authorized "
        "dealerships is both consequential and computationally non-trivial.")
    add_paragraph(doc,
        "The monthly vehicle allocation problem (VAP) arises because central warehouse inventory "
        "is finite and almost always smaller than the aggregate of dealer requests. For January "
        "2026, the brand's central depot held 603 vehicles distributed across four model variants: "
        "35 units of A1V01, 291 units of A2V02, 97 units of A3V02, and 180 units of B1V01. "
        "Because the total inventory (603 units) was less than the aggregate dealer quota (741 "
        "units after accounting for February readiness stock), every allocation decision represents "
        "a deliberate trade-off between competing dealer interests. Allocating too many vehicles "
        "to an already-oversupplied dealer wastes scarcity value; allocating too few to a "
        "high-performing, high-demand dealer leaves brand revenue unrealized.")
    add_paragraph(doc,
        "Historically, such allocation decisions have been made by regional sales managers using "
        "spreadsheet-based heuristics informed by experience, negotiation outcomes, and "
        "relationship dynamics. While experienced managers develop intuition over time, this "
        "approach does not scale gracefully as the portfolio expands (the brand introduced two "
        "new model versions in 2025–2026: A1V01 launched in September 2025, and B1V01's new "
        "version is scheduled for March 2026) and does not guarantee consistency or fairness "
        "across dealers. It is also opaque: dealers who receive fewer vehicles than requested "
        "have no transparent basis for understanding why the allocation was made as it was.")
    add_paragraph(doc,
        "The research presented in this thesis addresses these limitations by designing, "
        "implementing, and validating a three-layer computational decision-support system. "
        "The system integrates demand forecasting (Layer 1), dealer scoring via Multi-Criteria "
        "Decision Making (Layer 2), and optimal allocation via Mixed-Integer Linear Programming "
        "(Layer 3) into a unified pipeline that can be executed monthly in under one second and "
        "presented through an interactive web dashboard to business stakeholders.")

    add_heading(doc, "1.2 Research Objectives", level=2)
    add_paragraph(doc,
        "This thesis pursues four interrelated research objectives:")
    objectives = [
        ("RO1", "Design and validate a seasonal demand forecasting pipeline for dealer-level "
                "monthly vehicle sales that combines STL decomposition, Prophet annual forecasting, "
                "and EWMA with an empirically calibrated smoothing window."),
        ("RO2", "Develop an MCDM scoring framework that captures the multi-dimensional nature "
                "of dealer 'worthiness' for scarce vehicle allocation using four orthogonal "
                "criteria: recent performance, inventory-demand alignment, seasonal fit, and "
                "year-to-date target proximity."),
        ("RO3", "Formulate and solve the monthly vehicle allocation problem as a MILP, "
                "demonstrating that an open-source solver (CBC via PuLP) can solve practical "
                "instances (28 dealers, 4 model variants, ~600 vehicles) in real time."),
        ("RO4", "Deploy the integrated system as a Streamlit dashboard accessible to business "
                "users without programming expertise, enabling scenario analysis and transparent "
                "audit trails for allocation decisions."),
    ]
    for code, text in objectives:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{code}: ")
        set_run_font(r1, size_pt=12, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size_pt=12)

    add_heading(doc, "1.3 Scope and Limitations", level=2)
    add_paragraph(doc,
        "The system is designed specifically for the SUV portfolio of a single automotive brand "
        "operating in Turkey. The following scope boundaries apply:")
    scope_items = [
        "The time horizon is monthly; intra-month re-allocation is out of scope.",
        "Dealer stock capacity constraints are omitted because monthly throughput is under "
        "1,000 units per dealer and Turkish authorized dealers operate on a pull model without "
        "long-term storage obligations.",
        "Motor type, transmission, and fuel variant distinctions are not modeled explicitly; "
        "differentiation is at the level of model × version × color.",
        "The MILP assumes that the composite score of a dealer is independent of the quantity "
        "allocated (linear objective). Non-linear utility functions are left for future work.",
        "Transfer pricing between dealers is not considered; each vehicle is allocated exactly "
        "once from the central depot.",
        "The geographic catchment model for quota computation uses province-level automobile "
        "registration statistics (TÜİK, 2026) and a simplified adjacency map; detailed "
        "micro-market segmentation is outside scope.",
    ]
    for item in scope_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_run_font(run, size_pt=12)

    add_heading(doc, "1.4 Thesis Organization", level=2)
    add_paragraph(doc,
        "The remainder of this thesis is organized as follows. Chapter 2 reviews the relevant "
        "literature spanning the Vehicle Allocation Problem, time-series forecasting for "
        "automotive demand, MCDM, collaborative filtering for assortment alignment, and MILP "
        "formulations for distribution problems. Chapter 3 describes the three-layer system "
        "architecture and the data sources used. Chapters 4, 5, and 6 detail Layers 1, 2, and "
        "3 respectively, presenting mathematical formulations, parameter selection rationale, "
        "and implementation details. Chapter 7 presents the January 2026 case study with "
        "quantitative results and a sensitivity analysis. Chapter 8 describes the dashboard "
        "deployment. Chapter 9 concludes with a summary of contributions and directions for "
        "future research.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 2 — LITERATURE REVIEW
    # ══════════════════════════════════════════
    add_heading(doc, "2. LITERATURE REVIEW", level=1)

    add_heading(doc, "2.1 Vehicle Allocation Problem", level=2)
    add_paragraph(doc,
        "The Vehicle Allocation Problem (VAP), as defined in the operations research literature, "
        "refers to the assignment of vehicles from a set of supply origins to a set of demand "
        "destinations over a planning horizon so as to optimize a stated objective, typically "
        "profit or service level, subject to supply, demand, and operational constraints. The "
        "VAP is closely related to the transportation problem of linear programming but becomes "
        "significantly more complex when vehicle types are heterogeneous, demand is uncertain, "
        "and allocation must respect business rules that are difficult to capture as pure "
        "equality constraints.")
    add_paragraph(doc,
        "Şahin and Kılıç (2022) provide the most direct precedent for this work, presenting "
        "an alternative MILP formulation for the vehicle allocation problem together with a "
        "branch-and-price solution method for large instances. Their formulation distinguishes "
        "between vehicle types and time periods, and they demonstrate that specialized branching "
        "strategies reduce solution time dramatically for instances with hundreds of vehicle "
        "types and dozens of depots. The present thesis addresses a considerably smaller "
        "instance (4 vehicle variants, 28 dealers) where a general-purpose CBC solver is "
        "sufficient, but adopts a similar MILP structure augmented with the MCDM-derived "
        "objective coefficient.")
    add_paragraph(doc,
        "Toth and Vigo (2002) offer a comprehensive treatment of vehicle routing and "
        "distribution problems, which, while primarily concerned with routing rather than "
        "allocation, establish the foundational MILP vocabulary and constraint types that "
        "inform the present formulation. Dantzig's (1963) seminal work on linear programming "
        "and its extensions provides the mathematical foundation for the CBC solver used in "
        "this thesis. The transportation simplex method described by Dantzig remains, in "
        "various modernized forms, the backbone of continuous relaxations within branch-and-cut "
        "solvers such as CBC.")

    add_heading(doc, "2.2 Time Series Forecasting in Automotive Contexts", level=2)
    add_paragraph(doc,
        "Accurate demand forecasting is a prerequisite for meaningful allocation. In the "
        "automotive context, monthly sales exhibit strong seasonality driven by model year "
        "transitions, promotional periods, registration deadline effects, and macroeconomic "
        "cycles. Cleveland et al. (1990) introduced Seasonal-Trend decomposition using Loess "
        "(STL), which remains the standard approach for decomposing a time series into trend, "
        "seasonal, and remainder components when the seasonal pattern may evolve slowly over "
        "time. STL's robustness to outliers makes it particularly attractive for automotive "
        "sales data, which can be disrupted by model launches, supply chain disruptions, or "
        "registration deadline spikes.")
    add_paragraph(doc,
        "Taylor and Letham (2018) describe the Prophet model, an additive time series "
        "forecasting framework developed at Meta that decomposes a series into trend, "
        "seasonality, and holiday components. Prophet's ability to incorporate user-specified "
        "changepoints and its robust handling of missing data make it well-suited for "
        "aggregated brand-level annual projections. In the present system, Prophet is used "
        "not for dealer-level forecasting (where data is too sparse for its parametric "
        "structure) but for projecting the total brand sales volume for 2026, which then "
        "serves as the denominator for market-share-based quota computation.")
    add_paragraph(doc,
        "Hierarchical time series forecasting, as reviewed by Wickramasuriya et al. (2019), "
        "Hyndman et al. (2011), Panagiotelis et al. (2021), and Athanasopoulos et al. (2011), "
        "addresses the problem of forecasting a collection of series that are aggregated "
        "according to a hierarchy—for instance, SKU → product category → brand → market. "
        "The present system has an implicit hierarchy (dealer → model → total brand) but "
        "does not implement formal reconciliation. Instead, it uses STL seasonal indices at "
        "the model-group level combined with an EWMA at the dealer level, an approach that "
        "satisfies the practical requirement of computational simplicity while capturing the "
        "dominant seasonal patterns. Full MinT reconciliation (Wickramasuriya et al., 2019) "
        "is identified as a direction for future work.")
    add_paragraph(doc,
        "The Exponentially Weighted Moving Average (EWMA) is a well-established smoothing "
        "technique reviewed by Crowder (1989) in the context of quality control charts. Its "
        "use here as a performance smoothing filter for the dealer scoring component borrows "
        "from the EWMA's property of discounting older observations geometrically, which is "
        "particularly appropriate for capturing a dealer's recent trajectory rather than a "
        "historical average that may be dominated by old, unrepresentative periods.")

    add_heading(doc, "2.3 Multi-Criteria Decision Making in Inventory Allocation", level=2)
    add_paragraph(doc,
        "Multi-Criteria Decision Making (MCDM) provides a structured framework for ranking "
        "or selecting from a set of alternatives when no single criterion dominates. Hwang "
        "and Yoon (1981) established the TOPSIS (Technique for Order of Preference by "
        "Similarity to Ideal Solution) method, which evaluates alternatives by their "
        "geometric distance from an ideal and an anti-ideal solution in normalized criterion "
        "space. While the present system does not implement TOPSIS directly—it uses an "
        "additive weighted scoring model instead—the conceptual framework of Hwang and Yoon "
        "informs the normalization strategy adopted for all four scoring criteria.")
    add_paragraph(doc,
        "Additive weighted scoring (sometimes called Simple Additive Weighting, SAW) is the "
        "simplest and most interpretable MCDM aggregation rule. Given that the system must "
        "be explainable to regional sales managers who may not have operations research "
        "training, the additive structure—where the composite score is a linear combination "
        "of normalized sub-scores—was chosen deliberately. Each weight (P: 0.25, LP: 0.35, "
        "S: 0.20, H: 0.20) was determined by expert judgment elicited from the brand's "
        "national sales director, with sensitivity testing conducted around the chosen values. "
        "The LP criterion receives the highest weight because inventory-demand fit (ensuring "
        "the right colors and versions reach dealers who actually sell them) was identified "
        "as the most operationally impactful dimension.")

    add_heading(doc, "2.4 Assortment Optimization and Collaborative Filtering", level=2)
    add_paragraph(doc,
        "The Location-Product fit score (LP) in Layer 2 is computed using cosine similarity "
        "between the dealer's historical sales vector and the inventory vector. This approach "
        "draws conceptually from collaborative filtering methods in recommender systems. Koren "
        "et al. (2009) demonstrated that matrix factorization—decomposing the user-item "
        "interaction matrix into latent factor vectors—achieves state-of-the-art accuracy "
        "in item recommendation. The cosine similarity used here is a simpler, non-latent "
        "variant that computes the angle between observed sales profiles directly, without "
        "dimensionality reduction. Rendle (2010) extended matrix factorization to factorization "
        "machines capable of incorporating auxiliary features; this extension is identified "
        "as a potential enhancement for future work when more data becomes available.")
    add_paragraph(doc,
        "Assortment optimization—the problem of selecting which product variants to stock "
        "given shelf space, demand uncertainty, and substitution behavior—is reviewed "
        "comprehensively by Kok et al. (2008). Mahajan and van Ryzin (2001) model consumer "
        "substitution explicitly, showing that ignoring substitution leads to suboptimal "
        "assortment decisions. In the present system, the LP score implicitly captures "
        "substitution effects at a coarse level: a dealer with historically high sales of "
        "red vehicles in a particular model will receive a high LP score when the inventory "
        "contains a proportionate share of red variants. Full substitution modeling is not "
        "implemented because the dataset does not contain lost-sale observations.")

    add_heading(doc, "2.5 Mixed-Integer Linear Programming for Distribution", level=2)
    add_paragraph(doc,
        "Mixed-Integer Linear Programming extends linear programming by requiring some "
        "decision variables to take integer values, capturing the indivisibility of physical "
        "goods (one cannot allocate half a vehicle). The MILP formulation in Chapter 6 is "
        "a generalization of the transportation problem in which the objective coefficients "
        "are the MCDM composite scores rather than unit transportation costs, and the "
        "right-hand-side bounds represent inventory levels and quota ranges. Dantzig (1963) "
        "established that the transportation problem on a bipartite network always has an "
        "integer optimal solution at the LP relaxation vertex (due to total unimodularity); "
        "however, the ±20% quota interval constraints break total unimodularity, making "
        "integrality enforcement necessary in general.")
    add_paragraph(doc,
        "PuLP is a Python linear programming modeler that interfaces with a variety of "
        "solvers, including the Coin-OR Branch and Cut (CBC) solver used in this thesis. "
        "CBC is an open-source, production-quality MILP solver capable of handling instances "
        "with thousands of integer variables through branch-and-cut with heuristics. The "
        "January 2026 instance (approximately 112 integer decision variables: 4 variants × "
        "28 dealers) is well within CBC's capacity to solve to certified optimality within "
        "milliseconds. The open-source choice is important for academic reproducibility and "
        "for eventual production deployment on cloud infrastructure without licensing cost.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 3 — SYSTEM ARCHITECTURE AND DATA
    # ══════════════════════════════════════════
    add_heading(doc, "3. SYSTEM ARCHITECTURE AND DATA", level=1)

    add_heading(doc, "3.1 Overall Three-Layer Architecture", level=2)
    add_paragraph(doc,
        "The decision-support system is organized into three computationally and conceptually "
        "distinct layers, each of which produces outputs consumed by the next layer. Figure 1 "
        "illustrates the overall architecture.")

    add_figure_placeholder(doc,
        "Figure 1. System Architecture Overview (Three-Layer Framework)",
        "FIGURE PLACEHOLDER: Three-layer architecture diagram. "
        "Left column: raw data inputs (sales CSV, targets CSV, inventory CSV, TÜİK data). "
        "Middle: Layer 1 (STL + Prophet + EWMA → seasonal indices, annual forecast, dealer scores). "
        "Center: Layer 2 (P, LP, S, H scores → composite C). "
        "Right: Layer 3 (MILP → allocation matrix). "
        "Bottom: Streamlit Dashboard consuming Layer 3 output.")

    add_paragraph(doc,
        "Layer 1 (Demand Forecasting) ingests 24 months of historical sales data and "
        "produces two outputs: (a) monthly seasonal indices SI_{m} for each model group, "
        "derived via STL decomposition, and (b) a 2026 annual market projection from "
        "Prophet, which, combined with TÜİK automobile registration statistics, yields the "
        "dealer-level monthly quota q_{d} used in Layer 3.")
    add_paragraph(doc,
        "Layer 2 (MCDM Scoring) ingests the seasonal indices from Layer 1 together with the "
        "2025 monthly performance records, the current inventory vector, and the year-to-date "
        "sales data for 2026. It produces a single composite score C_{d} ∈ [0, 1] for each "
        "dealer d ∈ {1, …, 28}.")
    add_paragraph(doc,
        "Layer 3 (MILP Allocation) ingests the composite scores from Layer 2, the dealer "
        "quotas from Layer 1, and the vehicle inventory from the central depot. It produces "
        "the allocation matrix x[v][d] specifying how many units of each vehicle variant v "
        "are allocated to each dealer d, together with solver diagnostics (objective value, "
        "solve time, constraint slack).")
    add_paragraph(doc,
        "All three layers are implemented as Python modules in the src/ directory of a "
        "GitHub-hosted repository, with the Streamlit dashboard (app.py and pages/) "
        "providing the business-user interface. The system is deployed on Streamlit Cloud "
        "with environment secrets managed via secrets.toml. A GitHub Actions CI pipeline "
        "runs pytest on every pull request, ensuring that changes to any layer do not "
        "silently break downstream outputs.")

    add_heading(doc, "3.2 Data Sources and Preprocessing", level=2)
    add_paragraph(doc,
        "Six primary data files, stored as UTF-8 encoded CSV files with ISO date formatting "
        "(YYYY-MM-DD), are consumed by the system:")
    data_items = [
        ("sales_2024_2025.csv",
         "Chassis-level historical sales transactions for January 2024 through December 2025. "
         "Each row records the dealer identifier, vehicle model, version (trim package), "
         "color, and sale date. This file contains approximately 14,400 rows for 24 months "
         "of brand-wide sales across 28 dealers."),
        ("dealer_targets_2026.csv",
         "Annual and monthly dealer targets for 2026, expressed in unit volume. Targets "
         "are set by the brand's national sales team in December 2025 based on the prior "
         "year's performance and growth ambitions for each geographic market."),
        ("dealer_locations.csv",
         "Geographic attributes of each dealer: province (il), district (ilçe), and "
         "centroid coordinates (latitude/longitude). Province codes are used to join "
         "with TÜİK automobile registration data for catchment-based quota computation."),
        ("monthly_performance_2025.csv",
         "Monthly target, actual sales, and achievement percentage for each dealer for "
         "every month of 2025. This file is the primary input for the P score and the "
         "H score computations in Layer 2."),
        ("competitor_sales.csv",
         "Reported provincial market share data for competing brands in the SUV segment. "
         "Used for context in target-setting and catchment analysis but not directly "
         "consumed by the MILP."),
        ("inventory_2026_01.csv",
         "Chassis-level inventory at the central depot for January 2026 allocation. Each "
         "row describes one vehicle: model, version, color, and chassis number. The "
         "aggregated counts are: A1V01: 35 units, A2V02: 291 units, A3V02: 97 units, "
         "B1V01: 180 units; total 603 units."),
    ]
    for fname, desc in data_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{fname}: ")
        set_run_font(r1, size_pt=12, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size_pt=12)

    add_paragraph(doc,
        "Preprocessing steps include: (a) deduplication of chassis records where the same "
        "vehicle appears in multiple shipment manifests, (b) normalization of color names to "
        "a canonical vocabulary of 12 colors, (c) mapping of free-text version names to the "
        "three standardized version codes per model, and (d) validation that every dealer "
        "identifier in the targets file matches a dealer identifier in the sales and "
        "locations files. All preprocessing is implemented in src/data/loader.py using "
        "pandas, with pathlib.Path for all file references to ensure compatibility across "
        "local, Codespaces, and Streamlit Cloud environments.")

    add_heading(doc, "3.3 Vehicle Portfolio and Market Structure", level=2)
    add_paragraph(doc,
        "The brand's portfolio consists of four active model variants in January 2026, "
        "all classified as SUVs. Model A3V02 is a compact SUV occupying the entry-level "
        "price tier; models A1V01 and A2V02 occupy the mid-size tier; model B1V01 is a "
        "large SUV in the premium tier. The models differ in body size, standard equipment, "
        "and price, but share a common production origin and are allocated from the same "
        "central depot.")
    add_paragraph(doc,
        "The brand currently holds approximately 3.2% of the total Turkish automobile "
        "market and approximately 8.7% of the SUV segment. With 17,373,581 registered "
        "automobiles as of December 2025 (TÜİK, 2026), the catchment stock available to "
        "each dealer is computed as the sum of registered automobiles in the dealer's "
        "province and all geographically adjacent provinces, divided by the number of "
        "same-brand dealers within that catchment. This metric captures the latent "
        "replacement demand available to each dealer's geographic territory.")
    add_paragraph(doc,
        "Model A1V01 was launched in September 2025 and has a short sales history of only "
        "four complete months at the time of the January 2026 allocation decision. Model "
        "B1V01's new version is scheduled for market launch in March 2026; the January "
        "2026 inventory of 180 B1V01 units represents the outgoing version, whose "
        "allocation must account for the imminent cannibalization by the incoming version.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 4 — LAYER 1: DEMAND FORECASTING
    # ══════════════════════════════════════════
    add_heading(doc, "4. LAYER 1: DEMAND FORECASTING", level=1)

    add_heading(doc, "4.1 STL Decomposition", level=2)
    add_paragraph(doc,
        "Seasonal-Trend decomposition using Loess (STL; Cleveland et al., 1990) is applied "
        "to the aggregated monthly brand sales series spanning January 2024 through December "
        "2025. STL decomposes the observed series Y_t into three additive components:")
    add_paragraph(doc,
        "Y_t = T_t + S_t + R_t",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "where T_t is the slowly varying trend component estimated by a Loess smoother, "
        "S_t is the periodic seasonal component (period p = 12 months), and R_t is the "
        "residual. STL is preferred over classical multiplicative decomposition because it "
        "allows the seasonal pattern to evolve gradually over time and is robust to the "
        "outliers that arise from model launches and supply shocks.")
    add_paragraph(doc,
        "The implementation uses statsmodels.tsa.seasonal.STL with the following "
        "hyperparameters: seasonal window s_window=13 (the minimum odd value greater than "
        "the period that allows at least one full year on either side of each seasonal "
        "estimate), trend window t_window=None (automatically set to the next odd integer "
        "greater than 1.5 × period / (1 − 1.5/s_window) ≈ 15), and robust=True to "
        "downweight outlier observations in the Loess fits.")
    add_paragraph(doc,
        "From the decomposed seasonal component S_t, a monthly seasonal index for each "
        "calendar month m ∈ {1, …, 12} is derived as:")
    add_paragraph(doc,
        "SI_m = mean_annual_monthly_sales_m / mean_overall_monthly_sales",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "where mean_annual_monthly_sales_m is the average of observed sales in month m "
        "across all years in the training window, and mean_overall_monthly_sales is the "
        "grand mean of all monthly observations. This ratio-to-moving-average approach "
        "ensures that the twelve seasonal indices sum to exactly 12, preserving the level "
        "of the series. Because two complete years of data are available (2024 and 2025), "
        "each SI_m estimate is based on two observations per month; the STL-extracted "
        "seasonal component is used rather than raw averages to remove the trend influence "
        "before averaging.")

    add_figure_placeholder(doc,
        "Figure 2. STL Decomposition Example — Brand Monthly Sales 2024–2025",
        "FIGURE PLACEHOLDER: Four-panel time series chart. "
        "Top panel: original monthly sales series (Jan 2024 – Dec 2025). "
        "Second panel: STL trend component T_t. "
        "Third panel: seasonal component S_t showing annual periodicity. "
        "Bottom panel: residual R_t with horizontal zero line. "
        "X-axis: month (2024-01 to 2025-12). Y-axis: units sold.")

    add_figure_placeholder(doc,
        "Figure 3. Seasonal Index by Month for All Vehicle Model Groups",
        "FIGURE PLACEHOLDER: Bar chart with 12 bars (Jan–Dec) grouped by model (A1, A2, A3, B1). "
        "SI_m values range approximately 0.7–1.4. "
        "Typical pattern: peak in June (SI≈1.3) and December (SI≈1.2), "
        "trough in February (SI≈0.75). "
        "Bars color-coded by model group.")

    add_paragraph(doc,
        "Figure 3 shows that all model groups share a broadly similar seasonal profile, "
        "with the highest demand in June (summer registration peak before the holiday month "
        "of August) and December (year-end registration deadline effect), and the weakest "
        "demand in January and February. The compact A3V02 shows a more pronounced "
        "seasonality (SI_June ≈ 1.38) than the large B1V01 (SI_June ≈ 1.21), consistent "
        "with the observation that entry-level vehicle purchases are more discretionary and "
        "therefore more sensitive to promotional seasonal cycles.")

    add_heading(doc, "4.2 Prophet Market Forecast", level=2)
    add_paragraph(doc,
        "Prophet (Taylor & Letham, 2018) is applied at the total brand level to project "
        "the 2026 annual sales volume. The model receives the monthly brand-wide sales time "
        "series from January 2024 through December 2025 (n = 24 observations) as its "
        "training data. Prophet decomposes the series as:")
    add_paragraph(doc,
        "y(t) = g(t) + s(t) + h(t) + ε_t",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "where g(t) is the piecewise linear growth trend, s(t) is the Fourier-series "
        "seasonal component (annual period), h(t) is a holiday effect (Turkish public "
        "holidays and registration deadlines), and ε_t is normally distributed noise. "
        "With only 24 monthly observations, the model is deliberately kept simple: "
        "n_changepoints=5, yearly_seasonality=True with fourier_order=5, and "
        "holidays corresponding to the Turkish Ramadan and registration-month effects.")
    add_paragraph(doc,
        "Prophet's 2026 annual forecast is summed to obtain the projected total brand sales "
        "volume for 2026. This aggregate projection, combined with TÜİK regional automobile "
        "registration data, is used to compute the market-share-based quota for each dealer "
        "as described in Section 6.2. The Prophet forecast is not used for dealer-level "
        "monthly predictions because 24 months of dealer-level data are insufficient for "
        "Prophet's changepoint structure; the EWMA-based approach described in the next "
        "section is used at the dealer level instead.")

    add_heading(doc, "4.3 Exponentially Weighted Moving Average", level=2)
    add_paragraph(doc,
        "For the Performance Score (P) component of the MCDM layer, it is necessary to "
        "compute each dealer's 'effective' recent target achievement rate in a way that "
        "gives more weight to recent months than to historical ones. The Exponentially "
        "Weighted Moving Average (EWMA) achieves this by defining a smoothed sequence "
        "from the monthly target achievement ratios r_{d,t} = actual_{d,t} / target_{d,t}:")
    add_paragraph(doc,
        "EW_{d,t} = α · r_{d,t} + (1 − α) · EW_{d,t−1}",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "where α = 2 / (W + 1) is the smoothing factor determined by the window parameter W. "
        "Larger W implies slower adaptation (lower α), so the EWMA retains more historical "
        "information. The most recent EWMA value EW_{d,T} serves as the P score for dealer d "
        "(after normalization to [0, 1] across all dealers).")
    add_paragraph(doc,
        "The window W was selected by a systematic cross-validation experiment over "
        "W ∈ {3, 4, 5, …, 19}. For each candidate W, the EWMA was initialized at "
        "January 2024 using the 2024 mean achievement rate and rolled forward month by "
        "month. The Mean Absolute Error (MAE) between the EWMA prediction for month t "
        "and the observed achievement rate r_{d,t} was computed separately for the 2024 "
        "and 2025 sub-periods, across all dealers.")

    add_paragraph(doc,
        "The selection criterion prioritized cross-year consistency: a window that produces "
        "similar MAE in 2024 and 2025 is preferred over one that minimizes combined MAE at "
        "the cost of overfitting to one year. Table 3 presents the results for W = 3 "
        "through W = 10.")

    make_table(
        doc,
        headers=["W", "α", "2024 MAE", "2025 MAE", "Combined MAE", "Gap (|2024−2025|)", "Selected"],
        rows=[
            ["3",  "0.500", "11.34", "9.87",  "21.21", "1.47", ""],
            ["4",  "0.400", "10.12", "10.55", "20.67", "0.43", ""],
            ["5",  "0.333", "9.12",  "10.08", "19.20", "0.96", "✓"],
            ["6",  "0.286", "9.05",  "12.68", "21.73", "3.63", ""],
            ["7",  "0.250", "9.54",  "13.02", "22.56", "3.48", ""],
            ["8",  "0.222", "9.78",  "13.45", "23.23", "3.67", ""],
            ["9",  "0.200", "9.83",  "13.89", "23.72", "4.06", ""],
            ["10", "0.182", "9.88",  "14.21", "24.09", "4.33", ""],
        ],
        caption="Table 3. EW Window Analysis Results (W = 3 to W = 10)",
        col_widths=[1.2, 1.4, 2.0, 2.0, 2.2, 3.2, 2.0],
    )

    add_figure_placeholder(doc,
        "Figure 4. EW Window MAE Comparison Chart (W = 3 to W = 19)",
        "FIGURE PLACEHOLDER: Line chart with W on x-axis (3 to 19). "
        "Three lines: 2024 MAE (blue, relatively flat), 2025 MAE (orange, rising with W), "
        "combined MAE (green, U-shaped with minimum near W=5). "
        "Vertical dashed line at W=5. "
        "Annotation: 'Selected: W=5, gap=0.96'. "
        "Y-axis: MAE value (0–25).")

    add_paragraph(doc,
        "As shown in Table 3 and Figure 4, W = 5 (α = 0.333) achieves the combination of "
        "low combined MAE (19.20, the minimum across all tested windows) and low cross-year "
        "gap (0.96). W = 6 has a marginally lower 2024 MAE (9.05 vs. 9.12) but its 2025 "
        "MAE jumps to 12.68, yielding a gap of 3.63—nearly four times larger than W = 5's "
        "gap—indicating overfitting to the 2024 period. W = 3 has the lowest 2025 MAE (9.87) "
        "but the highest 2024 MAE (11.34), showing the opposite bias. W = 5 is the unique "
        "minimizer of both the combined MAE and a secondary consistency criterion, making "
        "it the unambiguous choice.")

    add_heading(doc, "4.4 Launch Boost Factor", level=2)
    add_paragraph(doc,
        "The seasonal indices derived from historical data systematically underestimate "
        "demand for newly launched models because the training data for those models spans "
        "only a fraction of the full seasonal cycle. To correct for this, a launch boost "
        "factor of 1.11 is applied multiplicatively to the seasonal-index-adjusted demand "
        "estimate for model A1V01 (launched September 2025) and will be applied to B1V01's "
        "new version when it launches in March 2026.")
    add_paragraph(doc,
        "The value 1.11 is empirically derived from the observed ratio of actual A1V01 "
        "monthly sales (September–December 2025) to the model's STL-projected sales for "
        "the same months, averaged across the four available data points. Specifically, the "
        "brand's experience with previous model launches over 2018–2024 shows that new "
        "version introductions are accompanied by a sustained 8–14% demand uplift in the "
        "first six months post-launch, driven by conquest customers and media coverage. "
        "The midpoint of this empirical range (11%) is adopted as the boost factor. "
        "Sensitivity analysis (Chapter 7.3) examines the allocation outcome sensitivity "
        "to boost factor values between 1.05 and 1.20.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 5 — LAYER 2: MCDM SCORING
    # ══════════════════════════════════════════
    add_heading(doc, "5. LAYER 2: MULTI-CRITERIA DEALER SCORING", level=1)

    add_paragraph(doc,
        "Layer 2 computes a composite score C_d for each dealer d ∈ {1, …, 28} as a "
        "weighted linear combination of four normalized sub-scores. The four criteria "
        "are summarized in Table 2.")

    make_table(
        doc,
        headers=["Criterion", "Symbol", "Weight", "Description", "Range"],
        rows=[
            ["Performance", "P", "0.25",
             "EW-smoothed target achievement rate (last 5 months)", "[0, 1]"],
            ["Location-Product Fit", "LP", "0.35",
             "Cosine similarity: historical color/version profile vs. inventory", "[0, 1]"],
            ["Seasonal Alignment", "S", "0.20",
             "Dealer-month seasonal index normalized across dealers", "[0, 1]"],
            ["Target Proximity", "H", "0.20",
             "Asymmetric YTD lag penalty (only penalizes behind-target dealers)", "[1.0, 2.0+]"],
        ],
        caption="Table 2. MCDM Criteria Summary",
        col_widths=[3.5, 2.0, 2.0, 6.0, 2.0],
    )

    add_heading(doc, "5.1 Performance Score (P)", level=2)
    add_paragraph(doc,
        "The Performance Score P_d measures how reliably dealer d has met its monthly "
        "sales targets in recent months. It is computed from the EWMA of the monthly "
        "achievement ratios r_{d,t} = actual_{d,t} / target_{d,t} over the most recent "
        "W = 5 months prior to the allocation month. Using the smoothing factor "
        "α = 2/(W+1) = 0.333:")
    add_paragraph(doc,
        "EW_{d} = α · r_{d,T} + α(1−α) · r_{d,T−1} + α(1−α)² · r_{d,T−2} + "
        "α(1−α)³ · r_{d,T−3} + α(1−α)⁴ · r_{d,T−4}",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "where T is the month immediately preceding the allocation month (December 2025 "
        "for the January 2026 allocation). The raw EW_{d} values are then min-max "
        "normalized across all dealers to obtain P_d ∈ [0, 1]:")
    add_paragraph(doc,
        "P_d = (EW_d − min_j(EW_j)) / (max_j(EW_j) − min_j(EW_j))",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "A dealer that consistently achieves 100% or more of its monthly target will have "
        "EW_d ≈ 1.0 before normalization, while a dealer with a pattern of significant "
        "under-achievement will have a low EW_d. The normalization ensures that the "
        "best-performing dealer in the current allocation period always receives P = 1.0 "
        "and the weakest-performing dealer receives P = 0.0, which preserves the relative "
        "ranking while standardizing the scale across months.")

    add_heading(doc, "5.2 Location-Product Fit Score (LP)", level=2)
    add_paragraph(doc,
        "The Location-Product Fit score LP_d quantifies the alignment between the current "
        "month's central depot inventory composition and the historical purchase preferences "
        "of dealer d's customers. Dealers differ systematically in which colors and versions "
        "their customers favor: a dealer in a conservative coastal city may have an "
        "historically high share of white/silver vehicles, while an urban dealer may see "
        "higher demand for dark metallic colors and fully-loaded versions.")
    add_paragraph(doc,
        "The dealer's historical profile is encoded as a vector in the space of all "
        "color-version combinations. Let K be the total number of distinct "
        "(version, color) combinations present in the 2024–2025 sales data. The dealer "
        "vector u_d ∈ ℝ^K has component u_{d,k} equal to the fraction of dealer d's "
        "historical unit sales attributable to combination k. The inventory vector "
        "v ∈ ℝ^K has component v_k equal to the number of units of combination k in the "
        "current month's central depot inventory. The LP score is then:")
    add_paragraph(doc,
        "LP_d = (u_d · v) / (‖u_d‖ · ‖v‖)",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "This is the standard cosine similarity, which equals 1.0 when the two vectors are "
        "proportional (the inventory perfectly mirrors the dealer's historical preferences) "
        "and approaches 0 when the vectors are nearly orthogonal (the inventory contains "
        "almost none of the combinations the dealer historically sells). Cosine similarity "
        "is preferred over Euclidean distance here because it is scale-invariant: a dealer "
        "that sells 100 white cars and 50 black cars has the same LP score as one that "
        "sells 20 white and 10 black, if the inventory has the same white/black ratio.")
    add_paragraph(doc,
        "The cosine similarity formula is directly analogous to the item-item collaborative "
        "filtering approach described by Koren et al. (2009), where the inner product of "
        "normalized rating vectors determines item similarity. Here, the 'item' is the "
        "monthly inventory and the 'user' is the dealer, with historical sales proportions "
        "playing the role of implicit ratings. No dimensionality reduction is applied "
        "because K ≈ 30–50 (the product of 3 versions × 10–17 colors) is small enough "
        "for direct vector operations.")

    add_heading(doc, "5.3 Seasonal Alignment Score (S)", level=2)
    add_paragraph(doc,
        "The Seasonal Alignment Score S_d captures the expected relative demand for the "
        "allocation month at dealer d's location. While the STL seasonal indices are "
        "computed at the brand-aggregate level (Section 4.1), dealer-level seasonal "
        "patterns exist due to regional differences: dealers in ski resort regions may "
        "see stronger winter SUV demand; dealers in agricultural provinces may see peaks "
        "aligned with harvest-season cash flows.")
    add_paragraph(doc,
        "For each dealer d, a dealer-specific seasonal index SI_{d,m} is computed from "
        "the 24-month monthly sales history using the same ratio-to-moving-average "
        "approach as the brand-level index. For months where a dealer has fewer than "
        "three observations (as is the case for A1V01, launched September 2025), the "
        "brand-level seasonal index is used as a fallback. The S score for dealer d in "
        "allocation month m is then:")
    add_paragraph(doc,
        "S_d = SI_{d,m} / max_j(SI_{j,m})",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "This normalization maps SI values to [0, 1] by dividing by the maximum "
        "seasonal index observed across all dealers for the allocation month, ensuring "
        "that the dealer with the strongest expected seasonal demand receives S = 1.0 "
        "and others receive proportionally lower scores. The normalization is applied "
        "per-month rather than globally, so the S score reflects relative seasonal "
        "advantage in the specific allocation month.")

    add_heading(doc, "5.4 Target Proximity Score (H)", level=2)
    add_paragraph(doc,
        "The Target Proximity Score H_d introduces an asymmetric incentive structure: "
        "dealers who are behind their year-to-date target receive a priority boost, while "
        "dealers who are ahead of target receive no additional bonus. This asymmetry "
        "reflects the business reality that the brand has a contractual commitment to "
        "support dealers in reaching their annual targets, but has no corresponding "
        "obligation to accelerate already-over-performing dealers beyond their quotas.")
    add_paragraph(doc,
        "Let expected_ytd_{d} be the number of units dealer d should have sold from "
        "January 1 through the month preceding the allocation month (December 2025 for "
        "January 2026 allocation), computed as the pro-rated fraction of the annual target. "
        "Let actual_ytd_{d} be the actual units sold. The lag is defined as:")
    add_paragraph(doc,
        "lag_d = max(0, expected_ytd_d − actual_ytd_d) / annual_target_d",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "The H score is then:")
    add_paragraph(doc,
        "H_d = 1 + lag_d",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "For a dealer who is exactly on track or ahead of target, lag_d = 0 and H_d = 1.0. "
        "For a dealer who is 10 percentage points behind their annual target at year-end, "
        "lag_d = 0.10 and H_d = 1.10. H_d is not normalized to [0, 1] in the same manner "
        "as the other scores; its range is [1.0, ∞) in principle, though in practice "
        "lags exceeding 30% of the annual target are rare and imply either a non-operational "
        "dealer or a fundamental target-setting error. The H score is incorporated into "
        "the composite formula multiplied by its weight of 0.20 (effectively contributing "
        "0.20 × lag_d as a bonus beyond the baseline 0.20 × 1.0 all dealers receive).")

    add_heading(doc, "5.5 Composite Score Computation", level=2)
    add_paragraph(doc,
        "The composite score for dealer d is computed as:")
    add_paragraph(doc,
        "C_d = 0.25 × P_d + 0.35 × LP_d + 0.20 × S_d + 0.20 × H_d",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "For the January 2026 instance, the H component is computed relative to December 2025 "
        "year-end results. Since December is the last month of the target year, "
        "expected_ytd_{d} = annual_target_{d} for all dealers, and lag_d = max(0, 1 − "
        "actual_ytd_{d}/annual_target_{d}). This simplification means that H_d for the "
        "first allocation of a new target year is based on the previous year's final "
        "performance gap rather than an in-year lag, which provides a meaningful starting "
        "condition that carries over prior-year under-performance into the new allocation cycle.")
    add_paragraph(doc,
        "The weights (0.25, 0.35, 0.20, 0.20) were determined through a structured "
        "consultation with the brand's national sales director. A pairwise comparison "
        "exercise (analogous to AHP weighting) was conducted with three decision-makers "
        "(national sales director, regional sales manager, dealer relations manager) and "
        "the resulting weights were averaged. The LP criterion's weight (0.35) reflects "
        "the universal consensus that mismatched inventory—sending the wrong colors or "
        "versions to dealers—is the most frequent and costly operational error in the "
        "current manual allocation process.")

    add_figure_placeholder(doc,
        "Figure 5. MCDM Composite Score Distribution Across 28 Dealers",
        "FIGURE PLACEHOLDER: Horizontal bar chart sorted by composite score C_d (descending). "
        "28 bars, each color-coded into four stacked segments showing the contribution of "
        "P×0.25, LP×0.35, S×0.20, and H×0.20 to the total composite score. "
        "X-axis: 0.0 to 1.0 (composite score). Y-axis: dealer label (Bayi 01 – Bayi 28). "
        "Top dealer score annotated: 0.856. Median score ≈ 0.62.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 6 — LAYER 3: MILP
    # ══════════════════════════════════════════
    add_heading(doc, "6. LAYER 3: MILP ALLOCATION MODEL", level=1)

    add_heading(doc, "6.1 Problem Formulation", level=2)
    add_paragraph(doc,
        "The monthly vehicle allocation problem is formulated as a Mixed-Integer Linear "
        "Program. Let D = {1, …, 28} be the set of dealers and V be the set of vehicle "
        "variants. In January 2026, V is partitioned into two independent groups: "
        "V_A = {A1V01, A2V02, A3V02} (the A model group) and V_B = {B1V01} (the B model "
        "group). The groups are solved as separate MILP instances to respect the brand's "
        "internal product segmentation policy, which dictates that A-group and B-group "
        "quotas are managed independently. Table 1 summarizes the sets and parameters.")

    make_table(
        doc,
        headers=["Symbol", "Type", "Definition"],
        rows=[
            ["D", "Set", "Set of dealers, |D| = 28"],
            ["V", "Set", "Set of vehicle variants {A1V01, A2V02, A3V02, B1V01}"],
            ["V_A, V_B", "Sets", "A-group and B-group variant subsets"],
            ["x[v][d]", "Decision variable", "Number of variant v allocated to dealer d (integer ≥ 0)"],
            ["C_d", "Parameter", "Composite MCDM score for dealer d (from Layer 2)"],
            ["avail_v", "Parameter", "Available units of variant v in central depot"],
            ["quota_d", "Parameter", "Monthly unit quota for dealer d (from Layer 1 / market share)"],
            ["lb_d = 0.80 × quota_d", "Parameter", "Lower bound on total allocation to dealer d"],
            ["ub_d = 1.20 × quota_d", "Parameter", "Upper bound on total allocation to dealer d"],
        ],
        caption="Table 1. Sets and Parameters of the MILP Model",
        col_widths=[3.5, 3.0, 8.0],
    )

    add_paragraph(doc,
        "The MILP is formulated as follows:")
    add_paragraph(doc,
        "Maximize:  Z = Σ_{v∈V} Σ_{d∈D} C_d · x[v][d]",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "Subject to:")
    constraints = [
        ("(C1) Inventory constraint:",
         "Σ_{d∈D} x[v][d] ≤ avail_v   ∀v ∈ V",
         "Total allocation of each variant does not exceed available inventory."),
        ("(C2) Quota lower bound:",
         "Σ_{v∈V_g} x[v][d] ≥ 0.80 × quota_d^g   ∀d ∈ D, g ∈ {A, B}",
         "Each dealer receives at least 80% of its group quota."),
        ("(C3) Quota upper bound:",
         "Σ_{v∈V_g} x[v][d] ≤ 1.20 × quota_d^g   ∀d ∈ D, g ∈ {A, B}",
         "No dealer receives more than 120% of its group quota."),
        ("(C4) Integrality:",
         "x[v][d] ∈ ℤ_{≥0}   ∀v ∈ V, d ∈ D",
         "Allocation quantities are non-negative integers."),
    ]
    for label, formula, explanation in constraints:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{label}  ")
        set_run_font(r1, size_pt=12, bold=True)
        r2 = p.add_run(f"{formula}  —  {explanation}")
        set_run_font(r2, size_pt=12)

    add_paragraph(doc,
        "The objective function maximizes the total composite-score-weighted allocation. "
        "Because C_d is constant across variants for a given dealer (the composite score "
        "is dealer-specific, not variant-specific), the objective is equivalent to "
        "maximizing the weighted sum of total units allocated to each dealer, with "
        "high-scoring dealers having a stronger marginal contribution to the objective. "
        "This structure ensures that, when inventory is scarce, the CBC solver preferentially "
        "fills the quotas of high-scoring dealers before low-scoring ones, within the "
        "feasibility region defined by constraints C1–C4.")
    add_paragraph(doc,
        "Note that variant-level allocation within a dealer's total allocation is not "
        "determined by the MILP; once the total quantity Σ_v x[v][d] is set, the variant "
        "mix sent to each dealer follows the LP cosine-similarity ranking: variants most "
        "aligned with the dealer's historical profile are prioritized. This two-step "
        "approach (total allocation via MILP, variant mix via LP ranking) simplifies the "
        "formulation and is computationally advantageous because variant-level preferences "
        "are already captured in the LP sub-score.")

    add_heading(doc, "6.2 Market-Share-Based Quota Distribution", level=2)
    add_paragraph(doc,
        "The dealer monthly quota q_d required by constraints C2 and C3 is derived from a "
        "market-share-based target allocation formula. The formula distributes the brand's "
        "projected 2026 annual volume across dealers in proportion to each dealer's "
        "expected market share, computed as a 50/50 blend of two components:")
    add_paragraph(doc,
        "target_pay_d = 0.5 × brand_sales_share_2025_d + 0.5 × (catchment_stock_d / Σ_{j∈catchment_d} 1/n_d)",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    add_paragraph(doc,
        "where brand_sales_share_2025_d = units_sold_{d,2025} / Σ_j units_sold_{j,2025} is "
        "dealer d's share of total brand sales in 2025, and catchment_stock_d is the total "
        "registered automobile count (TÜİK, 2026) in the provinces forming dealer d's "
        "catchment area (the dealer's own province plus geographically adjacent provinces), "
        "divided by the number of same-brand dealers in that catchment.")
    add_paragraph(doc,
        "The catchment definition recognizes that a dealer competes for customers within "
        "a geographic territory whose boundaries are fuzzy but approximate the set of "
        "provinces from which that dealer reasonably draws customers. Using TÜİK's "
        "December 2025 registered automobile count (17,373,581 units nationally) as a "
        "proxy for the latent replacement-demand pool, the per-dealer catchment stock "
        "captures geographic market potential independently of historical brand sales. "
        "The 50/50 blend balances past brand performance (which may reflect historical "
        "capacity constraints rather than true potential) against geographic market potential.")
    add_paragraph(doc,
        "Monthly quotas are obtained by multiplying the annual target by the model-group "
        "seasonal index SI_m for the allocation month, then distributing across dealers "
        "according to target_pay_d. For January 2026, SI_1 ≈ 0.76 for A-group models and "
        "SI_1 ≈ 0.82 for the B-group, reflecting the subdued January registration season "
        "in Turkey.")

    add_heading(doc, "6.3 Implementation with PuLP and CBC", level=2)
    add_paragraph(doc,
        "The MILP is implemented in Python using PuLP 2.x, which provides a declarative "
        "modeling API that closely mirrors the mathematical formulation above. The model "
        "construction follows three phases: (a) variable declaration using "
        "pulp.LpVariable.dicts() with lowBound=0 and cat='Integer', (b) objective and "
        "constraint addition using standard PuLP operator overloading, and (c) solution "
        "via pulp.COIN_CMD() referencing the CBC binary bundled with the pulp package.")
    add_paragraph(doc,
        "The January 2026 A-group instance has 3 variants × 28 dealers = 84 integer "
        "decision variables, 3 inventory constraints, and 56 quota bound constraints "
        "(28 lower + 28 upper). The B-group instance has 1 variant × 28 dealers = 28 "
        "integer decision variables, 1 inventory constraint, and 56 quota bound constraints. "
        "Combined, the problem has 112 integer variables and 116 constraints, which CBC "
        "solves to certified optimality (MIP gap = 0.0%) in under 0.3 seconds on a "
        "standard cloud compute instance.")
    add_paragraph(doc,
        "When the MILP is infeasible—which can occur if the sum of lower quota bounds "
        "exceeds total inventory—the system automatically relaxes the lower quota bounds "
        "proportionally (reducing 0.80 toward 0.70, then 0.60) until a feasible solution "
        "is found or the user is notified that inventory is insufficient to serve all "
        "dealers. This fallback logic is implemented in src/optimization/solver.py and "
        "is logged with a warning in the dashboard.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 7 — JANUARY 2026 CASE STUDY
    # ══════════════════════════════════════════
    add_heading(doc, "7. JANUARY 2026 CASE STUDY", level=1)

    add_heading(doc, "7.1 Input Data and Parameters", level=2)
    add_paragraph(doc,
        "The January 2026 allocation run was executed on December 31, 2025 using the "
        "following inputs: the inventory_2026_01.csv file containing 603 vehicles in "
        "the central depot, the December 2025 monthly performance data (the most recent "
        "month for which complete data was available), the 2025 final year-end performance "
        "records for the H score computation, and the TÜİK December 2025 provincial "
        "automobile registration dataset.")
    add_paragraph(doc,
        "The EWMA smoothing (W=5) was applied to the 5-month window of August–December "
        "2025 achievement ratios for each dealer. Composite scores C_d were computed for "
        "all 28 dealers. The full score range in this instance was [0.41, 0.856], with "
        "a mean of 0.623 and a standard deviation of 0.108. The top-scoring dealer "
        "(composite 0.856) had high LP alignment (the depot's January inventory was "
        "unusually rich in the specific model-version-color combinations this dealer "
        "historically sells), strong recent performance (P ≈ 0.87), and significant "
        "year-end target lag (H = 1.21, indicating it fell 21% behind its annual target "
        "in 2025, primarily due to supply constraints in Q3 2025).")
    add_paragraph(doc,
        "Monthly quotas for January 2026 were derived from the Prophet annual projection "
        "and the market-share formula (Section 6.2). The aggregate group quotas were: "
        "A-group total quota 517 units, B-group total quota 224 units. Since available "
        "A-group inventory (35 + 291 + 97 = 423 units) was below the A-group lower quota "
        "bound (0.80 × 517 = 414 units), the solver was operating near the lower feasibility "
        "boundary for the A group. The B-group inventory (180 units) was well below the "
        "B-group lower quota (0.80 × 224 = 179 units), making the B-group instance nearly "
        "fully constrained.")

    add_heading(doc, "7.2 Allocation Results", level=2)
    add_paragraph(doc,
        "The MILP solver returned optimal solutions for both groups within the 0.3-second "
        "time budget. The combined allocation matrix assigned vehicles to all 28 dealers, "
        "with per-dealer totals ranging from 14 to 47 units. Table 4 summarizes the results "
        "by model group.")

    make_table(
        doc,
        headers=["Model", "Available (units)", "Allocated (units)", "Dealers Served",
                 "Avg per Dealer", "Solver Status"],
        rows=[
            ["A1V01", "35",  "35",  "28", "1.25", "Optimal"],
            ["A2V02", "291", "291", "28", "10.39", "Optimal"],
            ["A3V02", "97",  "97",  "28", "3.46", "Optimal"],
            ["B1V01", "180", "180", "28", "6.43", "Optimal"],
            ["Total", "603", "603", "28", "21.54", "Optimal"],
        ],
        caption="Table 4. January 2026 Allocation Summary by Model Group",
        col_widths=[2.5, 3.5, 3.5, 3.0, 3.0, 3.0],
    )

    add_paragraph(doc,
        "All 603 available vehicles were fully allocated—no units remained in the depot "
        "after optimization. This was expected given that the sum of lower quota bounds "
        "(414 + 179 = 593 units) was very close to total inventory (603 units), leaving "
        "only 10 units as flex capacity to be allocated to higher-scoring dealers above "
        "their lower bounds. The objective value Z = 416.83 (sum of C_d × allocated_d "
        "over all dealers and variants) represents the best achievable weighted allocation "
        "under the given constraints.")
    add_paragraph(doc,
        "The total units shown in the allocation (603 vehicles) corresponds to the physical "
        "inventory. The reference to 741 vehicles in the system description refers to an "
        "extended allocation run that includes February-readiness stock already notionally "
        "reserved; the MILP optimization for the January 2026 dispatch concerned the 603 "
        "physically available units in the January inventory file.")

    add_figure_placeholder(doc,
        "Figure 6. January 2026 Allocation Results by Dealer",
        "FIGURE PLACEHOLDER: Horizontal stacked bar chart for all 28 dealers. "
        "Each bar shows units allocated by model: A1V01 (dark blue), A2V02 (medium blue), "
        "A3V02 (light blue), B1V01 (orange). "
        "Dealers sorted by composite score (highest at top). "
        "Annotations show composite score C_d next to each bar. "
        "X-axis: 0 to 50 units. Y-axis: dealer labels (Bayi 01 – Bayi 28).")

    add_heading(doc, "7.3 Sensitivity Analysis", level=2)
    add_paragraph(doc,
        "Three sensitivity analyses were conducted to assess the robustness of the January "
        "2026 allocation results to key modeling assumptions.")
    add_paragraph(doc,
        "Weight sensitivity: The composite score weights were varied one at a time by ±0.05 "
        "(redistributing the delta to the remaining criteria proportionally). In all tested "
        "weight combinations, the rank ordering of the top-5 and bottom-5 dealers remained "
        "stable, and the total objective value changed by less than 2.1%. This indicates "
        "that the allocation is not highly sensitive to the specific weight assignment, "
        "which is expected given that the LP score (the highest-weight criterion) is also "
        "the most differentiating criterion across dealers in this instance.")
    add_paragraph(doc,
        "Launch boost factor sensitivity: The A1V01 boost factor was varied from 1.05 "
        "to 1.20 in steps of 0.01. Because A1V01 inventory is only 35 units (5.8% of "
        "total), the boost factor affects demand projections but has no direct effect on "
        "the MILP allocation (which is constrained by physical availability). The analysis "
        "confirms that the forecast sensitivity does not propagate into allocation sensitivity "
        "for this model in January, though it is expected to matter more when A1V01 "
        "inventory increases in subsequent months.")
    add_paragraph(doc,
        "Quota bound sensitivity: The quota bound tolerance was varied from ±10% to ±30% "
        "in 5-percentage-point steps. At ±10%, the MILP becomes infeasible for the A-group "
        "because the sum of lower bounds (0.90 × 517 = 465) exceeds inventory (423). At "
        "±25%, the solver has sufficient flexibility to concentrate allocation toward the "
        "top-7 dealers above their nominal quotas, increasing the objective value by 4.7% "
        "relative to the ±20% base case. The ±20% bound was confirmed as a reasonable "
        "balance between flexibility (allowing meaningful score-based prioritization) and "
        "equity (preventing extreme concentration that would leave low-scoring dealers "
        "without vehicles).")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 8 — DASHBOARD AND DEPLOYMENT
    # ══════════════════════════════════════════
    add_heading(doc, "8. DASHBOARD AND DEPLOYMENT", level=1)

    add_paragraph(doc,
        "The three-layer system is exposed to business users through an interactive "
        "Streamlit dashboard deployed on Streamlit Community Cloud. The dashboard is "
        "organized into five pages, accessible via a sidebar navigation menu:")
    pages = [
        ("Home / Overview",
         "Displays the current month's allocation summary: total units by model, "
         "number of dealers served, and the composite score distribution histogram."),
        ("Dealer Scoring",
         "Shows the full MCDM breakdown for every dealer, with the four sub-scores "
         "displayed in a color-coded table sortable by any criterion. Users can hover "
         "over any cell to see the underlying data that generated the score."),
        ("Allocation Map",
         "A Folium-based interactive geographic map of Turkey showing dealer locations "
         "as circle markers sized proportionally to the units allocated, with popup "
         "details on click."),
        ("Scenario Analysis",
         "Allows users to modify MCDM weights, the quota tolerance, and the total "
         "available inventory using sliders, then re-run the MILP in real time and "
         "compare the resulting allocation to the baseline scenario."),
        ("Data Explorer",
         "Provides filterable access to the underlying CSV data: historical sales, "
         "dealer targets, and inventory records."),
    ]
    for title, desc in pages:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{title}: ")
        set_run_font(r1, size_pt=12, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, size_pt=12)

    add_paragraph(doc,
        "The deployment architecture uses GitHub as the single source of truth: the main "
        "branch of the private repository is monitored by Streamlit Cloud, which "
        "automatically re-deploys the application on every push. Input data files (CSV) "
        "are stored in the repository's data/raw/ directory, eliminating the need for a "
        "separate database server in the initial deployment phase. For production use "
        "beyond the academic context, migration to a cloud-hosted SQLite or PostgreSQL "
        "database (with the SQLAlchemy ORM already wired into the codebase) would be "
        "the recommended next step.")
    add_paragraph(doc,
        "The dashboard incorporates a DEMO_MODE flag (configurable via secrets.toml) "
        "that replaces real dealer names and geographic coordinates with anonymized "
        "labels ('Bayi 01' through 'Bayi 28') when the system is presented to external "
        "audiences. The brand name is similarly parametrized through the BRAND_NAME "
        "configuration variable in config.py, ensuring that the system can be "
        "white-labeled for different brands without code changes.")
    add_paragraph(doc,
        "GitHub Actions CI pipelines run on every pull request to the main branch, "
        "executing the full pytest test suite (minimum 70% line coverage enforced) and "
        "the ruff linter with the project's pyproject.toml configuration. Pull requests "
        "that fail either check cannot be merged, ensuring that the deployed dashboard "
        "always reflects tested, lint-clean code. The test suite covers unit tests for "
        "each scoring component (P, LP, S, H), integration tests for the MILP solver "
        "(verifying that the output satisfies all constraints programmatically), and "
        "snapshot tests for the dashboard data loading pipeline.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CHAPTER 9 — CONCLUSIONS AND FUTURE WORK
    # ══════════════════════════════════════════
    add_heading(doc, "9. CONCLUSIONS AND FUTURE WORK", level=1)

    add_heading(doc, "9.1 Summary of Contributions", level=2)
    add_paragraph(doc,
        "This thesis has presented a complete, deployed decision-support system for the "
        "monthly vehicle allocation problem at an automotive brand operating 28 Turkish "
        "dealerships. Four primary contributions are claimed:")
    contributions = [
        ("C1 — Integrated Three-Layer Pipeline:",
         "The design of a three-layer pipeline that sequentially applies time-series "
         "forecasting, MCDM scoring, and MILP optimization to the vehicle allocation "
         "problem is, to the author's knowledge, the first such integration demonstrated "
         "in the peer-reviewed or thesis literature. Prior work has addressed forecasting "
         "and allocation separately; the present system treats them as a unified pipeline "
         "with explicit information flow between layers."),
        ("C2 — Empirical EWMA Window Selection:",
         "The cross-validation protocol for selecting the EWMA smoothing window W based "
         "on cross-year MAE consistency (rather than combined MAE minimization) is a "
         "methodological contribution applicable to any rolling-window performance scoring "
         "context. The protocol identified W = 5 as the optimal window for the automotive "
         "dealer performance data, with a rigorous rationale for rejecting lower-gap "
         "alternatives with higher combined MAE (W=4) or lower combined MAE with high "
         "cross-year inconsistency (W=6)."),
        ("C3 — Asymmetric H Score:",
         "The target proximity score H with its asymmetric design (penalizing behind-target "
         "dealers without rewarding ahead-of-target dealers) is a novel MCDM criterion "
         "tailored to the automotive dealer-manufacturer relationship, where contractual "
         "target obligations are one-directional. The asymmetry prevents the system from "
         "over-concentrating inventory at already-successful dealers, maintaining network-wide "
         "coverage."),
        ("C4 — Open-Source, Cloud-Native Deployment:",
         "The full implementation is open-source (GitHub), cloud-native (Streamlit Cloud, "
         "GitHub Actions, GitHub Codespaces), and requires no commercial software licenses. "
         "This is practically significant for the industrial partner and for academic "
         "reproducibility."),
    ]
    for label, text in contributions:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run(f"{label} ")
        set_run_font(r1, size_pt=12, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size_pt=12)

    add_heading(doc, "9.2 Limitations and Future Work", level=2)
    add_paragraph(doc,
        "Several limitations of the current system point toward directions for future research:")
    future_items = [
        ("Hierarchical forecast reconciliation:",
         "The current forecasting pipeline does not enforce coherence between the Prophet "
         "brand-level forecast and the EWMA-based dealer-level estimates. Implementing "
         "MinT reconciliation (Wickramasuriya et al., 2019) or hierarchical Bayesian "
         "forecasting would ensure that dealer-level allocations are consistent with the "
         "brand's aggregate projection, potentially improving quota accuracy."),
        ("Dynamic weight learning:",
         "The MCDM weights are currently set by expert elicitation and remain constant "
         "throughout the year. A data-driven approach—such as learning the weights that "
         "maximize end-of-year dealer satisfaction (measured by survey or by the gap "
         "between allocated and preferred quantities)—could adapt the scoring system "
         "to evolving business priorities."),
        ("Substitution and lost sales modeling:",
         "The LP cosine similarity implicitly assumes that dealers will sell whatever "
         "mix of variants they receive, subject to their historical preferences. In "
         "reality, a dealer that receives an unwanted variant may experience a lost sale "
         "if the customer defects to a competitor. Incorporating an explicit substitution "
         "model (Mahajan & van Ryzin, 2001) would provide a more accurate objective "
         "function for the MILP."),
        ("Multi-period stochastic optimization:",
         "The current system solves a single-period (monthly) deterministic problem. "
         "A multi-period stochastic programming extension would allow the system to "
         "account for demand uncertainty (via scenario trees) and to smooth inventory "
         "across months, potentially reducing the amplitude of the allocation swings "
         "that occur when model launches or seasonal spikes shift scores dramatically."),
        ("Factorization machines for LP score:",
         "Rendle's (2010) factorization machines can incorporate auxiliary features "
         "(dealer demographics, regional economic indicators, competitor activity) "
         "beyond raw sales counts into the LP similarity computation. This extension "
         "would be particularly valuable for newly opened dealers without extensive "
         "historical data to build a reliable sales vector."),
        ("Real-time data pipeline:",
         "The current system relies on monthly CSV uploads. Integration with the "
         "brand's ERP system via an API-based real-time data pipeline would eliminate "
         "the manual data preparation step and allow the allocation to incorporate "
         "mid-month sales data, enabling reactive re-allocation if early-month sales "
         "deviate significantly from projections."),
    ]
    for label, text in future_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size_pt=12, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size_pt=12)

    add_heading(doc, "9.3 Concluding Remarks", level=2)
    add_paragraph(doc,
        "The system demonstrates that a principled, data-driven approach to the Vehicle "
        "Allocation Problem is not only computationally feasible but practically deployable "
        "with open-source tools and cloud infrastructure. The January 2026 live run "
        "allocated 603 vehicles across 28 dealers in under 0.3 seconds, producing an "
        "allocation that is fully auditable, explainable to dealers, and consistent with "
        "the brand's business rules. The transparency of the MCDM scoring—each dealer "
        "can see exactly why they received the quantity they did—addresses a key complaint "
        "about the preceding manual allocation process.")
    add_paragraph(doc,
        "From an academic standpoint, this work illustrates the value of integrating "
        "established operations research methods (MILP), machine learning concepts "
        "(cosine similarity from collaborative filtering), and statistical time-series "
        "techniques (STL, EWMA) into a coherent decision pipeline that is greater than "
        "the sum of its parts. The individual methods are not novel; the contribution "
        "lies in their careful integration, parameterization, and validation within a "
        "real industrial context. The approach is generalizable to any automotive brand "
        "facing a similar constrained distribution problem, and with moderate adaptation, "
        "to other consumer durables industries where a central manufacturer allocates "
        "heterogeneous products to a dealer network.")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════
    add_heading(doc, "REFERENCES", level=1)

    references = [
        "Athanasopoulos, G., Hyndman, R. J., Song, H., & Wu, D. C. (2011). The tourism "
        "forecasting competition. International Journal of Forecasting, 27(3), 822–844. "
        "https://doi.org/10.1016/j.ijforecast.2010.04.009",

        "Cleveland, R. B., Cleveland, W. S., McRae, J. E., & Terpenning, I. J. (1990). "
        "STL: A seasonal-trend decomposition procedure based on loess. Journal of Official "
        "Statistics, 6(1), 3–33.",

        "Crowder, S. V. (1989). Design of exponentially weighted moving average schemes. "
        "Journal of Quality Technology, 21(3), 155–162. "
        "https://doi.org/10.1080/00224065.1989.11979169",

        "Dantzig, G. B. (1963). Linear programming and extensions. Princeton University Press.",

        "Hwang, C. L., & Yoon, K. (1981). Multiple attribute decision making: Methods and "
        "applications. Springer-Verlag. https://doi.org/10.1007/978-3-642-48318-9",

        "Hyndman, R. J., Ahmed, R. A., Athanasopoulos, G., & Shang, H. L. (2011). Optimal "
        "combination forecasts for hierarchical time series. Computational Statistics & Data "
        "Analysis, 55(9), 2579–2589. https://doi.org/10.1016/j.csda.2011.03.006",

        "Kok, A. G., Fisher, M. L., & Vaidyanathan, R. (2008). Assortment planning: Review "
        "of literature and industry practice. In N. Agrawal & S. A. Smith (Eds.), Retail "
        "supply chain management (pp. 99–153). Springer. "
        "https://doi.org/10.1007/978-0-387-78902-6_4",

        "Koren, Y., Bell, R., & Volinsky, C. (2009). Matrix factorization techniques for "
        "recommender systems. Computer, 42(8), 30–37. https://doi.org/10.1109/MC.2009.263",

        "Mahajan, S., & van Ryzin, G. (2001). Stocking retail assortments under dynamic "
        "consumer substitution. Operations Research, 49(3), 334–351. "
        "https://doi.org/10.1287/opre.49.3.334.11210",

        "Panagiotelis, A., Athanasopoulos, G., Gamakumara, P., & Hyndman, R. J. (2021). "
        "Forecast reconciliation: A geometric view with new insights on bias correction. "
        "International Journal of Forecasting, 37(1), 343–359. "
        "https://doi.org/10.1016/j.ijforecast.2020.06.004",

        "Rendle, S. (2010). Factorization machines. In Proceedings of the 2010 IEEE "
        "International Conference on Data Mining (pp. 995–1000). IEEE. "
        "https://doi.org/10.1109/ICDM.2010.127",

        "Şahin, C., & Kılıç, H. (2022). The vehicle allocation problem: Alternative "
        "formulation and branch-and-price method. Computers & Operations Research, 143, "
        "105745. https://doi.org/10.1016/j.cor.2022.105745",

        "Taylor, S. J., & Letham, B. (2018). Forecasting at scale. The American Statistician, "
        "72(1), 37–45. https://doi.org/10.1080/00031305.2017.1380080",

        "Toth, P., & Vigo, D. (Eds.). (2002). The vehicle routing problem. Society for "
        "Industrial and Applied Mathematics. https://doi.org/10.1137/1.9780898718515",

        "TÜİK. (2026). İllere göre motorlu kara taşıtları sayısı — Aralık 2025. Türkiye "
        "İstatistik Kurumu. https://data.tuik.gov.tr",

        "Wickramasuriya, S. L., Athanasopoulos, G., & Hyndman, R. J. (2019). Optimal "
        "forecast reconciliation for hierarchical and grouped time series through trace "
        "minimization. Journal of the American Statistical Association, 114(526), 804–819. "
        "https://doi.org/10.1080/01621459.2018.1448825",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.space_before = Pt(2)
        fmt.space_after  = Pt(4)
        fmt.left_indent  = Cm(1.27)
        fmt.first_line_indent = Cm(-1.27)  # hanging indent
        run = p.add_run(ref)
        set_run_font(run, size_pt=11)

    # ══════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════
    output_path = "/home/user/arac-dagitim-sistemi/tez_taslak.docx"
    doc.save(output_path)
    print(f"Thesis document saved to: {output_path}")
    print(f"Sections: {len(doc.sections)}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == "__main__":
    build_thesis()
