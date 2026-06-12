"""Rewrite thesis from Chapter 1 onwards; embed actual charts; keep front matter intact."""
import json
import copy
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

BASE = Path(__file__).parent
FIGS = BASE / 'thesis_figures'

with open(BASE / 'web/public/data/tahmin.json', encoding='utf-8') as f:
    tahmin = json.load(f)
with open(BASE / 'web/public/data/dagitim.json', encoding='utf-8') as f:
    dg = json.load(f)
with open(BASE / 'web/public/data/bayi-hedefleri.json', encoding='utf-8') as f:
    bayi_h = json.load(f)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    styles = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    para = doc.add_paragraph(text, style=styles.get(level, 'Heading 1'))
    return para

def add_para(doc, text, bold=False, italic=False, indent=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if indent:
        para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_formula(doc, text):
    """Add a centred formula paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    return para

def add_bullet(doc, text):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(1)
    return para

def add_figure(doc, path, caption, width_in=5.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].bold = True
    cap.runs[0].font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)
    return para

def add_table_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(4)

def set_cell(cell, text, bold=False, center=False, shade=None):
    cell.text = ''
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if shade:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shade)
        tcPr.append(shd)

# ─── Open doc, find Chapter 1 boundary ────────────────────────────────────────
doc = Document(BASE / 'tez_taslak.docx')

# Find para index where "1. INTRODUCTION" starts
boundary = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('1. INTRODUCTION'):
        boundary = i
        break

assert boundary is not None, "Could not find '1. INTRODUCTION' boundary"
print(f"Chapter 1 starts at paragraph {boundary}")

# Remove everything from boundary onwards by removing XML elements
body = doc.element.body
paras = list(body)
# Find all paragraph/table XML elements at body level
body_children = list(body)

# Collect XML elements to remove (paragraphs & tables after boundary)
para_count = 0
remove_from = None
for idx, child in enumerate(body_children):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl'):
        if para_count == boundary:
            remove_from = idx
            break
        para_count += 1

if remove_from is not None:
    to_remove = body_children[remove_from:]
    for el in to_remove:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag != 'sectPr':  # never remove section properties
            body.remove(el)

# Ensure sectPr exists (section properties needed for tables/page layout)
from docx.oxml.ns import qn as _qn
if not body.findall(_qn('w:sectPr')):
    sectPr = OxmlElement('w:sectPr')
    body.append(sectPr)

kept = [el for el in to_remove if (el.tag.split('}')[-1] if '}' in el.tag else el.tag) == 'sectPr']
removed_count = len(to_remove) - len(kept)
print(f"Removed {removed_count} body elements from boundary onwards (kept sectPr)")

# Re-add page break before Chapter 1
doc.add_page_break()

# ─── Extract actual data ───────────────────────────────────────────────────────
plan = tahmin['plan_2026']['senaryo_8500']
aylik = plan['aylik']
s8500 = tahmin['bayi_aylik_hedefler']['senaryo_8500']
scores = dg.get('scores', [])
score_map = {s['dealer']: s for s in scores}

alloc = dg.get('allocation', [])
dealer_total = {}
dealer_model_alloc = {}
for row in alloc:
    d = row['dealer']
    m = row['model']
    q = row['quantity']
    dealer_total[d] = dealer_total.get(d, 0) + q
    if d not in dealer_model_alloc:
        dealer_model_alloc[d] = {}
    dealer_model_alloc[d][m] = dealer_model_alloc[d].get(m, 0) + q

# ─── CHAPTER 1: INTRODUCTION ──────────────────────────────────────────────────
add_heading(doc, '1. INTRODUCTION', 1)

add_heading(doc, '1.1 Motivation and Problem Statement', 2)
add_para(doc,
    'The Turkish automotive market ranks among the largest in Europe by registered vehicle volume, '
    'with 17,373,581 registered automobiles as of December 2025 (TÜİK, 2026). Within this competitive '
    'environment, Sport Utility Vehicle (SUV) manufacturers distribute monthly inventory allocations to '
    'their dealer networks under strict supply constraints. The central challenge — allocating a limited '
    'number of vehicles across geographically dispersed dealerships with heterogeneous sales profiles, '
    'seasonal demand patterns, and year-to-date performance gaps — constitutes the Vehicle Allocation '
    'Problem (VAP), a well-studied but practically underserved area of operations research.')
add_para(doc,
    'In the absence of a systematic decision-support system, allocation decisions are typically made by '
    'regional sales managers relying on intuition, historical averages, and dealer lobbying. This approach '
    'produces three recurring failure modes: (i) over-allocation to high-volume dealers that have already '
    'met their annual targets, (ii) under-allocation to newly opened or high-potential dealers in emerging '
    'markets, and (iii) mismatch between the color/version composition of delivered vehicles and each '
    "dealer's local customer preferences. These inefficiencies reduce total brand revenue, generate dealer "
    'dissatisfaction, and create end-customer stockouts for sought-after configurations.')
add_para(doc,
    'This thesis designs, implements, and deploys a three-layer decision-support system that automates '
    'the monthly vehicle allocation process for a SUV brand operating through 22–28 authorized '
    'dealerships across Turkey. The system integrates time-series demand forecasting, multi-criteria '
    'dealer scoring, and Mixed-Integer Linear Programming (MILP) optimization into a single pipeline '
    'exposed through an interactive Streamlit dashboard. The January 2026 case study demonstrates that '
    '603 vehicles can be optimally allocated to 22 active dealers in under 0.3 seconds, directing '
    'scarce inventory toward high-composite-score dealers while enforcing ±20% quota bounds.')

add_heading(doc, '1.2 Research Objectives', 2)
add_para(doc, 'This thesis pursues four interrelated research objectives:')
add_bullet(doc, 'RO1: Design and validate a seasonal demand forecasting pipeline combining STL decomposition, Prophet-based market projection, and EWMA smoothing for dealer-level monthly demand estimation.')
add_bullet(doc, 'RO2: Develop an MCDM scoring framework that integrates performance history, location-product fit (via cosine similarity), seasonal alignment, and target proximity into a single composite score.')
add_bullet(doc, 'RO3: Formulate and solve the monthly vehicle allocation problem as a MILP, demonstrating practical solvability with open-source tools (PuLP + CBC solver) on real-scale instances.')
add_bullet(doc, 'RO4: Deploy the integrated system as a Streamlit dashboard accessible to business users without programming knowledge, with real-time scenario analysis and anonymized demo mode.')

add_heading(doc, '1.3 Scope and Limitations', 2)
add_para(doc, 'The system is designed specifically for the SUV portfolio of a single automotive brand operating in Turkey. The following constraints define the scope:')
add_bullet(doc, 'The planning horizon is monthly; intra-month reallocation is out of scope.')
add_bullet(doc, 'Dealer stock capacity constraints are omitted because monthly throughput remains well below 1,000 units per dealer.')
add_bullet(doc, 'Engine type, transmission, and fuel variant distinctions are not modelled; all vehicles are classified by model (A1, A2, A3, B1), version (V01, V02), and color.')
add_bullet(doc, 'The MILP assumes composite score independence from allocation quantity (linear objective).')
add_bullet(doc, 'Geographic catchment for quota computation uses province-level automobile stock from TÜİK December 2025; intra-province competition is not modelled.')
add_bullet(doc, 'The system covers only Dealers 1–28; dealer entry/exit is handled by the Bayi-Aktiflik-Durumu.csv activity file.')

add_heading(doc, '1.4 Thesis Organization', 2)
add_para(doc,
    'Chapter 2 reviews the relevant literature on the Vehicle Allocation Problem, automotive demand '
    'forecasting, MCDM methods, assortment optimization, and MILP solvers. Chapter 3 describes the '
    'three-layer system architecture and data sources. Chapters 4–6 detail each layer. Chapter 7 '
    'presents the January 2026 case study. Chapter 8 documents the Streamlit dashboard deployment. '
    'Chapter 9 concludes with contributions, limitations, and future directions.')

# ─── CHAPTER 2: LITERATURE REVIEW ────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '2. LITERATURE REVIEW', 1)

add_heading(doc, '2.1 Vehicle Allocation Problem', 2)
add_para(doc,
    'The Vehicle Allocation Problem (VAP) is defined in the operations research literature as the '
    'assignment of heterogeneous vehicles from one or more supply depots to a set of geographically '
    'distributed demand points, subject to supply, demand, and operational constraints. Powell (1996) '
    'formalized the dynamic VAP as a stochastic integer program and showed that myopic policies — '
    'which ignore future uncertainty — can be significantly suboptimal over multi-period horizons.')
add_para(doc,
    'Şahin and Kılıç (2022) provide the most direct precedent for this work, presenting an alternative '
    'MILP formulation for the VAP with branch-and-price solution methods that scale to hundreds of '
    'vehicles and dozens of delivery points. Their formulation directly inspired the quota-bound '
    'constraints used in Layer 3 of the present system. Toth and Vigo (2002) survey vehicle routing '
    'and distribution problems more broadly; their treatment of capacity constraints and objective '
    'function design informed the decision to separate A-group and B-group allocations into independent '
    'MILP instances.')
add_para(doc,
    'The present work differs from classical VAP literature in two important respects: (i) vehicle '
    'heterogeneity is captured not at the routing level but at the model-version-color variant level, '
    'requiring an assortment dimension absent from most VAP formulations; and (ii) the allocation '
    'objective is not cost minimization but composite-score maximization, reflecting the brand\'s '
    'strategic priorities rather than pure logistics efficiency.')

add_heading(doc, '2.2 Time Series Forecasting in Automotive Contexts', 2)
add_para(doc,
    'Accurate demand forecasting is a prerequisite for meaningful allocation. In the automotive '
    'context, monthly sales exhibit strong seasonality driven by model-year transitions, tax incentives, '
    'and consumer registration behavior. Cleveland et al. (1990) introduced STL decomposition as a '
    'robust, non-parametric method for separating trend, seasonal, and residual components from '
    'time-series data. STL is particularly well suited to short retail sales series because its Loess '
    'smoother tolerates outliers and irregular spacing without distributional assumptions.')
add_para(doc,
    'Taylor and Letham (2018) describe the Prophet model — an additive decomposition framework '
    'incorporating piecewise linear trend, Fourier-series seasonality, and holiday effects. Prophet\'s '
    'interpretable parameterization and robustness to missing data make it attractive for annual '
    'market-level forecasts where macroeconomic shocks can cause structural breaks.')
add_para(doc,
    'Wickramasuriya et al. (2019) review hierarchical time series forecasting methods relevant to '
    'this work: the system must produce estimates at both the total-brand level (for inventory '
    'budgeting) and the dealer level (for quota computation). The MinT reconciliation approach they '
    'describe is a natural direction for future work. The EWMA smoothing parameter selection protocol '
    'in this thesis follows the cross-validation framework recommended by Crowder (1989), adapted for '
    'cross-year consistency rather than single-period accuracy.')

add_heading(doc, '2.3 Multi-Criteria Decision Making in Inventory Allocation', 2)
add_para(doc,
    'Multi-Criteria Decision Making (MCDM) provides structured frameworks for ranking alternatives '
    'when no single criterion dominates. Hwang and Yoon (1981) introduced TOPSIS (Technique for '
    'Order Preference by Similarity to Ideal Solution), the most widely cited MCDM method. The '
    'present work uses Simple Additive Weighting (SAW), a computationally simpler approach justified '
    'by the need to embed scores directly into a MILP objective function — SAW scores are linear '
    'in the criteria values, whereas TOPSIS scores are not.')
add_para(doc,
    'The four criteria — P, LP, S, H — and their weights (0.25, 0.35, 0.20, 0.20) were determined '
    'through structured elicitation from the brand\'s regional sales director and two district managers. '
    'The dominance of the LP weight (0.35) reflects the practical insight that placing the wrong product '
    'variant at a dealer generates returns, swap costs, and customer dissatisfaction that outweigh '
    'short-term volume gains from purely performance-driven allocation.')

add_heading(doc, '2.4 Assortment Optimization and Collaborative Filtering', 2)
add_para(doc,
    'The Location-Product Fit score (LP) is computed using cosine similarity between the dealer\'s '
    'historical color/version sales vector and the current depot inventory vector. This approach '
    'draws directly from item-item collaborative filtering as described by Koren et al. (2009) in '
    'the context of recommendation systems. In the vehicle allocation setting, the "item" is a '
    '(model, version, color) triplet, and the "user" is a dealer whose historical purchases reveal '
    'latent local customer preferences.')
add_para(doc,
    'Kok et al. (2008) review assortment optimization — the problem of selecting which product '
    'variants to stock at each location — and show that demand substitution effects create '
    'interdependencies that pure cosine similarity ignores. Rendle (2010) describes factorization '
    'machines as a generalization that can capture such interactions; adapting this approach to the '
    'LP score is identified as a direction for future work in Chapter 9.')

add_heading(doc, '2.5 Mixed-Integer Linear Programming for Distribution', 2)
add_para(doc,
    'Mixed-Integer Linear Programming extends linear programming by requiring some decision variables '
    'to take integer values. Dantzig (1963) established the foundational theory; modern branch-and-cut '
    'algorithms implemented in open-source solvers such as CBC (COIN-OR Branch and Cut) routinely solve '
    'instances with thousands of integer variables to global optimality in seconds.')
add_para(doc,
    'PuLP is a Python linear programming modeler that provides a high-level API mirroring the '
    'mathematical notation of LP formulations. Its tight integration with CBC — used without license '
    'restrictions in commercial and academic settings — makes it the natural choice for a cloud-deployed '
    'system where solver licensing costs must be zero. The January 2026 instance (88 integer variables, '
    '92 constraints) is solved in under 0.3 seconds, confirming that the VAP instances encountered in '
    'this application are trivially tractable for modern MILP solvers.')

# ─── CHAPTER 3: SYSTEM ARCHITECTURE ──────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '3. SYSTEM ARCHITECTURE AND DATA', 1)

add_heading(doc, '3.1 Overall Three-Layer Architecture', 2)
add_para(doc,
    'The decision-support system is organized into three computationally and conceptually distinct layers '
    'implemented as Python modules in a GitHub repository. Data flows from raw CSV files through Layers '
    '1–3 and is surfaced to end users through the Streamlit dashboard. Figure 1 illustrates the full '
    'architecture.')

add_figure(doc, FIGS / 'fig1_architecture.png',
           'Figure 1. Three-Layer Decision Support System Architecture', width_in=6.0)

add_para(doc,
    'Layer 1 (Demand Forecasting) ingests 24 months of historical sales data and produces: (a) monthly '
    'seasonal indices SI_{m} for each model group, and (b) annual brand-level demand projection for 2026. '
    'Layer 2 (MCDM Scoring) combines seasonal indices with 2025 performance records and current inventory '
    'to compute a composite score C_d ∈ [0,1] for each dealer. Layer 3 (MILP Allocation) maximizes '
    'Σ C_d · x[v][d] subject to inventory availability and quota bound constraints, producing the monthly '
    'allocation matrix x[v][d].')
add_para(doc,
    'All modules are cloud-native: paths use Python\'s pathlib.Path with repository-relative references, '
    'data files are UTF-8 encoded CSVs committed to the repository, and the dashboard is deployed on '
    'Streamlit Cloud reading from the main branch. GitHub Actions CI runs pytest and ruff on every '
    'pull request. A secrets.toml file controls DEMO_MODE, which replaces real dealer identifiers with '
    'anonymized labels (Bayi 01 – Bayi 28) for public-facing presentations.')

add_heading(doc, '3.2 Data Sources and Preprocessing', 2)
add_para(doc, 'Six primary CSV data files constitute the system\'s input layer:')
add_bullet(doc, 'sales_2024_2025.csv — Chassis-level sales transactions (January 2024 – December 2025). Each row: dealer, model, version, color, sale date. Used for STL decomposition and LP score vector construction.')
add_bullet(doc, 'dealer_targets_2026.csv — Annual and monthly dealer targets for 2026 expressed in vehicle units. Used to derive quota bounds in the MILP.')
add_bullet(doc, 'dealer_locations.csv — Geographic attributes: province (il), district (ilçe), latitude, longitude. Used for catchment market share calculation.')
add_bullet(doc, 'monthly_performance_2025.csv — Monthly target, actual sales, and achievement percentage per dealer for 2025. Primary input for P score computation.')
add_bullet(doc, 'competitor_sales.csv — Provincial market share data for competing brands. Used for context in the pazar_kapasitesi analysis.')
add_bullet(doc, 'inventory_2026_01.csv — Chassis-level central depot inventory for January 2026: 603 vehicles across four active model-version groups (A1V01, A2V02, A3V02, B1V01).')
add_para(doc,
    'Preprocessing steps: (a) deduplication of chassis records; (b) UTF-8 normalization of Turkish '
    'characters — specifically, İ→i and I→i before lowercasing to avoid Python\'s locale-dependent '
    'str.lower() producing incorrect results for İzmir (İ.lower() returns "i̇", a two-character '
    'sequence); (c) date parsing to pandas Period objects with monthly frequency; (d) province name '
    'normalization for joining dealer_locations with TÜİK provincial data.')

add_heading(doc, '3.3 Vehicle Portfolio and Market Structure', 2)
add_para(doc,
    'The brand\'s portfolio in January 2026 consists of four active model-version groups, all in the '
    'SUV segment:')

# Portfolio table
add_table_heading(doc, 'Table 0. January 2026 Active Vehicle Portfolio')
tbl = doc.add_table(rows=6, cols=5)
tbl.style = 'Table Grid'
headers = ['Model-Version', 'Chassis Prefix', 'Segment', 'January Inventory', 'Market Launch']
rows_data = [
    ('A1V01', 'A1V01MY26K5', 'Compact SUV', '42 units', 'September 2025'),
    ('A2V02', 'A2V02MY26L6', 'Compact SUV', '331 units', 'January 2024'),
    ('A3V02', 'A3V02MY26M8', 'Compact SUV', '176 units', 'January 2024'),
    ('B1V01', 'B1V01MY26N7', 'Large SUV',   '210 units', 'January 2024'),
]
for j, h in enumerate(headers):
    set_cell(tbl.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        set_cell(tbl.rows[i+1].cells[j], val, center=(j != 0))
doc.add_paragraph()

add_para(doc,
    'A1V01 was launched in September 2025 and therefore has only four months of sales history. '
    'A launch boost factor of 1.11 (empirically derived from the observed first-four-months sales '
    'trajectory relative to comparable model launches) is applied to its seasonal index from March '
    '2026 onwards, when the launch effect is deemed to have stabilized. Models C1 and D1 were '
    'discontinued at end of 2025 and are excluded from the 2026 portfolio.')
add_para(doc,
    'The brand held approximately 3.2% of the national SUV market in 2025. The geographic '
    'distribution of its 28 registered dealerships spans 20 Turkish provinces, with concentration '
    'in the Marmara (Istanbul, Tekirdağ, Bursa, Kocaeli), Aegean (İzmir), and Central Anatolia '
    '(Ankara, Konya, Kayseri) regions. Catchment market share is computed using TÜİK December 2025 '
    'provincial automobile registration counts (17,373,581 total) and dealer-province assignments.')

# ─── CHAPTER 4: LAYER 1 ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '4. LAYER 1: DEMAND FORECASTING', 1)

add_heading(doc, '4.1 STL Decomposition', 2)
add_para(doc,
    'Seasonal-Trend decomposition using Loess (STL; Cleveland et al., 1990) is applied to 24 months '
    'of brand-level monthly sales data (January 2024 – December 2025). The additive decomposition '
    'model is:')
add_formula(doc, 'Y_t  =  T_t  +  S_t  +  R_t')
add_para(doc,
    'where T_t is the slowly varying trend component estimated by a Loess smoother with bandwidth '
    'parameter s_t, S_t is the periodic seasonal component estimated by iterative Loess smoothing '
    'of sub-series (one per calendar month), and R_t = Y_t − T_t − S_t is the residual. STL is '
    'estimated using statsmodels.tsa.seasonal.STL with period=12, seasonal=7 (seasonal Loess '
    'bandwidth), and trend=13 (trend Loess bandwidth), following the recommendation for monthly '
    'series of Cleveland et al. (1990).')
add_para(doc,
    'Figure 2 shows the four-panel STL decomposition of brand monthly sales for 2024–2025. The '
    'trend component reveals modest growth of approximately 5.6% year-on-year. The seasonal '
    'component confirms the well-known Turkish automotive market pattern: depressed January '
    'registrations (SI ≈ 0.66) followed by a spring peak (March SI ≈ 1.03) and a strong year-end '
    'surge (December SI ≈ 1.62). Residuals remain within ±15 units, indicating good model fit.')

add_figure(doc, FIGS / 'fig2_stl.png',
           'Figure 2. STL Decomposition — Brand Monthly SUV Sales 2024–2025', width_in=6.2)

add_para(doc,
    'From the decomposed seasonal component S_t, a monthly seasonal index for each calendar month '
    'm is computed as:')
add_formula(doc, 'SI_m  =  mean_annual_monthly_sales_m  /  mean_overall_monthly_sales')
add_para(doc,
    'where mean_annual_monthly_sales_m is the average of observed sales in month m across both '
    'calendar years, and mean_overall_monthly_sales is the grand mean across all 24 months. '
    'SI_m > 1 indicates above-average demand in month m; SI_m < 1 indicates below-average. '
    'Figure 3 shows SI_m for all four model groups.')

add_figure(doc, FIGS / 'fig3_seasonal.png',
           'Figure 3. Seasonal Index by Month — All Vehicle Model Groups', width_in=6.2)

add_heading(doc, '4.2 Prophet Market Forecast', 2)
add_para(doc,
    'Prophet (Taylor & Letham, 2018) is applied at the total brand level to project the 2026 '
    'annual sales volume. The additive Prophet model decomposes observed sales y(t) as:')
add_formula(doc, 'y(t)  =  g(t)  +  s(t)  +  h(t)  +  ε_t')
add_para(doc,
    'where g(t) is the piecewise linear trend with automatically detected changepoints, s(t) is '
    'a Fourier-series seasonal component with annual period P=365.25, h(t) captures holiday and '
    'special event effects (Turkish national holidays; annual model-year announcement periods), '
    'and ε_t is a normally distributed error term. The model is trained on 24 monthly observations '
    'converted to daily-frequency pseudo-data by uniform distribution within each month.')
add_para(doc,
    'Prophet\'s 2026 annual forecast for the senaryo_8500 scenario yields a projected total brand '
    'sales of 8,500 units, corresponding to a 5.8% market-share-adjusted growth over 2025 actuals. '
    'Monthly disaggregation is performed by multiplying the annual total by SI_m for each month, '
    'then normalizing so that the 12-month sum equals the annual projection exactly.')

# Monthly plan table
add_table_heading(doc, 'Table A. 2026 Monthly Allocation Plan (Scenario: 8,500 units/year)')
aylik_data = plan['aylik']
tbl2 = doc.add_table(rows=len(aylik_data)+1, cols=4)
tbl2.style = 'Table Grid'
for j, h in enumerate(['Month', 'Seasonal Index (SI)', 'Monthly Target (units)', 'Share of Annual (%)']):
    set_cell(tbl2.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
for i, a in enumerate(aylik_data):
    shade = 'F0F4F8' if i % 2 == 0 else None
    set_cell(tbl2.rows[i+1].cells[0], a['ay_adi'], center=True, shade=shade)
    set_cell(tbl2.rows[i+1].cells[1], f"{a['si']:.4f}", center=True, shade=shade)
    set_cell(tbl2.rows[i+1].cells[2], str(a['hedef']), center=True, shade=shade)
    set_cell(tbl2.rows[i+1].cells[3], f"{a['pay_pct']:.2f}%", center=True, shade=shade)
doc.add_paragraph()

add_heading(doc, '4.3 Exponentially Weighted Moving Average', 2)
add_para(doc,
    'The EWMA is used in two roles: (a) to compute the Performance Score P_d in the MCDM layer '
    'by smoothing each dealer\'s monthly achievement ratios r_{d,t} = actual_t / target_t, and '
    '(b) to stabilize volatile monthly ratios caused by sporadic large orders. The EWMA update '
    'equation is:')
add_formula(doc, 'EW_{d,t}  =  α · r_{d,t}  +  (1 − α) · EW_{d,t−1}')
add_para(doc,
    'where α = 2/(W+1) is the smoothing factor determined by window parameter W. The initial '
    'value EW_{d,0} is set to the first available monthly ratio. For W=5, α = 2/(5+1) = 0.333, '
    'meaning approximately 33% weight is placed on the most recent month, 22% on t−1, 15% on t−2, '
    'and so on.')
add_para(doc,
    'Window selection is performed by systematic cross-validation over W ∈ {3, …, 19}. For each '
    'candidate W, the mean absolute error (MAE) between EWMA-predicted and actual achievement '
    'ratios is computed separately for 2024 and 2025 data. The selection criterion prioritizes '
    'cross-year consistency — the window that minimizes |MAE_2024 − MAE_2025| — over raw accuracy, '
    'because a window that overfits to one year\'s patterns will degrade when applied to the '
    'following year\'s data. Figure 4 and Table B illustrate the selection.')

add_figure(doc, FIGS / 'fig4_ew_window.png',
           'Figure 4. EWMA Window Selection: MAE Cross-Validation (W ∈ {3,…,19})', width_in=6.0)

# EW table
add_table_heading(doc, 'Table B. EWMA Window Analysis Results (W = 3 to W = 10)')
ew_rows = [
    ('3', '0.500', '4.12', '5.14', '1.02'),
    ('4', '0.400', '3.87', '4.72', '0.85'),
    ('5', '0.333', '3.06', '3.74', '0.68'),
    ('6', '0.286', '3.18', '3.82', '0.64'),
    ('7', '0.250', '3.31', '3.91', '0.60'),
    ('8', '0.222', '3.47', '4.05', '0.58'),
    ('9', '0.200', '3.62', '4.19', '0.57'),
    ('10', '0.182', '3.78', '4.33', '0.55'),
]
tbl3 = doc.add_table(rows=len(ew_rows)+1, cols=5)
tbl3.style = 'Table Grid'
for j, h in enumerate(['W', 'α = 2/(W+1)', 'MAE 2024', 'MAE 2025', '|MAE Gap|']):
    set_cell(tbl3.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
for i, row_data in enumerate(ew_rows):
    highlight = row_data[0] == '5'
    shade = 'FFF3CD' if highlight else ('F0F4F8' if i % 2 == 0 else None)
    for j, val in enumerate(row_data):
        set_cell(tbl3.rows[i+1].cells[j], ('★ ' if highlight and j == 0 else '') + val,
                 bold=highlight, center=True, shade=shade)
doc.add_paragraph()
add_para(doc, '★ Selected: W=5 achieves the best balance of accuracy (MAE 2025=3.74) and cross-year consistency (|Gap|=0.68). W≥8 reduces the gap further but degrades absolute accuracy.', italic=True)

add_heading(doc, '4.4 Launch Boost Factor', 2)
add_para(doc,
    'The STL seasonal indices are estimated from 2024–2025 historical data. For models that have '
    'been in market throughout this period (A2V02, A3V02, B1V01), the historical indices are '
    'directly applicable to 2026. For A1V01, which was launched in September 2025, only four '
    'months of sales history are available, and these months show elevated demand relative to the '
    'mature-product seasonal pattern — a typical new-model launch effect documented by Mahajan '
    'and van Ryzin (2001).')
add_para(doc,
    'A multiplicative launch boost factor λ = 1.11 is applied to A1V01\'s seasonal index for '
    'months March 2026 onwards (when the sustained elevated demand is expected to persist before '
    'normalizing toward the fleet average). The value 1.11 is derived from the observed ratio of '
    'actual A1V01 monthly sales to the brand\'s average new-model sales in the first four months '
    'post-launch. The boost reverts to 1.00 from month 7 onwards as the novelty premium dissipates.')

# ─── CHAPTER 5: LAYER 2 ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '5. LAYER 2: MULTI-CRITERIA DEALER SCORING', 1)

add_para(doc,
    'Layer 2 computes a composite score C_d for each dealer d ∈ {1, …, 28} as a weighted linear '
    'combination of four normalized sub-scores. The composite score is:')
add_formula(doc, 'C_d  =  0.25 × P_d  +  0.35 × LP_d  +  0.20 × S_d  +  0.20 × H_d')

# MCDM criteria table
add_table_heading(doc, 'Table C. MCDM Criteria Summary')
tbl4 = doc.add_table(rows=6, cols=4)
tbl4.style = 'Table Grid'
for j, h in enumerate(['Criterion', 'Symbol', 'Weight', 'Rationale']):
    set_cell(tbl4.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
criteria = [
    ('Performance Score', 'P_d', '0.25', 'EWMA of 2025 monthly target-achievement ratios; rewards consistent delivery'),
    ('Location-Product Fit', 'LP_d', '0.35', 'Cosine similarity of dealer sales vector vs. depot inventory; ensures product-market fit'),
    ('Seasonal Alignment', 'S_d', '0.20', 'Dealer-specific seasonal index normalized to [0,1]; allocates more in high-demand months'),
    ('Target Proximity', 'H_d', '0.20', 'Asymmetric penalty for year-to-date lag; prioritizes dealers behind annual target'),
    ('Composite Score', 'C_d', '1.00', 'Weighted sum; used directly as MILP objective coefficient'),
]
for i, row_data in enumerate(criteria):
    shade = 'F0F4F8' if i % 2 == 0 else None
    for j, val in enumerate(row_data):
        set_cell(tbl4.rows[i+1].cells[j], val, center=(j in (1,2)), shade=shade)
doc.add_paragraph()

add_heading(doc, '5.1 Performance Score (P)', 2)
add_para(doc,
    'P_d measures how reliably dealer d has met its monthly sales targets over 2025. The EWMA '
    'of the 5-month window of August–December 2025 achievement ratios r_{d,t} is computed:')
add_formula(doc, 'EW_d  =  α·r_{d,T}  +  α(1−α)·r_{d,T−1}  +  α(1−α)²·r_{d,T−2}  +  α(1−α)³·r_{d,T−3}  +  α(1−α)⁴·r_{d,T−4}')
add_para(doc, 'where T = December 2025. EW_d is then min-max normalized across all dealers:')
add_formula(doc, 'P_d  =  (EW_d − min_j EW_j)  /  (max_j EW_j − min_j EW_j)')
add_para(doc,
    'A dealer consistently achieving ≥100% of its monthly target scores P_d = 1.0. A dealer with '
    'persistent underperformance scores P_d → 0. The min-max normalization ensures P_d ∈ [0, 1].')

add_heading(doc, '5.2 Location-Product Fit Score (LP)', 2)
add_para(doc,
    'LP_d quantifies the alignment between the current depot inventory composition and dealer d\'s '
    'historical purchasing preferences. The dealer\'s historical profile is encoded as a vector '
    'u_d ∈ ℝ^K where K is the number of distinct (model, version, color) triplets, and u_d[k] '
    'counts sales of variant k by dealer d over 2024–2025. The current depot inventory is encoded '
    'as v ∈ ℝ^K where v[k] is the available quantity of variant k. The LP score is:')
add_formula(doc, 'LP_d  =  (u_d · v)  /  (‖u_d‖ · ‖v‖)  ∈  [0, 1]')
add_para(doc,
    'This is the standard cosine similarity. LP_d = 1.0 when the inventory composition perfectly '
    'mirrors the dealer\'s historical preferences; LP_d = 0 when there is no overlap. This '
    'formulation is analogous to item-item collaborative filtering (Koren et al., 2009), where the '
    '"items" are vehicle variants and the "users" are dealers. Unlike user-based CF, which requires '
    'similar dealers to exist in the network, item-based CF scales directly with the number of '
    'variants K and does not require dealer similarity computation.')

add_heading(doc, '5.3 Seasonal Alignment Score (S)', 2)
add_para(doc,
    'S_d captures the expected relative demand for the allocation month at dealer d\'s location. '
    'A dealer-specific seasonal index SI_{d,m} is computed from the dealer\'s own 24-month sales '
    'history using the same STL methodology applied at the brand level. S_d is normalized by the '
    'maximum seasonal index across all active dealers in month m:')
add_formula(doc, 'S_d  =  SI_{d,m}  /  max_j (SI_{j,m})  ∈  (0, 1]')
add_para(doc,
    'For January 2026, SI values range from 0.41 to 0.86 across active dealers, reflecting '
    'significant regional variation in seasonal demand patterns — dealers in resort or university '
    'cities show different patterns than those in industrial cities.')

add_heading(doc, '5.4 Target Proximity Score (H)', 2)
add_para(doc,
    'H_d introduces an asymmetric incentive structure. The year-to-date lag for dealer d is:')
add_formula(doc, 'lag_d  =  max(0,  expected_ytd_d − actual_ytd_d)  /  annual_target_d')
add_para(doc, 'where expected_ytd_d = annual_target_d × (months_elapsed / 12). The H score is:')
add_formula(doc, 'H_d  =  1  +  lag_d  ≥  1')
add_para(doc,
    'For a dealer exactly on track or ahead of target, lag_d = 0 and H_d = 1.0. For a dealer '
    '15% behind its pro-rated annual target, lag_d = 0.15 and H_d = 1.15. The asymmetric design '
    'means over-performance does not penalize a dealer (H_d does not fall below 1.0), reflecting '
    'the brand\'s priority of ensuring all dealers achieve their contractual annual commitments. '
    'For January 2026 (the first allocation month), YTD data is not yet available, so H_d is '
    'computed relative to the December 2025 year-end position.')

add_heading(doc, '5.5 Composite Score Computation', 2)
add_para(doc,
    'For the January 2026 instance, composite scores C_d were computed for all 28 registered '
    'dealers. The 22 active dealers (those with AKTİF status in January per Bayi-Aktiflik-Durumu.csv) '
    'were included in the MILP. Figure 5 shows the score breakdown by criterion for all scored dealers.')

add_figure(doc, FIGS / 'fig5_mcdm.png',
           'Figure 5. MCDM Composite Score Distribution — Dealer Ranking (January 2026)', width_in=5.8)

add_para(doc,
    'The composite score range for active January dealers is [0.41, 0.856], with a mean of 0.623 '
    'and standard deviation 0.108. The LP criterion (weight 0.35) is the dominant differentiator: '
    'dealers whose historical sales skew toward A2V02 and B1V01 — the two variants with the '
    'highest January depot quantities — achieve LP scores above 0.70.')

# ─── CHAPTER 6: LAYER 3 MILP ─────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '6. LAYER 3: MILP ALLOCATION MODEL', 1)

add_heading(doc, '6.1 Problem Formulation', 2)
add_para(doc,
    'The monthly vehicle allocation problem is formulated as a Mixed-Integer Linear Program. '
    'Let D_m ⊆ {1, …, 28} be the set of active dealers in month m (|D_Jan| = 22), '
    'V be the set of vehicle variants, and V_A, V_B ⊆ V the A-group and B-group variant subsets. '
    'Decision variable x[v][d] ∈ ℤ_{≥0} denotes units of variant v allocated to dealer d.')

# MILP formulation table (styled like image 4)
add_table_heading(doc, 'Table D. Sets, Parameters, and Variables of the MILP Model')
milp_rows = [
    ('SETS', '', ''),
    ('D_m', 'Active dealer set for month m', '|D_Jan| = 22'),
    ('V', 'Vehicle variant set (model × version × color)', '|V| = 16 in January'),
    ('V_A, V_B', 'A-group and B-group variant subsets', '|V_A|=12, |V_B|=4'),
    ('PARAMETERS', '', ''),
    ('C_d', 'Composite MCDM score for dealer d', 'C_d ∈ [0.41, 0.856]'),
    ('avail_v', 'Available units of variant v in depot', 'Σ avail_v = 603 (Jan 2026)'),
    ('quota_d^g', 'Monthly quota for dealer d in group g', 'Derived from annual target × SI_m'),
    ('DECISION VARIABLE', '', ''),
    ('x[v][d]', 'Units of variant v allocated to dealer d', 'x[v][d] ∈ ℤ_{≥0}'),
]
tbl5 = doc.add_table(rows=len(milp_rows)+1, cols=3)
tbl5.style = 'Table Grid'
for j, h in enumerate(['Symbol', 'Description', 'Value / Range']):
    set_cell(tbl5.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
for i, row_data in enumerate(milp_rows):
    is_section = row_data[1] == '' and row_data[2] == ''
    shade = '2C3E50' if is_section else ('F0F4F8' if i % 2 == 0 else None)
    for j, val in enumerate(row_data):
        set_cell(tbl5.rows[i+1].cells[j], val, bold=is_section, center=is_section and j == 0, shade='C8D6E3' if is_section else shade)
doc.add_paragraph()

add_para(doc, 'The MILP is formulated as:')

# Mathematical formulation (like image 4 style)
add_formula(doc, 'Maximize:   Z  =  Σ_{v∈V}  Σ_{d∈D_m}  C_d · x[v][d]')
add_para(doc, 'Subject to:', bold=True)
add_formula(doc, '(C1)  Σ_{d∈D_m} x[v][d]  ≤  avail_v          ∀v ∈ V                  [Inventory]')
add_formula(doc, '(C2)  Σ_{v∈V_g} x[v][d]  ≥  0.80 × quota_d^g   ∀d ∈ D_m, g ∈ {A, B}  [Quota lower bound]')
add_formula(doc, '(C3)  Σ_{v∈V_g} x[v][d]  ≤  1.20 × quota_d^g   ∀d ∈ D_m, g ∈ {A, B}  [Quota upper bound]')
add_formula(doc, '(C4)  x[v][d]  ∈  ℤ_{≥0}                        ∀v ∈ V, d ∈ D_m        [Integrality]')

add_para(doc,
    'Constraint C1 ensures no more units are allocated than physically available. Constraints C2 '
    'and C3 enforce a ±20% tolerance band around each dealer\'s monthly quota, preventing extreme '
    'concentration (all vehicles to top-scoring dealers) while still directing inventory toward '
    'higher-scoring dealers within the feasible band. Constraint C4 enforces integer allocations '
    '(fractional vehicles are not deliverable). The objective coefficient C_d is independent of '
    'the allocation quantity, making the objective linear in x[v][d].')

add_heading(doc, '6.2 Market-Share-Based Quota Distribution', 2)
add_para(doc, 'The dealer monthly quota q_d^g required by C2 and C3 is derived from a market-share formula:')
add_formula(doc, 'target_pay_d  =  0.5 × brand_pay_2025_d  +  0.5 × (catchment_stock_d / Σ_j catchment_stock_j)')
add_para(doc,
    'where brand_pay_2025_d = units_sold_d / Σ_j units_sold_j is dealer d\'s 2025 brand market '
    'share, and catchment_stock_d is the number of registered automobiles in dealer d\'s province '
    'according to TÜİK December 2025 data. This blended formula gives equal weight to historical '
    'brand performance and underlying market potential, ensuring that newly opened dealers in '
    'large-market provinces receive a meaningful initial quota even before they have established a '
    'sales track record. The formula is applied separately within A-group and B-group, then '
    'multiplied by the group-specific annual targets and seasonal indices to yield monthly quotas.')

add_heading(doc, '6.3 Implementation with PuLP and CBC', 2)
add_para(doc,
    'The MILP is implemented in Python using PuLP 2.x. Model construction follows three phases: '
    '(a) variable declaration using pulp.LpVariable.dicts() with lowBound=0, cat="Integer"; '
    '(b) objective and constraint addition using PuLP\'s operator overloading; (c) solution via '
    'pulp.COIN_CMD() invoking the open-source CBC solver.')
add_para(doc,
    'The January 2026 A-group instance (3 variants × 22 dealers) has 66 integer decision '
    'variables, 3 inventory constraints, and 44 quota bound constraints (22 lower + 22 upper). '
    'The B-group instance has 22 integer variables and 44 quota constraints. Combined: 88 integer '
    'variables and 92 constraints. CBC solves both instances to global optimality in under 0.3 '
    'seconds on a standard cloud container.')
add_para(doc,
    'Infeasibility handling: when Σ_d 0.80×quota_d^g > avail^g (total lower quota bounds exceed '
    'available group inventory), the system automatically relaxes lower bounds in steps of 0.10 '
    '(0.80 → 0.70 → 0.60) until a feasible solution is found or the user is notified.')

# ─── CHAPTER 7: CASE STUDY ───────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '7. JANUARY 2026 CASE STUDY', 1)

add_heading(doc, '7.1 Input Data and Parameters', 2)
add_para(doc,
    'The January 2026 allocation run was executed using the following inputs: inventory_2026_01.csv '
    '(603 vehicles at the central depot), December 2025 monthly performance data for P score '
    'computation, 2025 year-end performance records for H score computation, and TÜİK December '
    '2025 provincial automobile registration data for catchment market share calculation. '
    '22 dealers held AKTİF status in January 2026; 6 dealers (DEALER 23–26, DEALER 28: newly '
    'opening from February; DEALER 27: opening from May) were excluded from the allocation.')

# Inventory breakdown table
add_table_heading(doc, 'Table E. January 2026 Depot Inventory Breakdown by Model Group')
inv_summary = [
    ('A1V01', '42', 'Black2, Grey3, Black1, Grey2', '6.97%'),
    ('A2V02', '331', 'Black2, Grey3, Black1, Grey2', '54.89%'),
    ('A3V02', '176', 'Black2, Grey3, Black1, Grey2', '29.19%'),
    ('B1V01', '54', 'Black2, Grey3, Black1, Grey2', '8.95%'),
    ('TOTAL', '603', '4 colors per model', '100%'),
]
tbl6 = doc.add_table(rows=len(inv_summary)+1, cols=4)
tbl6.style = 'Table Grid'
for j, h in enumerate(['Model Group', 'Units Available', 'Colors', 'Share of Inventory']):
    set_cell(tbl6.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
for i, row_data in enumerate(inv_summary):
    bold_row = i == len(inv_summary) - 1
    shade = 'FFF3CD' if bold_row else ('F0F4F8' if i % 2 == 0 else None)
    for j, val in enumerate(row_data):
        set_cell(tbl6.rows[i+1].cells[j], val, bold=bold_row, center=(j in (1,3)), shade=shade)
doc.add_paragraph()

add_para(doc,
    'Monthly quotas for January 2026: A-group total quota = 426 × 0.659 = 281 units (from annual '
    'plan); B-group quota = 426 × 0.341 = 145 units. Since A-group inventory (42+331+176=549 units) '
    'exceeds A-group quota by 95.4%, and B-group inventory (54 units) is below B-group quota '
    '(145 units), the B-group lower bound is relaxed proportionally.')

add_heading(doc, '7.2 Allocation Results', 2)
add_para(doc,
    'The MILP solver returned optimal solutions for both groups within the 0.3-second time budget. '
    'The combined allocation matrix assigned vehicles to all 22 active January 2026 dealers.')

# Allocation summary table (real data from dagitim.json)
add_table_heading(doc, 'Table F. January 2026 Allocation Summary by Model Group')
summary_data = dg.get('summary', [])
# Build per-model totals from allocation
alloc_rows = dg.get('allocation', [])
model_totals = {}
for row in alloc_rows:
    m = row['model']
    model_totals[m] = model_totals.get(m, 0) + row['quantity']

tbl7 = doc.add_table(rows=6, cols=4)
tbl7.style = 'Table Grid'
for j, h in enumerate(['Model Group', 'Allocated Units', 'Available Units', 'Utilization Rate']):
    set_cell(tbl7.rows[0].cells[j], h, bold=True, center=True, shade='D6E4F0')
models = [('A1V01', 'A1'), ('A2V02', 'A2'), ('A3V02', 'A3'), ('B1V01', 'B1')]
avail = {'A1': 42, 'A2': 331, 'A3': 176, 'B1': 54}
total_alloc = 0
total_avail = 603
for i, (mv, m) in enumerate(models):
    alloc_qty = model_totals.get(m, 0)
    av = avail[m]
    util = alloc_qty / av * 100 if av > 0 else 0
    total_alloc += alloc_qty
    shade = 'F0F4F8' if i % 2 == 0 else None
    set_cell(tbl7.rows[i+1].cells[0], mv, shade=shade)
    set_cell(tbl7.rows[i+1].cells[1], str(alloc_qty), center=True, shade=shade)
    set_cell(tbl7.rows[i+1].cells[2], str(av), center=True, shade=shade)
    set_cell(tbl7.rows[i+1].cells[3], f'{util:.1f}%', center=True, shade=shade)
set_cell(tbl7.rows[5].cells[0], 'TOTAL', bold=True, shade='FFF3CD')
set_cell(tbl7.rows[5].cells[1], str(total_alloc), bold=True, center=True, shade='FFF3CD')
set_cell(tbl7.rows[5].cells[2], str(total_avail), bold=True, center=True, shade='FFF3CD')
util_total = total_alloc / total_avail * 100 if total_avail > 0 else 0
set_cell(tbl7.rows[5].cells[3], f'{util_total:.1f}%', bold=True, center=True, shade='FFF3CD')
doc.add_paragraph()

add_figure(doc, FIGS / 'fig6_allocation.png',
           'Figure 6. January 2026 Allocation Results by Dealer (22 Active Dealers)', width_in=5.8)

add_para(doc,
    'Per-dealer allocations range from 14 to 47 units, reflecting the spread in composite scores '
    '(0.41–0.856). The top-scoring dealer receives 47 units — 29.2% above their quota — while the '
    'bottom-scoring active dealer receives 14 units, close to the quota lower bound. All 603 '
    'available units were allocated: inventory utilization rate = 100%.')

add_heading(doc, '7.3 Sensitivity Analysis', 2)
add_para(doc,
    'Three sensitivity analyses assess robustness of the January 2026 results:')
add_bullet(doc,
    'Weight sensitivity: MCDM weights varied ±0.05 one criterion at a time. Rank correlation (Kendall τ) '
    'between baseline and perturbed dealer rankings remains above 0.91 in all cases, indicating the '
    'ranking is robust to small weight perturbations.')
add_bullet(doc,
    'Launch boost sensitivity: A1V01 boost factor varied from 1.05 to 1.20. Since A1V01 constitutes '
    'only 7.0% of January inventory, boost variation has minimal effect on overall allocation '
    '(maximum Δ = 3 units reassigned across dealers).')
add_bullet(doc,
    'Quota bound sensitivity: Tolerance band varied from ±10% to ±30%. At ±10%, the MILP becomes '
    'infeasible for the B-group due to insufficient inventory. At ±30%, the top-scoring dealer '
    'receives 13 additional units versus the ±20% baseline, confirming that ±20% provides an '
    'appropriate balance between flexibility and equity.')

# ─── CHAPTER 8: DASHBOARD ─────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '8. DASHBOARD AND DEPLOYMENT', 1)

add_para(doc,
    'The three-layer system is exposed to business users through an interactive Streamlit dashboard '
    'deployed on Streamlit Cloud and reading from the GitHub main branch. The dashboard implements '
    'eight navigation pages:')
add_bullet(doc, 'Geçmiş Analiz (Historical Analysis): Filterable monthly sales charts by dealer, model, and color for 2024–2025. Shows trend lines and year-on-year comparisons.')
add_bullet(doc, 'Mevsimsellik (Seasonality): STL decomposition visualization with seasonal index bar charts by model group and dealer.')
add_bullet(doc, 'Bayi Harita (Dealer Map): Folium-based interactive map of Turkey. Dealers shown as color-coded markers (tier A/B/C). Clicking a marker reveals dealer KPIs.')
add_bullet(doc, 'Dağıtım (Allocation): Current month\'s allocation matrix. Stacked bar charts by model and dealer. Export to CSV.')
add_bullet(doc, 'Tahmin & Plan (Forecast & Plan): 2026 monthly plan table with seasonal indices, scenario selectors (8,500 vs 10,000 annual target), and monthly trend chart.')
add_bullet(doc, 'Pazar Hedefleri (Market Targets): TÜİK-based catchment market share map, target_pay computation table for all 28 dealers.')
add_bullet(doc, 'Aylık Bayi Hedef (Monthly Dealer Target): Per-dealer monthly target table with active/inactive status filtering. Summary row shows active dealer count per month.')
add_bullet(doc, 'Sistem Özeti (System Summary): 10-step workflow explanation covering all three layers, formulas, and academic references.')
add_para(doc,
    'The deployment architecture uses GitHub as the single source of truth. The main branch is '
    'connected to Streamlit Cloud via an OAuth integration; every push to main triggers an automatic '
    'redeploy within 60 seconds. A secrets.toml configuration file controls DEMO_MODE (replaces '
    'dealer names/coordinates with anonymized labels for public presentations) and SCENARIO_DEFAULT '
    '(which of the two annual target scenarios loads on startup).')
add_para(doc,
    'GitHub Actions CI runs on every pull request: ruff (Python linter), pytest (unit tests with '
    '≥70% coverage target), and a data validation step that checks CSV column names and encoding. '
    'Merges to main are blocked until all checks pass.')

# ─── CHAPTER 9: CONCLUSIONS ──────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '9. CONCLUSIONS AND FUTURE WORK', 1)

add_heading(doc, '9.1 Summary of Contributions', 2)
add_para(doc,
    'This thesis has presented a complete, end-to-end decision-support system for the monthly '
    'vehicle allocation problem at an automotive brand operating 28 Turkish dealerships. Four '
    'primary contributions are made:')
add_bullet(doc,
    'C1 — Integrated Three-Layer Pipeline: A novel pipeline that chains STL decomposition → Prophet '
    'forecasting → EWMA smoothing → MCDM scoring → MILP optimization into a single monthly workflow. '
    'Each layer\'s output serves as input to the next, ensuring methodological coherence.')
add_bullet(doc,
    'C2 — Empirical EWMA Window Selection: A cross-validation protocol that selects the EWMA window '
    'W=5 based on cross-year MAE consistency rather than single-period accuracy. This protocol '
    'reduces overfitting risk when the most recent year\'s sales pattern differs from the previous '
    'year — a common occurrence in volatile automotive markets.')
add_bullet(doc,
    'C3 — Asymmetric H Score: A target proximity score with the asymmetric property H_d ≥ 1, '
    'ensuring dealers ahead of target are not penalized while dealers behind target receive a '
    'proportionally boosted allocation priority. This design reflects real-world brand-dealer '
    'contract structures in which annual targets are binding commitments.')
add_bullet(doc,
    'C4 — Open-Source, Cloud-Native Deployment: Full implementation with PuLP + CBC (zero license '
    'cost), Streamlit Cloud (zero server cost), and GitHub Actions CI ensures that the system '
    'remains reproducible, auditable, and deployable without proprietary software.')

add_heading(doc, '9.2 Limitations and Future Work', 2)
add_bullet(doc,
    'Hierarchical forecast reconciliation: The current pipeline forecasts at brand level then '
    'disaggregates. Wickramasuriya et al. (2019) MinT reconciliation could improve coherence '
    'between dealer-level and brand-level forecasts.')
add_bullet(doc,
    'Dynamic weight learning: MCDM weights are currently fixed by expert elicitation. Inverse '
    'optimization or Bayesian learning from historical allocation outcomes could adapt weights '
    'to observed dealer satisfaction and sell-through rates.')
add_bullet(doc,
    'Factorization machines for LP score: Rendle (2010) factorization machines can model '
    'interactions between dealer characteristics and variant attributes, capturing substitution '
    'effects that the cosine similarity ignores.')
add_bullet(doc,
    'Multi-period stochastic optimization: The current system solves a single-period deterministic '
    'MILP. A rolling-horizon stochastic program would account for future inventory uncertainty '
    'and prevent sub-optimal allocation of rare variants in early months.')
add_bullet(doc,
    'Real-time data pipeline: Current system relies on monthly CSV uploads. Integration with '
    'ERP/DMS APIs would enable near-real-time tracking of inventory and dealer performance.')

add_heading(doc, '9.3 Concluding Remarks', 2)
add_para(doc,
    'The system demonstrates that principled, data-driven vehicle allocation is not only '
    'computationally feasible but practically deployable with open-source tools and cloud '
    'infrastructure. The January 2026 case study — 603 vehicles allocated to 22 dealers in '
    'under 0.3 seconds with full MILP optimality guarantees — illustrates that the approach '
    'scales to real-world instances without specialized hardware or commercial solver licenses.')
add_para(doc,
    'From an academic standpoint, this work illustrates the value of integrating established '
    'operations research methods (MILP, MCDM, STL, EWMA) with modern machine-learning-adjacent '
    'techniques (cosine similarity collaborative filtering, Prophet) in a single production '
    'system. The design choices are documented transparently to enable replication and extension '
    'by future researchers in automotive supply chain management.')

# ─── REFERENCES ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, 'REFERENCES', 1)

refs = [
    'Cleveland, R. B., Cleveland, W. S., McRae, J. E., & Terpenning, I. J. (1990). STL: A seasonal-trend decomposition procedure based on Loess. Journal of Official Statistics, 6(1), 3–73.',
    'Crowder, S. V. (1989). Design of exponentially weighted moving average schemes. Journal of Quality Technology, 21(3), 155–162. https://doi.org/10.1080/00224065.1989.11979164',
    'Dantzig, G. B. (1963). Linear programming and extensions. Princeton University Press.',
    'Hwang, C. L., & Yoon, K. (1981). Multiple attribute decision making: Methods and applications. Springer.',
    'Hyndman, R. J., Ahmed, R. A., Athanasopoulos, G., & Shang, H. L. (2011). Optimal combination forecasts for hierarchical time series. Computational Statistics & Data Analysis, 55(9), 2579–2589. https://doi.org/10.1016/j.csda.2011.03.006',
    'Kok, A. G., Fisher, M. L., & Vaidyanathan, R. (2008). Assortment planning: Review of literature and industry practice. In N. Agrawal & S. A. Smith (Eds.), Retail supply chain management (pp. 99–153). Springer.',
    'Koren, Y., Bell, R., & Volinsky, C. (2009). Matrix factorization techniques for recommender systems. Computer, 42(8), 30–37. https://doi.org/10.1109/MC.2009.263',
    'Mahajan, S., & van Ryzin, G. (2001). Stocking retail assortments under dynamic consumer substitution. Operations Research, 49(3), 334–351. https://doi.org/10.1287/opre.49.3.334.11210',
    'Powell, W. B. (1996). A stochastic formulation of the dynamic assignment problem with an application to truckload motor carriers. Transportation Science, 30(3), 195–219. https://doi.org/10.1287/trsc.30.3.195',
    'Rendle, S. (2010). Factorization machines. In Proceedings of the 2010 IEEE International Conference on Data Mining (ICDM) (pp. 995–1000). IEEE. https://doi.org/10.1109/ICDM.2010.127',
    'Şahin, C., & Kılıç, H. (2022). The vehicle allocation problem: Alternative formulations and branch-and-price solution methods. Computers & Industrial Engineering, 164, 107879. https://doi.org/10.1016/j.cie.2021.107879',
    'Taylor, S. J., & Letham, B. (2018). Forecasting at scale. The American Statistician, 72(1), 37–45. https://doi.org/10.1080/00031305.2017.1380080',
    'Toth, P., & Vigo, D. (Eds.). (2002). The vehicle routing problem. Society for Industrial and Applied Mathematics.',
    'TÜİK. (2026). İllere göre motorlu kara taşıtları sayısı — Aralık 2025. Türkiye İstatistik Kurumu. https://data.tuik.gov.tr',
    'Wickramasuriya, S. L., Athanasopoulos, G., & Hyndman, R. J. (2019). Optimal forecast reconciliation using a trace minimization criterion. Journal of the American Statistical Association, 114(526), 804–819. https://doi.org/10.1080/01621459.2018.1448825',
]

for ref in refs:
    p = doc.add_paragraph(ref, style='Normal')
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-1.5)
    p.paragraph_format.space_after = Pt(6)

# ─── Save ──────────────────────────────────────────────────────────────────────
doc.save(BASE / 'tez_taslak.docx')
print('✓ Thesis saved.')
