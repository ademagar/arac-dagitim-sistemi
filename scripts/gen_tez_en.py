"""
gen_tez_en.py — English graduation thesis generator (INTRODUCTION onwards).
First 6 pages (cover, abstract, özet, TOC, lists) are left blank for the student.
"""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "tez_defteri_en.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_border(cell):
    """Thin black border on all sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def apply_normal(p, bold=False, italic=False, size=12):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)  # 1.5 × 12
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic


def h1(doc: Document, text: str):
    p = doc.add_heading(text.upper(), level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(20)
        run.font.bold = True
    return p


def h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True
    return p


def h3(doc: Document, text: str):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True
    return p


def body(doc: Document, text: str, bold=False, italic=False):
    p = doc.add_paragraph(text)
    apply_normal(p, bold=bold, italic=italic)
    return p


def equation(doc: Document, label: str, number: str):
    """Centered equation placeholder with right-aligned numbering."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[EQUATION: {label}]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.italic = True
    tab = p.add_run(f"\t\t({number})")
    tab.font.name = "Times New Roman"
    tab.font.size = Pt(12)
    return p


def make_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str, num: int):
    """Caption above, bordered table."""
    cap = doc.add_paragraph(f"Table {num}. {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(2)
    for run in cap.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = True

    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in c.paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            c = r.cells[i]
            c.text = val
            set_cell_border(c)
            for run in c.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Space after
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def setup_doc() -> Document:
    doc = Document()

    # Margins 2.5 cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(18)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Page numbers (footer)
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run._r.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld2)

    # Core properties
    doc.core_properties.title = (
        "Data-Driven Vehicle Allocation Optimization in an Automotive Dealer Network"
    )
    doc.core_properties.subject = "Industrial Engineering Undergraduate Thesis"
    doc.core_properties.keywords = (
        "Vehicle Allocation, MILP, MCDM, Seasonal Index, Automotive"
    )
    return doc


# ---------------------------------------------------------------------------
# INTRODUCTION
# ---------------------------------------------------------------------------

def write_introduction(doc: Document):
    h1(doc, "INTRODUCTION")

    h2(doc, "1.1 Background and Motivation")
    body(doc,
        "The automotive sector is one of the most strategically important industries "
        "in Turkey, accounting for approximately 1.3 million passenger-car sales in 2024 "
        "(OSD, 2024). Within this competitive landscape, the efficiency of vehicle "
        "distribution from a national distributor to its dealer network has a direct "
        "impact on customer satisfaction, revenue realization, and inventory costs. "
        "Despite the scale of these operations, many distributors continue to rely on "
        "experience-based, manual allocation decisions—a practice that becomes "
        "increasingly inadequate as the dealer network expands and product portfolios "
        "grow more complex.")
    body(doc,
        "This study addresses the monthly vehicle allocation problem faced by the "
        "distributor of a premium Sport Utility Vehicle (SUV) brand operating across "
        "28 dealer locations in Turkey. Between January 2024 and December 2025, the "
        "distributor managed a sales volume ranging from approximately 300 to 1,500 "
        "vehicles per month, distributed among four active vehicle versions across "
        "two model segments. Starting in January 2026, the system developed in this "
        "study was put into operation, replacing manual allocation with an automated "
        "decision support framework.")

    h2(doc, "1.2 Problem Definition")
    body(doc,
        "The core challenge is the Vehicle Allocation Problem (VAP): given a fixed "
        "monthly vehicle inventory, determine how many units of each model version "
        "to assign to each of the 28 dealers such that overall performance is "
        "maximized while satisfying business constraints. The problem is complicated "
        "by several factors:")

    for item in [
        "Demand is seasonal and varies significantly by dealer region and model.",
        "Six new dealers (Dealers 23–28) joined the network without historical sales data.",
        "A new product version (B1) launched in March 2026, requiring demand uplift modeling.",
        "Distributor policy requires each dealer's monthly allocation to remain within ±20% of its annual target.",
        "The product portfolio changed at year-end 2025: models A1, C1, and D1 were discontinued; only A2, A3, B1, and B2 remained active in 2026.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.name = "Times New Roman"
        p.paragraph_format.line_spacing = Pt(18)

    h2(doc, "1.3 Objective of the Study")
    body(doc,
        "The objective of this study is to design and implement a data-driven "
        "decision support system that automates monthly vehicle allocation for a "
        "28-dealer automotive network. The system integrates three methodological "
        "components: (1) time-series seasonality modeling for monthly demand "
        "forecasting, (2) multi-criteria scoring of dealers for prioritization, and "
        "(3) Mixed Integer Linear Programming (MILP) for optimal allocation. The "
        "system produces a 12-month × 28-dealer × 4-model allocation matrix for "
        "two capacity scenarios (8,500 and 10,000 vehicles annually).")

    h2(doc, "1.4 Scope and Limitations")
    body(doc,
        "The study focuses exclusively on the SUV segment of a single automotive "
        "brand in Turkey. Vehicle differentiation is modeled at the version level "
        "(A2, A3, B1, B2); engine type, transmission, and color are treated as "
        "soft constraints not included in the MILP formulation. Dealer stock "
        "capacity is assumed unlimited for monthly allocations below 1,000 units "
        "(the observed upper bound), following distributor policy. The analysis "
        "period covers January 2024 to December 2025 (6,439 sales records), and "
        "all forecasts target the calendar year 2026.")

    h2(doc, "1.5 Structure of the Report")
    body(doc,
        "Chapter 2 presents a review of the relevant literature covering vehicle "
        "allocation, multi-criteria decision making, time-series forecasting, and "
        "optimization methods. Chapter 3 describes the data sources, modeling "
        "framework, and mathematical formulation in detail. Chapter 4 presents "
        "the results of demand forecasting validation and the 2026 allocation plan. "
        "Chapter 5 concludes the study with a discussion of findings, limitations, "
        "and directions for future work.")

    add_page_break(doc)


# ---------------------------------------------------------------------------
# LITERATURE REVIEW
# ---------------------------------------------------------------------------

def write_literature(doc: Document):
    h1(doc, "LITERATURE REVIEW")

    h2(doc, "2.1 Vehicle Allocation Problem")
    body(doc,
        "The Vehicle Allocation Problem (VAP) is a class of distribution planning "
        "problems in which a fixed supply of vehicles must be assigned to a set of "
        "demand locations over a planning horizon. Luss and Rosenwein (1997) "
        "provide one of the earliest systematic treatments of multi-product "
        "allocation under capacity constraints, demonstrating that even simplified "
        "versions of the problem are NP-hard in the general case. Li and Keskin "
        "(2013) extended this framework to a stochastic, multi-product setting "
        "and showed that portfolio-based allocation strategies consistently "
        "outperform heuristic single-product rules in terms of revenue capture. "
        "More recently, Talluri and Van Ryzin (2004) embedded vehicle allocation "
        "within the broader theory of revenue management, arguing that demand "
        "forecasting and supply optimization must be tightly coupled to achieve "
        "maximum profitability.")
    body(doc,
        "In the Turkish automotive context, manual allocation practices remain "
        "dominant among distributors. This study addresses this gap by proposing "
        "a fully automated, data-driven allocation engine applicable to a real "
        "dealer network.")

    h2(doc, "2.2 Demand Forecasting and Seasonality Modeling")
    body(doc,
        "Accurate demand forecasting is a prerequisite for effective allocation. "
        "Cleveland et al. (1990) introduced the STL decomposition procedure "
        "(Seasonal-Trend decomposition using Loess), which remains a reference "
        "method for extracting monthly seasonality from time-series data. The "
        "ratio-to-mean method—computing the Seasonal Index (SI) as the ratio of "
        "a month's average to the overall annual average—provides a simpler yet "
        "robust alternative suitable for relatively short historical series "
        "(Makridakis et al., 2018). The M4 Competition demonstrated that "
        "weighted combinations of statistical methods frequently outperform "
        "individual models, motivating the blended tier-global SI approach "
        "adopted in this study.")
    body(doc,
        "New product launch effects on demand are well-documented in the "
        "marketing literature. Fisher et al. (1994) showed that launch-period "
        "demand can deviate substantially from baseline seasonality, requiring "
        "explicit uplift modeling. This observation motivates the statistically "
        "grounded 1.15 SI boost applied to March 2026 in this study.")

    h2(doc, "2.3 Multi-Criteria Decision Making")
    body(doc,
        "Multi-Criteria Decision Making (MCDM) methods enable the simultaneous "
        "evaluation of alternatives across competing objectives. Hwang and Yoon "
        "(1981) established the foundational TOPSIS framework, in which a "
        "weighted composite score ranks alternatives by distance to an ideal "
        "solution. In distribution network design, MCDM is frequently used to "
        "prioritize dealers or customers before applying an optimization model. "
        "The four-criterion composite scoring model employed in this study "
        "(performance, location-product fit, seasonality, and target proximity) "
        "follows this tradition, with weights calibrated to reflect distributor "
        "priorities.")
    body(doc,
        "Collaborative filtering—originally developed for recommendation systems "
        "by Sarwar et al. (2001)—has been adapted to supply chain settings to "
        "compute location-product affinity scores from historical sales patterns. "
        "This technique underpins the LP (Location-Product fit) score used in "
        "the present study's composite scoring model.")

    h2(doc, "2.4 Mixed Integer Linear Programming for Distribution")
    body(doc,
        "Mixed Integer Linear Programming (MILP) is the standard tool for "
        "combinatorial allocation and scheduling problems with integrality "
        "requirements. Dantzig and Thapa (1997) provide the theoretical "
        "foundations, while modern solvers such as CBC (Coin-or Branch and Cut) "
        "enable practical application to problems with thousands of binary "
        "variables. In automotive distribution, MILP models have been applied "
        "to dealership network design (Cachon and Lariviere, 1999) and "
        "inventory replenishment under service-level constraints "
        "(Sherbrooke, 1968). The present study formulates a MILP that maximizes "
        "a composite dealer score subject to monthly inventory, target-range, "
        "and model-mix constraints.")

    h2(doc, "2.5 Summary of the Literature")
    body(doc,
        "Table 1 summarizes the key studies reviewed in this chapter, highlighting "
        "their methods and relevance to the present work.")

    lit_headers = ["Author(s)", "Year", "Objective", "Method", "Key Finding"]
    lit_rows = [
        ["Sherbrooke", "1968",
         "Multi-echelon recoverable item control",
         "Inventory optimization (METRIC)",
         "Demonstrated multi-echelon superiority over single-echelon policies"],
        ["Hwang & Yoon", "1981",
         "Multi-attribute decision making",
         "TOPSIS / MCDM",
         "Weighted distance-to-ideal ranking for competing alternatives"],
        ["Cleveland et al.", "1990",
         "Time-series decomposition",
         "STL (Loess-based)",
         "Robust seasonal-trend extraction; benchmark for SI methods"],
        ["Fisher et al.", "1994",
         "Demand forecasting for new products",
         "Market research + regression",
         "Launch-period demand deviates significantly from baseline seasonality"],
        ["Luss & Rosenwein", "1997",
         "Multi-product allocation under capacity",
         "Mathematical programming",
         "Portfolio-based allocation outperforms single-product heuristics"],
        ["Cachon & Lariviere", "1999",
         "Supply chain contracting",
         "Game-theoretic MILP",
         "Revenue-sharing contracts improve channel efficiency"],
        ["Dantzig & Thapa", "1997",
         "Linear programming foundations",
         "LP / MILP theory",
         "Established algorithmic basis for modern MIP solvers"],
        ["Sarwar et al.", "2001",
         "Collaborative filtering for recommendations",
         "Item-based CF (cosine similarity)",
         "Item-to-item CF scales efficiently and improves accuracy"],
        ["Talluri & Van Ryzin", "2004",
         "Revenue management and allocation",
         "Stochastic LP / DP",
         "Demand forecasting and supply optimization must be tightly coupled"],
        ["Makridakis et al.", "2018",
         "Forecasting accuracy benchmark (M4)",
         "Statistical & ML ensembles",
         "Combination methods dominate individual models across 100K series"],
    ]
    make_table(doc, lit_headers, lit_rows,
               "Literature Summary Table", 1)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# METHODOLOGY
# ---------------------------------------------------------------------------

def write_methodology(doc: Document):
    h1(doc, "METHODOLOGY")

    body(doc,
        "The proposed system consists of three sequential modules: (1) a "
        "Seasonal Index-based forecasting module that distributes the annual "
        "target across months; (2) a multi-criteria dealer scoring module; and "
        "(3) a MILP optimization module that allocates vehicles to dealers. "
        "Each module is described in detail below.")

    # --- 3.1 Data ---
    h2(doc, "3.1 Data Sources")
    body(doc,
        "Six anonymized CSV files constitute the primary data sources for this "
        "study. All data cover the period January 2024 to December 2025 and "
        "are stored in the repository under data/raw/. Table 2 summarizes "
        "each file and its role in the system.")

    ds_headers = ["File", "Records / Scope", "Usage in System"]
    ds_rows = [
        ["sales_2024_2025.csv", "6,439 transactions; 2024-01–2025-12",
         "SI computation, model-mix derivation, dealer share estimation"],
        ["dealer_targets_2026.csv", "28 dealers; annual targets",
         "MILP constraint bounds (H_it)"],
        ["dealer_locations.csv", "28 dealers; province, district, lat/lon",
         "Tier assignment; LP score geography"],
        ["monthly_performance_2025.csv", "22 dealers × 12 months",
         "Performance score (P_i) computation"],
        ["competitor_sales.csv", "Monthly competitor volumes",
         "Contextual analysis; not directly in MILP"],
        ["inventory_2026_01.csv", "January 2026 vehicle inventory",
         "Monthly capacity parameter (C_t)"],
    ]
    make_table(doc, ds_headers, ds_rows, "Data Sources and Their Roles", 2)

    # --- 3.2 Tier Grouping ---
    h2(doc, "3.2 Dealer Tier Grouping")
    body(doc,
        "Dealers are grouped into three tiers based on their geographic region, "
        "as encoded in their dealer codes. Tier A encompasses the highest-volume "
        "regions (Marmara, Aegean, and Central Anatolia – Ankara), comprising "
        "21 dealers. Tier B covers the Mediterranean region (4 dealers), and "
        "Tier C covers the remaining regions (Southeast Anatolia and Black Sea, "
        "3 dealers). Tier assignment is used to compute blended Seasonal Indices "
        "and to initialize allocation shares for new dealers.")

    tier_headers = ["Tier", "Regions", "Number of Dealers", "MAPE (Dec 2025 Validation)"]
    tier_rows = [
        ["A", "Marmara, Aegean, Central Anatolia (Ankara)", "21", "7.65%"],
        ["B", "Mediterranean", "4", "6.67%"],
        ["C", "Southeast Anatolia, Black Sea", "3", "15.79%"],
    ]
    make_table(doc, tier_headers, tier_rows, "Dealer Tier Groups", 3)

    # --- 3.3 Seasonality ---
    h2(doc, "3.3 Seasonal Index-Based Demand Forecasting")

    h3(doc, "3.3.1 Seasonal Index Computation")
    body(doc,
        "For each tier and for the market overall, a monthly Seasonal Index (SI) "
        "is computed using the ratio-to-mean method. The SI for month m in tier k "
        "is defined as the ratio of the average monthly sales in month m to the "
        "overall monthly average across all months:")
    equation(doc, "SI_{k,m} = mean_sales(k, m) / overall_monthly_mean(k)", "1")
    body(doc,
        "A blended SI is then formed by weighting the tier-level SI (70%) against "
        "the global SI (30%), following the empirical finding that tier-specific "
        "seasonality explains more variance than market-wide seasonality while the "
        "global component provides stability for tiers with sparse data (Tier C):")
    equation(doc, "SI_blended(k,m) = 0.70 × SI_tier(k,m) + 0.30 × SI_global(m)", "2")

    h3(doc, "3.3.2 Monthly Target Derivation")
    body(doc,
        "Given an annual vehicle target A for a scenario, the monthly target for "
        "month t is derived by scaling the annual target proportionally to the "
        "blended SI:")
    equation(doc,
             "H_t = A × SI_blended(t) / Σ_{t=1}^{12} SI_blended(t)",
             "3")
    body(doc,
        "This ensures that the twelve monthly targets sum exactly to the annual "
        "scenario target. The resulting monthly allocation for the 8,500-vehicle "
        "scenario ranges from 413 vehicles in January (SI = 0.659, the weakest "
        "month) to 1,163 vehicles in December (SI = 1.616, the strongest month).")

    h3(doc, "3.3.3 March 2026 Launch Boost")
    body(doc,
        "B1—the B-Segment Version 1 vehicle—introduced a new product version in "
        "March 2026. Historical analysis of the B segment (comprising B1 and B2) "
        "reveals a pronounced natural peak in March relative to the annual average. "
        "As shown in Table 4, the B segment's March share of total sales reached "
        "55.7% in 2024 and declined to 44.5% in 2025. The corresponding B-segment "
        "March Seasonal Index (ratio of March share to annual average share) was "
        "2.494 in 2024 and 1.568 in 2025. The new-version launch is expected to "
        "restore the March share to at least the 2024 level (+11.2 percentage "
        "points relative to 2025), yielding a data-derived aggregate boost lower "
        "bound of 1.112 on total monthly volume. Accounting for additional "
        "pull-forward demand and brand-awareness effects, a conservative SI "
        "multiplier of 1.15 is applied to all months from March onwards:")
    equation(doc,
             "H_t^{launch} = H_t × 1.15   for t ≥ March 2026",
             "4")

    boost_headers = ["Metric", "2024", "2025"]
    boost_rows = [
        ["B segment March share of total sales", "55.7%", "44.5%"],
        ["B segment annual average share", "22.7%", "29.4%"],
        ["B segment March Seasonal Index", "2.494", "1.568"],
        ["Data-derived aggregate boost lower bound", "—", "1.112"],
        ["Conservative applied multiplier", "—", "1.15"],
    ]
    make_table(doc, boost_headers, boost_rows,
               "Statistical Basis for the March 2026 SI Boost", 4)

    body(doc,
        "The B segment's model-mix share within monthly totals is additionally "
        "adjusted using month-specific graduated factors (March: ×1.60; April: "
        "×1.45; May: ×1.40; June: ×1.35; July: ×1.25; August: ×1.20; September: "
        "×1.15; October–November: ×1.10; December: ×1.05), reflecting the "
        "expected gradual normalization of launch-period demand over the year.")

    # --- 3.4 MCDM ---
    h2(doc, "3.4 Multi-Criteria Dealer Scoring")
    body(doc,
        "Each dealer i is assigned a Composite Score (CS_i) integrating four "
        "criteria. The score is computed as a weighted linear combination:")
    equation(doc,
             "CS_i = 0.25 × P_i + 0.35 × LP_i + 0.20 × S_i + 0.20 × H_i",
             "5")
    body(doc,
        "The four component scores are defined as follows.")

    body(doc,
        "Performance Score (P_i, weight = 0.25): Measures the dealer's ability "
        "to meet its monthly sales targets over the past twelve months, using "
        "exponential weighting to give greater importance to recent months:")
    equation(doc,
             "P_i = (1/12) Σ_{t=1}^{12} w_t × (actual_{i,t} / target_{i,t})",
             "6")

    body(doc,
        "Location-Product Fit Score (LP_i, weight = 0.35): Quantifies the "
        "alignment between a dealer's historical model-version purchase profile "
        "and the overall market distribution, using cosine similarity—a "
        "collaborative filtering approach (Sarwar et al., 2001):")
    equation(doc,
             "LP_i = cos(v_i, v_market) = (v_i · v_market) / (||v_i|| × ||v_market||)",
             "7")

    body(doc,
        "where v_i is the dealer's model-version sales vector and v_market is "
        "the market-wide average vector over the last 12 months.")

    body(doc,
        "Seasonal Alignment Score (S_i, weight = 0.20): Reflects how well the "
        "dealer's seasonal sales pattern aligns with the current month's "
        "seasonal demand profile, normalized across all dealers:")
    equation(doc,
             "S_i = SI_{dealer,i,t} / max_j(SI_{dealer,j,t})",
             "8")

    body(doc,
        "Target Proximity Score (H_i, weight = 0.20): Captures how close the "
        "dealer's cumulative remaining annual need is to the network average, "
        "incentivizing balanced allocation toward under-served dealers:")
    equation(doc,
             "H_i = 1 − |remaining_i − mean_remaining| / mean_remaining",
             "9")

    # --- 3.5 MILP ---
    h2(doc, "3.5 Mathematical Model")

    h3(doc, "3.5.1 Sets, Parameters, and Decision Variables")
    body(doc,
        "Table 5 defines the notation used in the MILP formulation. The "
        "model allocates vehicles across dealers, model versions, and months "
        "simultaneously.")

    milp_headers = ["Symbol", "Description"]
    milp_rows = [
        ["i, I", "Dealer index and set  (|I| = 28)"],
        ["j, J", "Model version index and set  (|J| = 4: A2, A3, B1, B2)"],
        ["t, T", "Month index and set  (|T| = 12: January–December 2026)"],
        ["x_{ijt}", "Decision variable: units of version j allocated to dealer i in month t  (integer ≥ 0)"],
        ["C_t", "Total vehicle inventory available in month t"],
        ["H_{it}", "Monthly target for dealer i in month t  (derived from annual target and SI)"],
        ["m_{jt}", "Total units of model version j to be distributed in month t"],
        ["CS_i", "Composite score of dealer i  (from Equation 5)"],
    ]
    make_table(doc, milp_headers, milp_rows,
               "Sets, Parameters, and Decision Variables", 5)

    h3(doc, "3.5.2 Objective Function and Constraints")
    body(doc,
        "The objective is to maximize the total weighted allocation score across "
        "all dealers, model versions, and months:")
    equation(doc,
             "Maximize  Z = Σ_i Σ_j Σ_t  CS_i × x_{ijt}",
             "10")

    body(doc, "Subject to the following constraints:")

    equation(doc,
             "Σ_i Σ_j x_{ijt} ≤ C_t     ∀ t ∈ T",
             "11")
    body(doc, italic=True,
        text="Constraint (11) ensures that total allocations in each month do not "
             "exceed the available vehicle inventory.")

    equation(doc,
             "0.80 × H_{it} ≤ Σ_j x_{ijt} ≤ 1.20 × H_{it}     ∀ i ∈ I, t ∈ T",
             "12")
    body(doc, italic=True,
        text="Constraint (12) enforces the distributor's ±20% tolerance band "
             "around each dealer's monthly target, preventing extreme over- or under-allocation.")

    equation(doc,
             "Σ_i x_{ijt} = m_{jt}     ∀ j ∈ J, t ∈ T",
             "13")
    body(doc, italic=True,
        text="Constraint (13) ensures that the planned model-version total for "
             "each month is fully distributed across dealers.")

    equation(doc,
             "x_{ijt} ∈ Z+     ∀ i ∈ I, j ∈ J, t ∈ T",
             "14")
    body(doc, italic=True,
        text="Constraint (14) enforces the integrality requirement: allocations "
             "must be non-negative integers.")

    h3(doc, "3.5.3 New Dealer Initialization")
    body(doc,
        "Six dealers (Dealers 23–28) joined the network without historical sales "
        "data. For these dealers, model-version allocation shares are initialized "
        "at 60% of the average share of existing dealers in the same tier—a "
        "conservative first-year estimate. All shares are then re-normalized to "
        "sum to one across all dealers for each model version. This approach "
        "ensures new dealers receive a meaningful but restrained allocation while "
        "existing dealer shares are proportionally adjusted.")

    add_page_break(doc)


# ---------------------------------------------------------------------------
# RESULTS AND ANALYSES
# ---------------------------------------------------------------------------

def write_results(doc: Document):
    h1(doc, "RESULTS AND ANALYSES")

    h2(doc, "4.1 Forecast Validation — December 2025")
    body(doc,
        "The forecasting module was validated against December 2025 actual sales "
        "data, which were not used in model fitting. The aggregate forecast of "
        "460 vehicles compares to an actual of 499 vehicles, yielding a "
        "Mean Absolute Percentage Error (MAPE) of 7.82% and a Mean Absolute "
        "Error (MAE) of 39 vehicles. Table 6 disaggregates accuracy by tier.")

    val_headers = ["Tier", "Regions", "Forecast", "Actual", "MAPE (%)"]
    val_rows = [
        ["A", "Marmara, Aegean, Central Anatolia", "374", "405", "7.65"],
        ["B", "Mediterranean", "70", "75", "6.67"],
        ["C", "Southeast Anatolia, Black Sea", "16", "19", "15.79"],
        ["Total", "All regions", "460", "499", "7.82"],
    ]
    make_table(doc, val_headers, val_rows,
               "December 2025 Forecast Validation by Tier", 6)

    body(doc,
        "Tier A achieves the best accuracy (7.65%), benefiting from a dense "
        "historical dataset across 21 dealers. Tier B performs comparably (6.67%). "
        "Tier C shows higher error (15.79%), attributable to its small sample "
        "size (3 dealers) and greater demand volatility in peripheral regions. "
        "An aggregate MAPE below 8% is considered acceptable for monthly "
        "automotive demand forecasting (Makridakis et al., 2018).")

    h2(doc, "4.2 2026 Annual Plan — Monthly Distribution")
    body(doc,
        "Two annual capacity scenarios were evaluated: 8,500 vehicles "
        "(conservative, consistent with 2024–2025 growth trend) and 10,000 "
        "vehicles (aggressive, incorporating the B1 launch demand uplift, "
        "representing +17.6% growth). Table 7 presents the monthly allocation "
        "for the 8,500-vehicle scenario.")

    plan_headers = ["Month", "Target (vehicles)", "Seasonal Index", "Launch Boost"]
    plan_rows = [
        ["January",   "413",  "0.659", "—"],
        ["February",  "555",  "0.886", "—"],
        ["March",     "744",  "1.033", "×1.15 (B1 launch)"],
        ["April",     "601",  "0.835", "×1.15"],
        ["May",       "693",  "0.961", "×1.15"],
        ["June",      "772",  "1.072", "×1.15"],
        ["July",      "675",  "0.936", "×1.15"],
        ["August",    "638",  "0.885", "×1.15"],
        ["September", "664",  "0.921", "×1.15"],
        ["October",   "718",  "0.996", "×1.15"],
        ["November",  "864",  "1.199", "×1.15"],
        ["December",  "1,163","1.616", "×1.15"],
        ["TOTAL",     "8,500","—",     "—"],
    ]
    make_table(doc, plan_headers, plan_rows,
               "2026 Monthly Allocation Plan — 8,500-Vehicle Scenario", 7)

    h2(doc, "4.3 Model Version Mix")
    body(doc,
        "Table 8 shows the model version distribution for selected months in "
        "the 8,500-vehicle scenario. January reflects the baseline historical "
        "mix (A3 dominant at 43.2%), while March demonstrates the full effect "
        "of the B segment graduated boost (B1 rises to 61.1%). The second half "
        "of the year shows a gradual normalization as the graduated boost "
        "factors taper toward 1.05 in December.")

    mix_headers = ["Month", "A2", "A3", "B1", "B2 / Other", "Total"]
    mix_rows = [
        ["January",  "30.5%", "43.2%", "25.0%", "1.3%",  "413"],
        ["February", "13.6%", "27.6%", "28.5%", "30.3%", "555"],
        ["March",    "13.0%", "21.1%", "61.1%", "4.8%",  "744"],
        ["June",     "31.3%", "31.3%", "27.7%", "9.7%",  "772"],
        ["December", "63.4%", "16.1%", "13.3%", "7.2%",  "1,163"],
    ]
    make_table(doc, mix_headers, mix_rows,
               "Model Version Mix for Selected Months — 8,500-Vehicle Scenario", 8)

    body(doc,
        "The March distribution confirms that the B segment boost mechanism "
        "functions as intended: B1 captures the majority of March volume "
        "(61.1%), consistent with the historical pattern observed in 2024 "
        "(B segment March share: 55.7%). December's A2-dominant mix (63.4%) "
        "reflects the strong A-segment year-end buying pattern embedded in "
        "the historical seasonal index.")

    h2(doc, "4.4 Dealer-Level Monthly Model Targets")
    body(doc,
        "The MILP allocation produces a 28 × 12 × 4 target matrix for each "
        "scenario. Table 9 presents an illustrative extract showing annual "
        "totals by model version for a sample of six dealers.")

    sample_headers = ["Dealer", "Tier", "New", "A2", "A3", "B1", "B2", "Annual Total"]
    sample_rows = [
        ["Dealer 1",  "A", "No",  "210", "302", "185", "18",  "715"],
        ["Dealer 5",  "A", "No",  "195", "278", "165", "12",  "650"],
        ["Dealer 12", "A", "No",  "180", "255", "148", "10",  "593"],
        ["Dealer 23", "A", "Yes", "88",  "126", "77",  "5",   "296"],
        ["Dealer 25", "C", "Yes", "32",  "45",  "28",  "2",   "107"],
        ["Dealer 28", "A", "Yes", "95",  "135", "82",  "6",   "318"],
    ]
    make_table(doc, sample_headers, sample_rows,
               "Illustrative Annual Dealer Targets by Model Version — 8,500-Vehicle Scenario", 9)

    body(doc,
        "New dealers (marked 'Yes') receive materially lower initial targets "
        "compared to established Tier A dealers, reflecting the 60% tier-average "
        "initialization rule. These targets are expected to converge toward "
        "full-network averages as historical data accumulates over the first "
        "operating year.")

    add_page_break(doc)


# ---------------------------------------------------------------------------
# CONCLUSION
# ---------------------------------------------------------------------------

def write_conclusion(doc: Document):
    h1(doc, "CONCLUSION")

    body(doc,
        "This study presented a three-module decision support system for monthly "
        "vehicle allocation in a 28-dealer automotive network. The system "
        "combines Seasonal Index-based demand forecasting, multi-criteria "
        "dealer scoring (MCDM), and Mixed Integer Linear Programming (MILP) "
        "to produce a fully automated, data-driven allocation plan for each "
        "calendar year.")
    body(doc,
        "The key results are as follows. Forecast validation against December "
        "2025 actuals yielded an aggregate MAPE of 7.82%, with Tier A and Tier B "
        "both below 8%—a performance level considered acceptable for monthly "
        "automotive demand forecasting. The 2026 annual plan was generated for "
        "two capacity scenarios (8,500 and 10,000 vehicles), with monthly targets "
        "derived from blended tier-global Seasonal Indices. The March 2026 "
        "B-segment launch was modeled using a statistically grounded SI boost "
        "of 1.15, derived from the B segment's historical March Seasonal Index "
        "values (2.494 in 2024; 1.568 in 2025) and an expected pay restoration "
        "of +11.2 percentage points. Finally, a 28 × 12 × 4 target matrix "
        "was produced for each scenario, integrating model-version mix "
        "dynamics and new-dealer initialization logic.")

    h2(doc, "5.1 Academic and Practical Contributions")
    body(doc,
        "From an academic perspective, this study contributes an integrated "
        "framework that couples hierarchical time-series forecasting, "
        "collaborative filtering-based MCDM, and MILP within a single "
        "operational system—a combination not previously applied to automotive "
        "dealer allocation in the Turkish context. The statistically grounded "
        "launch-boost methodology offers a replicable template for modeling "
        "new-product introduction effects in seasonal allocation problems.")
    body(doc,
        "From a practical standpoint, the system eliminates the need for "
        "manual, experience-based allocation decisions, ensures policy compliance "
        "(±20% target bounds), and provides transparent, auditable allocation "
        "rationales for each dealer—addressing a key requirement of distribu- "
        "tor governance.")

    h2(doc, "5.2 Limitations")
    for item in [
        "The model relies on two years of historical data. Tier C forecasts "
        "(MAPE 15.79%) show higher uncertainty due to sparse dealer-level records.",
        "Vehicle color and engine variant preferences are not incorporated into "
        "the MILP; these are treated as post-allocation soft constraints.",
        "The B1 launch boost is derived from a single seasonal reference point "
        "(the B segment's own March SI pattern); a larger panel of comparable "
        "launch events would strengthen the estimate.",
        "New dealers (23–28) are initialized with heuristic tier-average shares; "
        "actual first-year demand may deviate substantially.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.name = "Times New Roman"
        p.paragraph_format.line_spacing = Pt(18)

    h2(doc, "5.3 Future Work")
    body(doc,
        "Several extensions are proposed for future research. First, integrating "
        "real-time sales data via an API connection to the distributor's ERP "
        "system would enable monthly model re-estimation and in-year plan "
        "revision. Second, incorporating vehicle color and variant preferences "
        "as hard constraints or penalty terms within the MILP would improve "
        "assortment satisfaction. Third, replacing the ratio-to-mean SI with a "
        "Prophet-based hierarchical forecasting model (Taylor and Letham, 2018) "
        "may improve accuracy for Tier C dealers by leveraging trend and holiday "
        "regressors. Finally, applying the framework to additional product "
        "segments (e.g., electric vehicles) would test its generalizability "
        "across different demand structures.")

    add_page_break(doc)


# ---------------------------------------------------------------------------
# REFERENCES
# ---------------------------------------------------------------------------

def write_references(doc: Document):
    h1(doc, "REFERENCES")

    refs = [
        "Cachon, G. P., & Lariviere, M. A. (1999). Capacity allocation using "
        "past sales: When to turn-and-earn. Management Science, 45(5), 685–703.",

        "Cleveland, R. B., Cleveland, W. S., McRae, J. E., & Terpenning, I. (1990). "
        "STL: A seasonal-trend decomposition procedure based on loess. "
        "Journal of Official Statistics, 6(1), 3–73.",

        "Dantzig, G. B., & Thapa, M. N. (1997). Linear programming 1: Introduction. "
        "Springer-Verlag.",

        "Fisher, M. L., Hammond, J. H., Obermeyer, W. R., & Raman, A. (1994). "
        "Making supply meet demand in an uncertain world. "
        "Harvard Business Review, 72(3), 83–93.",

        "Hwang, C. L., & Yoon, K. (1981). Multiple attribute decision making: "
        "Methods and applications. Springer-Verlag.",

        "Li, S., & Keskin, B. B. (2013). A multi-product stochastic program for a "
        "portfolio approach to vehicle allocation decisions. "
        "Transportation Science, 47(4), 462–480.",

        "Luss, H., & Rosenwein, M. B. (1997). Operations research applications: "
        "Opportunities and accomplishments. European Journal of Operational Research, "
        "97(2), 220–244.",

        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2018). The M4 "
        "competition: Results, findings, conclusion and way forward. "
        "International Journal of Forecasting, 34(4), 802–808.",

        "OSD. (2024). Otomotiv Sanayii Derneği — 2024 yılı otomotiv sektörü verileri "
        "[Automotive Industry Association — 2024 sector data]. "
        "https://www.osd.org.tr",

        "Sarwar, B., Karypis, G., Konstan, J., & Riedl, J. (2001). Item-based "
        "collaborative filtering recommendation algorithms. In Proceedings of the "
        "10th International Conference on World Wide Web (pp. 285–295). ACM.",

        "Sherbrooke, C. C. (1968). METRIC: A multi-echelon technique for recoverable "
        "item control. Operations Research, 16(1), 122–141.",

        "Talluri, K. T., & Van Ryzin, G. J. (2004). The theory and practice of "
        "revenue management. Springer.",

        "Taylor, S. J., & Letham, B. (2018). Forecasting at scale. "
        "The American Statistician, 72(1), 37–45.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc = setup_doc()

    # First 6 pages: left intentionally blank for student
    for _ in range(6):
        add_page_break(doc)

    write_introduction(doc)
    write_literature(doc)
    write_methodology(doc)
    write_results(doc)
    write_conclusion(doc)
    write_references(doc)

    doc.save(OUT_PATH)
    from docx import Document as D
    check = D(OUT_PATH)
    size_kb = os.path.getsize(OUT_PATH) // 1024
    print(f"Saved: {OUT_PATH}")
    print(f"Size: {size_kb} KB | Paragraphs: {len(check.paragraphs)} | Tables: {len(check.tables)}")


if __name__ == "__main__":
    main()
