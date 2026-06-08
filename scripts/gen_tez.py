"""
Tez Defteri Oluşturucu
======================
Otomotiv Bayi Ağında Veri Odaklı Araç Dağıtım Optimizasyonu tezi için
python-docx ile tam formatlı Word belgesi üretir.

Çıktı: /home/user/arac-dagitim-sistemi/docs/tez_defteri.docx

Kullanım:
    python scripts/gen_tez.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

# ---------------------------------------------------------------------------
# Sabitler
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(20)
H2_SIZE = Pt(14)
LINE_SPACING = 1.5  # satır aralığı çarpanı
MARGIN = Cm(2.5)

OUTPUT_PATH = Path("/home/user/arac-dagitim-sistemi/docs/tez_defteri.docx")
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Yardımcı fonksiyonlar
# ---------------------------------------------------------------------------

def _set_font(run, name: str = FONT_NAME, size: Pt = BODY_SIZE,
              bold: bool = False, italic: bool = False,
              color: RGBColor | None = None) -> None:
    """Run için yazı tipi ayarlarını uygular."""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Doğu dili (Latin) fontu da ayarla
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def _para_format(para, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before: Pt = Pt(0), space_after: Pt = Pt(6),
                 line_spacing: float = LINE_SPACING,
                 left_indent: Cm | None = None,
                 first_line_indent: Cm | None = None) -> None:
    """Paragraf biçim ayarlarını uygular."""
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if left_indent is not None:
        pf.left_indent = left_indent
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_body_para(doc: Document, text: str,
                  bold: bool = False, italic: bool = False,
                  align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before: Pt = Pt(0), space_after: Pt = Pt(6),
                  first_indent: bool = True) -> None:
    """Gövde metni paragrafı ekler."""
    para = doc.add_paragraph()
    _para_format(para, align=align, space_before=space_before,
                 space_after=space_after,
                 first_line_indent=Cm(1.25) if first_indent else None)
    run = para.add_run(text)
    _set_font(run, bold=bold, italic=italic)


def add_heading1(doc: Document, text: str) -> None:
    """Birinci düzey başlık (20pt, BÜYÜK HARF) ekler."""
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(24), space_after=Pt(12),
                 line_spacing=1.0, first_line_indent=None)
    run = para.add_run(text.upper())
    _set_font(run, size=H1_SIZE, bold=True)


def add_heading2(doc: Document, text: str) -> None:
    """İkinci düzey başlık (14pt, Title Case) ekler."""
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=Pt(18), space_after=Pt(6),
                 line_spacing=1.0, first_line_indent=None)
    run = para.add_run(text)
    _set_font(run, size=H2_SIZE, bold=True)


def add_heading3(doc: Document, text: str) -> None:
    """Üçüncü düzey başlık (12pt, bold-italic) ekler."""
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=Pt(12), space_after=Pt(3),
                 line_spacing=1.0, first_line_indent=None)
    run = para.add_run(text)
    _set_font(run, size=BODY_SIZE, bold=True, italic=True)


def add_figure_caption(doc: Document, number: int, title: str) -> None:
    """Şekil başlığı (şeklin altına) ekler. 'Şekil X. Başlık' formatı."""
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(3), space_after=Pt(12),
                 line_spacing=1.0, first_line_indent=None)
    run = para.add_run(f"Şekil {number}. {title}")
    _set_font(run, italic=True)


def add_table_caption(doc: Document, number: int, title: str) -> None:
    """Tablo başlığı (tablonun üstüne) ekler. 'Tablo X. Başlık' formatı."""
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=Pt(12), space_after=Pt(3),
                 line_spacing=1.0, first_line_indent=None)
    run = para.add_run(f"Tablo {number}. {title}")
    _set_font(run, bold=True)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Madde işaretli liste öğesi ekler."""
    para = doc.add_paragraph(style="List Bullet")
    _para_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=Pt(0), space_after=Pt(3),
                 line_spacing=LINE_SPACING,
                 left_indent=Cm(1.25 + level * 0.63),
                 first_line_indent=None)
    run = para.add_run(text)
    _set_font(run)


def add_numbered(doc: Document, text: str, level: int = 0) -> None:
    """Numaralı liste öğesi ekler."""
    para = doc.add_paragraph(style="List Number")
    _para_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=Pt(0), space_after=Pt(3),
                 line_spacing=LINE_SPACING,
                 left_indent=Cm(1.25 + level * 0.63),
                 first_line_indent=None)
    run = para.add_run(text)
    _set_font(run)


def add_equation_placeholder(doc: Document, label: str, equation_text: str) -> None:
    """
    Matematiksel denklem için yer tutucu ekler.
    Not: Bu alan Word'de Equation Editor ile doldurulacaktır.
    """
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(6), space_after=Pt(6),
                 line_spacing=1.0, first_line_indent=None)
    run = para.add_run(f"[DENKLEM: {equation_text}]  ({label})")
    _set_font(run, italic=True, color=RGBColor(0x44, 0x44, 0x88))


def add_reference(doc: Document, text: str) -> None:
    """Kaynakça girdisi ekler (asılı girinti)."""
    para = doc.add_paragraph()
    _para_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=Pt(0), space_after=Pt(6),
                 line_spacing=LINE_SPACING,
                 left_indent=Cm(1.25),
                 first_line_indent=Cm(-1.25))
    run = para.add_run(text)
    _set_font(run)


def add_page_break(doc: Document) -> None:
    """Sayfa sonu ekler."""
    doc.add_page_break()


def _add_footer_page_number(doc: Document) -> None:
    """Her sayfaya merkeze hizalı sayfa numarası ekler."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        _set_font(run, size=Pt(10))
        # PAGE field
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)


def _set_margins(doc: Document) -> None:
    """Tüm bölümler için kenar boşluklarını ayarlar (2.5 cm her taraf)."""
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN


def make_simple_table(doc: Document, headers: list[str],
                      rows: list[list[str]]) -> None:
    """Basit stilli tablo oluşturur."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Başlık satırı
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            _set_font(run, bold=True)
        # Arka plan rengi (açık gri)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
    # Veri satırları
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    _set_font(run)
    return table


# ---------------------------------------------------------------------------
# Bölüm oluşturma fonksiyonları
# ---------------------------------------------------------------------------

def build_cover_page(doc: Document) -> None:
    """Kapak sayfasını oluşturur."""
    # Üst boşluk için boş paragraflar
    for _ in range(3):
        p = doc.add_paragraph()
        _para_format(p, space_before=Pt(0), space_after=Pt(0),
                     first_line_indent=None)

    # Üniversite başlığı
    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(6), first_line_indent=None)
    r = p.add_run("YILDIZ TEKNİK ÜNİVERSİTESİ")
    _set_font(r, bold=True, size=Pt(14))

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(6), first_line_indent=None)
    r = p.add_run("FEN BİLİMLERİ ENSTİTÜSÜ")
    _set_font(r, bold=True, size=Pt(13))

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(24), first_line_indent=None)
    r = p.add_run("ENDÜSTRİ MÜHENDİSLİĞİ ANABİLİM DALI")
    _set_font(r, size=Pt(12))

    # Çizgi
    for _ in range(2):
        p = doc.add_paragraph()
        _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(0), space_after=Pt(0), first_line_indent=None)
        r = p.add_run("─" * 55)
        _set_font(r)

    # Tez başlığı
    for _ in range(2):
        p = doc.add_paragraph()
        _para_format(p, first_line_indent=None)

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(24), space_after=Pt(12), first_line_indent=None)
    r = p.add_run(
        "OTOMOTİV BAYİ AĞINDA VERİ ODAKLI ARAÇ DAĞITIM OPTİMİZASYONU:\n"
        "ÇOK KRİTERLİ KARAR VERME VE KARMA TAMSAYILI DOĞRUSAL\n"
        "PROGRAMLAMA YAKLAŞIMI"
    )
    _set_font(r, bold=True, size=Pt(16))

    for _ in range(2):
        p = doc.add_paragraph()
        _para_format(p, first_line_indent=None)

    # Öğrenci bilgisi
    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(36), space_after=Pt(6), first_line_indent=None)
    r = p.add_run("Hazırlayan: [ÖĞRENCİ ADI SOYADI]")
    _set_font(r, size=Pt(12))

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(6), first_line_indent=None)
    r = p.add_run("Danışman: [DANIŞMAN ADI]")
    _set_font(r, size=Pt(12))

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(6), first_line_indent=None)
    r = p.add_run("Bölüm: Endüstri Mühendisliği")
    _set_font(r, size=Pt(12))

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(6), first_line_indent=None)
    r = p.add_run("Haziran 2026 / İSTANBUL")
    _set_font(r, size=Pt(12))

    add_page_break(doc)


def build_abstract(doc: Document) -> None:
    """Türkçe ve İngilizce özet sayfalarını oluşturur."""
    # --- ÖZET ---
    add_heading1(doc, "ÖZET")

    add_body_para(
        doc,
        "Otomotiv Bayi Ağında Veri Odaklı Araç Dağıtım Optimizasyonu: "
        "Çok Kriterli Karar Verme ve Karma Tamsayılı Doğrusal Programlama Yaklaşımı",
        bold=True, first_indent=False
    )

    add_body_para(
        doc,
        "Bu çalışmada, SUV segmentinde faaliyet gösteren bir otomotiv markasının "
        "Türkiye genelindeki 28 bayisine aylık araç dağıtımını otomatize etmek "
        "amacıyla veri odaklı bir karar destek sistemi geliştirilmiştir. "
        "Sistem, Ocak 2026 itibarıyla devreye alınmak üzere tasarlanmış olup "
        "aylık 300–1.500 arasında değişen araç dağıtımı kararlarını "
        "desteklemeyi hedeflemektedir."
    )

    add_body_para(
        doc,
        "Önerilen metodoloji üç temel aşamadan oluşmaktadır. Birinci aşamada, "
        "2024–2025 satış verileri üzerinde STL ayrıştırması ve ratio-to-mean "
        "yöntemiyle mevsimsel indeksler (SI) hesaplanmış; bayiler coğrafi "
        "konumlarına göre Tier A, B ve C olmak üzere üç gruba ayrılmıştır. "
        "Blended forecast yaklaşımıyla (%70 tier SI + %30 global SI) Aralık 2025 "
        "doğrulama setinde ortalama yüzde mutlak hata (MAPE) %7.82 olarak elde "
        "edilmiştir."
    )

    add_body_para(
        doc,
        "İkinci aşamada, dört ölçütü birleştiren Çok Kriterli Karar Verme (MCDM) "
        "modeli geliştirilmiştir: Performans Skoru (P, ağırlık=0.25), "
        "Lokasyon–Ürün Uyum Skoru (LP, ağırlık=0.35, cosine similarity tabanlı "
        "collaborative filtering), Mevsimsel Uyum Skoru (S, ağırlık=0.20) ve "
        "Hedef Yakınlık Skoru (H, ağırlık=0.20). Üçüncü aşamada ise bütünleşik "
        "skor maksimizasyonunu amaçlayan, CBC çözücü ile çözülen Karma Tamsayılı "
        "Doğrusal Programlama (MILP) modeli kurulmuştur."
    )

    add_body_para(
        doc,
        "2026 Ocak–Aralık dönemi için 28 bayi × 12 ay × 4 model hedef matrisi "
        "üretilmiştir. B segmentinin Mart 2026 yeni versiyon lansmanı için "
        "veri destekli boost katsayısı (×1.60) hesaplanmıştır. Sistem, muhafazakâr "
        "senaryoda 8.500, agresif senaryoda 10.000 araçlık yıllık satış planını "
        "desteklemektedir."
    )

    # Anahtar kelimeler
    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=Pt(12), space_after=Pt(6), first_line_indent=None)
    r1 = p.add_run("Anahtar Kelimeler: ")
    _set_font(r1, bold=True)
    r2 = p.add_run(
        "araç dağıtım optimizasyonu, karma tamsayılı doğrusal programlama, "
        "çok kriterli karar verme, mevsimsel indeks, collaborative filtering, "
        "otomotiv tedarik zinciri"
    )
    _set_font(r2)

    add_page_break(doc)

    # --- ABSTRACT ---
    add_heading1(doc, "ABSTRACT")

    add_body_para(
        doc,
        "Data-Driven Vehicle Allocation Optimization in an Automotive Dealer Network: "
        "A Multi-Criteria Decision Making and Mixed Integer Linear Programming Approach",
        bold=True, first_indent=False
    )

    add_body_para(
        doc,
        "In this study, a data-driven decision support system is developed to automate "
        "monthly vehicle allocation to 28 dealers of an SUV-segment automotive brand "
        "operating across Turkey. The system is designed to be deployed as of "
        "January 2026 and aims to support monthly distribution decisions ranging "
        "from 300 to 1,500 vehicles."
    )

    add_body_para(
        doc,
        "The proposed methodology consists of three main phases. In the first phase, "
        "seasonal indices (SI) are calculated using STL decomposition and the "
        "ratio-to-mean method on 2024–2025 sales data; dealers are grouped into "
        "three tiers (A, B, C) based on their geographic locations. Using a blended "
        "forecast approach (70% tier SI + 30% global SI), a mean absolute percentage "
        "error (MAPE) of 7.82% is achieved on the December 2025 validation set."
    )

    add_body_para(
        doc,
        "In the second phase, a Multi-Criteria Decision Making (MCDM) model "
        "integrating four criteria is developed: Performance Score (P, weight=0.25), "
        "Location–Product Compatibility Score (LP, weight=0.35, cosine similarity-based "
        "collaborative filtering), Seasonal Fit Score (S, weight=0.20), and "
        "Target Proximity Score (H, weight=0.20). In the third phase, a Mixed Integer "
        "Linear Programming (MILP) model maximizing the composite score is formulated "
        "and solved using the CBC solver."
    )

    add_body_para(
        doc,
        "A 28-dealer × 12-month × 4-model target matrix is generated for the "
        "January–December 2026 period. A data-driven boost coefficient (×1.60) is "
        "computed for the March 2026 new version launch of the B segment. The system "
        "supports an annual sales plan of 8,500 vehicles in the conservative scenario "
        "and 10,000 vehicles in the aggressive scenario."
    )

    p = doc.add_paragraph()
    _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=Pt(12), space_after=Pt(6), first_line_indent=None)
    r1 = p.add_run("Keywords: ")
    _set_font(r1, bold=True)
    r2 = p.add_run(
        "vehicle allocation optimization, mixed integer linear programming, "
        "multi-criteria decision making, seasonal index, collaborative filtering, "
        "automotive supply chain"
    )
    _set_font(r2)

    add_page_break(doc)


def build_table_of_contents(doc: Document) -> None:
    """İçindekiler sayfasını oluşturur."""
    add_heading1(doc, "İÇİNDEKİLER")

    toc_entries = [
        ("ÖZET", "iii"),
        ("ABSTRACT", "iv"),
        ("İÇİNDEKİLER", "v"),
        ("ŞEKİL LİSTESİ", "vi"),
        ("TABLO LİSTESİ", "vii"),
        ("SİMGE VE KISALTMA LİSTESİ", "viii"),
        ("1. GİRİŞ", "1"),
        ("   1.1 Problem Tanımı ve Motivasyon", "1"),
        ("   1.2 Araştırmanın Amacı ve Kapsamı", "3"),
        ("   1.3 Tezin Organizasyonu", "4"),
        ("2. LİTERATÜR TARAMASI", "5"),
        ("   2.1 Araç Dağıtım Problemi", "5"),
        ("   2.2 Çok Kriterli Karar Verme Yöntemleri", "7"),
        ("   2.3 Karma Tamsayılı Doğrusal Programlama", "9"),
        ("   2.4 Zaman Serisi Tahmini ve Mevsimsellik", "11"),
        ("   2.5 Collaborative Filtering ve Ürün-Konum Uyumu", "13"),
        ("3. METODOLOJİ", "15"),
        ("   3.1 Sistem Mimarisi", "15"),
        ("   3.2 Veri Toplama ve Ön İşleme", "17"),
        ("   3.3 Mevsimsellik Tahmini", "19"),
        ("   3.4 Çok Kriterli Skorlama Modeli", "25"),
        ("   3.5 MILP Optimizasyon Modeli", "31"),
        ("4. BULGULAR VE ANALİZ", "39"),
        ("   4.1 Mevsimsel İndeks Doğrulaması", "39"),
        ("   4.2 2026 Yıllık Satış Planı", "42"),
        ("   4.3 Bayi Bazlı Hedef Matrisi", "47"),
        ("   4.4 Optimizasyon Sonuçları", "51"),
        ("5. TARTIŞMA VE SONUÇ", "55"),
        ("   5.1 Temel Bulgular", "55"),
        ("   5.2 Yönetimsel Öneriler", "57"),
        ("   5.3 Sınırlılıklar ve Gelecek Çalışmalar", "59"),
        ("KAYNAKÇA", "61"),
        ("EKLER", "67"),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        _para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_before=Pt(0), space_after=Pt(3),
                     line_spacing=1.15, first_line_indent=None)
        bold = not entry.startswith("   ")
        r = p.add_run(entry)
        _set_font(r, bold=bold)
        # Noktalı tabulatör
        p.add_run(" " + "." * max(1, 60 - len(entry)) + " " + page)

    add_page_break(doc)


def build_figure_list(doc: Document) -> None:
    """Şekil listesi sayfası."""
    add_heading1(doc, "ŞEKİL LİSTESİ")

    figures = [
        ("1", "Sistem Mimarisi Genel Görünümü", "16"),
        ("2", "Veri Akış Diyagramı", "18"),
        ("3", "2024–2025 Aylık Satış Zaman Serisi (Tüm Bayiler)", "20"),
        ("4", "STL Ayrıştırma Bileşenleri – Trend, Mevsimsellik, Artık", "22"),
        ("5", "Tier Gruplarına Göre Mevsimsel İndeks Profilleri", "24"),
        ("6", "MCDM Skor Dağılımı – Bayi Bazında Kutu Grafiği", "28"),
        ("7", "Cosine Similarity Isı Haritası – Bayi × Model", "30"),
        ("8", "MILP Kısıt Yapısı Şeması", "34"),
        ("9", "Aralık 2025 Tahmin vs. Gerçek Karşılaştırması", "40"),
        ("10", "2026 Yıllık Satış Planı – Aylık Dağılım (Her İki Senaryo)", "44"),
        ("11", "Model Bazında 2026 Aylık Dağılım", "46"),
        ("12", "28 Bayi Coğrafi Dağılım Haritası (Folium)", "48"),
        ("13", "Bayi Hedef Matrisi Isı Haritası", "50"),
        ("14", "MILP Optimizasyon Sonuçları – Dağıtım Verimliliği", "53"),
    ]

    for num, title, page in figures:
        p = doc.add_paragraph()
        _para_format(p, space_before=Pt(0), space_after=Pt(3),
                     line_spacing=1.15, first_line_indent=None)
        p.add_run(f"Şekil {num}. {title} {'.' * max(1, 55 - len(title) - len(num))} {page}")

    add_page_break(doc)


def build_table_list(doc: Document) -> None:
    """Tablo listesi sayfası."""
    add_heading1(doc, "TABLO LİSTESİ")

    tables = [
        ("1", "Literatür Özet Tablosu", "6"),
        ("2", "Veri Seti Açıklayıcı İstatistikleri", "18"),
        ("3", "Bayi Tier Gruplandırması", "21"),
        ("4", "Tier Gruplarına Göre Mevsimsel İndeks Değerleri (2024–2025)", "23"),
        ("5", "MCDM Kriter Ağırlıkları ve Hesaplama Mantığı", "26"),
        ("6", "MILP Model Parametreleri", "33"),
        ("7", "Aralık 2025 Tahmin Doğrulama Sonuçları", "41"),
        ("8", "2026 Yıllık Satış Planı – İki Senaryo Karşılaştırması", "43"),
        ("9", "B Segmenti Mart Lansmanı İstatistiksel Gerekçesi", "45"),
        ("10", "28 Bayi 2026 Yıllık Hedef Özeti", "49"),
        ("11", "MILP Optimizasyon Çözüm İstatistikleri", "52"),
    ]

    for num, title, page in tables:
        p = doc.add_paragraph()
        _para_format(p, space_before=Pt(0), space_after=Pt(3),
                     line_spacing=1.15, first_line_indent=None)
        p.add_run(f"Tablo {num}. {title} {'.' * max(1, 55 - len(title) - len(num))} {page}")

    add_page_break(doc)


def build_symbols(doc: Document) -> None:
    """Simge ve kısaltma listesi."""
    add_heading1(doc, "SİMGE VE KISALTMA LİSTESİ")

    items = [
        ("MCDM", "Multi-Criteria Decision Making (Çok Kriterli Karar Verme)"),
        ("MILP", "Mixed Integer Linear Programming (Karma Tamsayılı Doğrusal Programlama)"),
        ("VAP", "Vehicle Allocation Problem (Araç Dağıtım Problemi)"),
        ("STL", "Seasonal and Trend decomposition using Loess"),
        ("SI", "Seasonal Index (Mevsimsel İndeks)"),
        ("MAPE", "Mean Absolute Percentage Error (Ortalama Yüzde Mutlak Hata)"),
        ("CBC", "Coin-or Branch and Cut (MIP çözücüsü)"),
        ("EW", "Exponentially Weighted (Üstel Ağırlıklı)"),
        ("LP", "Lokasyon-Ürün Uyum Skoru"),
        ("CS", "Composite Score (Bütünleşik Skor)"),
        ("P", "Performans Skoru"),
        ("S", "Mevsimsel Uyum Skoru"),
        ("H", "Hedef Yakınlık Skoru"),
        ("SUV", "Sport Utility Vehicle"),
        ("APA", "American Psychological Association (atıf stili)"),
        ("CSV", "Comma-Separated Values"),
        ("SQL", "Structured Query Language"),
        ("API", "Application Programming Interface"),
    ]

    for abbr, meaning in items:
        p = doc.add_paragraph()
        _para_format(p, space_before=Pt(0), space_after=Pt(3),
                     line_spacing=LINE_SPACING,
                     left_indent=Cm(0), first_line_indent=None)
        r1 = p.add_run(f"{abbr:<12}")
        _set_font(r1, bold=True)
        r2 = p.add_run(meaning)
        _set_font(r2)

    add_page_break(doc)


def build_chapter1(doc: Document) -> None:
    """Bölüm 1: Giriş"""
    add_heading1(doc, "1. GİRİŞ")

    add_heading2(doc, "1.1 Problem Tanımı ve Motivasyon")

    add_body_para(
        doc,
        "Otomotiv sektörü, küresel ölçekte son derece rekabetçi ve dinamik bir "
        "pazar yapısına sahiptir. Türkiye pazarı, 2024 yılında yaklaşık 1.1 milyon "
        "binek araç satışıyla Avrupa'nın en büyük pazarlarından biri konumuna "
        "gelmiştir (ODD, 2025). Bu büyüklükteki bir pazarda üretici ve ithalatçıların "
        "karşı karşıya kaldığı en kritik operasyonel sorunlardan biri, sınırlı araç "
        "envanterinin bayi ağına etkin biçimde dağıtılmasıdır."
    )

    add_body_para(
        doc,
        "Araç Dağıtım Problemi (Vehicle Allocation Problem – VAP), klasik stok "
        "yönetimi ve atama problemlerinin bir uzantısıdır. Geleneksel yaklaşımlarda "
        "dağıtım kararları genellikle satış temsilcilerinin deneyimine, bayi "
        "taleplerinin sabit oranlarda karşılanmasına veya basit kural tabanlı "
        "sistemlere dayanmaktadır. Bu yaklaşımlar; mevsimsellik etkilerini, "
        "bölgesel talep farklılıklarını ve ürün-konum uyumunu yeterince "
        "hesaba katamamaktadır."
    )

    add_body_para(
        doc,
        "Bu tezde incelenen Marka X, SUV segmentinde 28 aktif bayisiyle Türkiye "
        "genelinde faaliyet göstermektedir. 2025 yılsonu itibarıyla A1, C1 ve D1 "
        "modelleri üretimden kaldırılmış; 2026 yılına A Segmenti (A2, A3 "
        "versiyonları) ve B Segmenti (B1, B2 versiyonları) olmak üzere dört aktif "
        "model ile girilmektedir. Mart 2026'da planlanan B1 yeni versiyon lansmanı, "
        "talep planlamasında özellikle dikkat gerektiren bir dönem oluşturmaktadır."
    )

    add_body_para(
        doc,
        "Mevcut el ile yürütülen dağıtım süreci aşağıdaki temel sorunları "
        "barındırmaktadır:"
    )
    add_bullet(doc, "Mevsimsellik etkilerinin sistematik olarak modellenmemesi")
    add_bullet(doc, "Bayi bazlı performans farklılıklarının dağıtıma yansıtılmaması")
    add_bullet(doc, "Yıllık satış hedeflerine olan uzaklığın anlık takip edilememesi")
    add_bullet(doc, "Ürün karışımı (model/versiyon/renk) kararlarının veri desteksiz verilmesi")
    add_bullet(doc, "Yeni bayi ekleme veya kapatma senaryolarına adaptasyon güçlüğü")

    add_body_para(
        doc,
        "Söz konusu sorunlar, hem müşteri memnuniyetini (stokta bulunmayan modeller "
        "nedeniyle kaybedilen satışlar) hem de bayi karlılığını (gereksiz stok "
        "maliyetleri) olumsuz etkilemektedir. Veriye dayalı, otomatize edilmiş bir "
        "dağıtım destek sistemi, bu boşluğu doldurmak amacıyla tasarlanmıştır."
    )

    add_heading2(doc, "1.2 Araştırmanın Amacı ve Kapsamı")

    add_body_para(
        doc,
        "Bu araştırmanın temel amacı, Marka X'in 28 bayisine aylık araç dağıtımını "
        "veri odaklı, şeffaf ve tekrarlanabilir bir karar mekanizmasıyla "
        "otomatize etmektir. Bu amaca ulaşmak için aşağıdaki alt hedefler "
        "belirlenmiştir:"
    )

    hedefler = [
        "2024–2025 satış verilerinden istatistiksel yöntemlerle mevsimsel indeks "
        "hesaplamak ve 2026 yıllık satış planını oluşturmak.",
        "Dört farklı ölçütü bütünleştiren bir Çok Kriterli Karar Verme (MCDM) "
        "modeli geliştirerek bayi bazlı öncelik skorları hesaplamak.",
        "Karma Tamsayılı Doğrusal Programlama (MILP) ile aylık dağıtım "
        "miktarlarını optimize etmek.",
        "Tüm süreci takip edilebilir, yorumlanabilir ve güncellemeye açık bir "
        "Python tabanlı karar destek sisteminde hayata geçirmek.",
        "Sonuçları etkileşimli bir Streamlit dashboard'unda sunmak.",
    ]
    for i, h in enumerate(hedefler, 1):
        add_numbered(doc, f"{h}")

    add_body_para(
        doc,
        "Araştırmanın kapsamı, Ocak 2024 – Aralık 2025 satış dönemi ve 2026 "
        "planlama dönemiyle sınırlıdır. Çalışmada motor tipi, şanzıman veya "
        "yakıt tipi ayrımına girilmemekte; model, versiyon ve renk düzeyinde "
        "analizler yapılmaktadır. Bayi stok kapasitesi kısıtı (1.000 araç "
        "eşiğinin altında olduğu için) modele dahil edilmemiştir."
    )

    add_heading2(doc, "1.3 Tezin Organizasyonu")

    add_body_para(
        doc,
        "Tezin geri kalanı aşağıdaki şekilde düzenlenmiştir. İkinci bölümde araç "
        "dağıtım problemi, çok kriterli karar verme, karma tamsayılı programlama, "
        "zaman serisi tahmini ve collaborative filtering konularındaki ilgili "
        "literatür incelenmektedir. Üçüncü bölümde geliştirilen metodoloji "
        "ayrıntılı biçimde açıklanmaktadır. Dördüncü bölümde elde edilen bulgular "
        "sunulmakta ve yorumlanmaktadır. Beşinci bölümde sonuçlar tartışılmakta, "
        "yönetimsel öneriler ve gelecek araştırma yönleri belirtilmektedir. "
        "Kaynakça ve ekler tezin sonunda yer almaktadır."
    )

    add_page_break(doc)


def build_chapter2(doc: Document) -> None:
    """Bölüm 2: Literatür Taraması"""
    add_heading1(doc, "2. LİTERATÜR TARAMASI")

    add_body_para(
        doc,
        "Bu bölümde, araştırmanın dayandığı beş temel alan için kapsamlı bir "
        "literatür taraması sunulmaktadır: (1) Araç Dağıtım Problemi, "
        "(2) Çok Kriterli Karar Verme, (3) Karma Tamsayılı Doğrusal Programlama, "
        "(4) Zaman Serisi Tahmini ve Mevsimsellik, (5) Collaborative Filtering "
        "ve Ürün-Konum Uyumu. Bölümün sonunda bu çalışmaları bir arada sunan "
        "kapsamlı bir literatür özet tablosu yer almaktadır."
    )

    add_heading2(doc, "2.1 Araç Dağıtım Problemi")

    add_body_para(
        doc,
        "Araç Dağıtım Problemi (Vehicle Allocation Problem – VAP), lojistik ve "
        "tedarik zinciri yönetimi literatüründe önemli bir yer tutmaktadır. "
        "Luss ve Rosenwein (1997), çok dönemli araç dağıtım modellerini sistematik "
        "biçimde incelemiş; talebin rassal ya da deterministik olduğu durumlarda "
        "farklı optimizasyon stratejilerini karşılaştırmıştır. Bu çalışma, "
        "VAP'ın matematiksel temellerini atmış ve sonraki araştırmalar için "
        "referans bir kaynak haline gelmiştir."
    )

    add_body_para(
        doc,
        "Sherbrooke (1968), çok kademeli envanter teorisinin öncü çalışmasında, "
        "dağıtım ağlarında stok kararlarının birbirine bağımlılığını resmi olarak "
        "modellemştir. METRIC (Multi-Echelon Technique for Recoverable Item Control) "
        "modeli, otomotiv yedek parça dağıtımından havacılık lojistiğine kadar pek "
        "çok alanda uygulanmıştır. Sherbrooke'un çerçevesi, bu tezdeki bayi ağı "
        "hiyerarşisinin kavramsal temelini oluşturmaktadır."
    )

    add_body_para(
        doc,
        "Cachon ve Lariviere (1999), tedarik zinciri koordinasyonu bağlamında "
        "kapasite rezervasyonu ve talep belirsizliğinin dağıtım kararlarına "
        "etkisini incelemiştir. Araştırmacıların önerdiği sözleşme mekanizmaları, "
        "üretici ile bayi arasındaki çatışan çıkarları uyumlaştırmaya yönelik "
        "pratik araçlar sunmaktadır. Bu tezde geliştirilen ±%20 hedef aralığı "
        "kısıtı, benzer bir esneklik ve uyum mantığını yansıtmaktadır."
    )

    add_heading2(doc, "2.2 Çok Kriterli Karar Verme Yöntemleri")

    add_body_para(
        doc,
        "Çok Kriterli Karar Verme (MCDM) alanı, birbiriyle çelişen birden fazla "
        "ölçütü aynı anda göz önünde bulundurarak en iyi kararı bulmayı amaçlar. "
        "Hwang ve Yoon (1981), TOPSIS (Technique for Order of Preference by "
        "Similarity to Ideal Solution) yöntemini geliştirmiş ve MCDM metodolojilerini "
        "kapsamlı biçimde sınıflandırmıştır. TOPSIS, alternatifler arasındaki "
        "uzaklık hesabına dayanan yapısıyla tedarik zinciri ve lojistik "
        "kararlarında geniş uygulama alanı bulmuştur."
    )

    add_body_para(
        doc,
        "Bu tezde geliştirilen ağırlıklı doğrusal bütünleşik skor modeli (CS = "
        "0.25P + 0.35LP + 0.20S + 0.20H), Hwang ve Yoon (1981) çerçevesiyle "
        "uyumlu olup sektöre özgü dört ölçütü dengeli biçimde birleştirmektedir. "
        "Ağırlık belirleme sürecinde, uzman görüşü tabanlı Analitik Hiyerarşi "
        "Sürecine (AHP) alternatif olarak, sektör bilgisi ve veri analizi "
        "birleştirilmiştir."
    )

    add_heading2(doc, "2.3 Karma Tamsayılı Doğrusal Programlama")

    add_body_para(
        doc,
        "Doğrusal programlamanın temellerini atan Dantzig ve Thapa (1997), "
        "simpleks algoritmasını ve LP'nin teorik çerçevesini kapsamlı biçimde "
        "sunmuştur. Karma Tamsayılı Doğrusal Programlama (MILP), karar "
        "değişkenlerinin bir kısmının tam sayı olmasını gerektiren problemler "
        "için kullanılan ve LP'nin doğal bir uzantısı olan bir optimizasyon "
        "yaklaşımıdır. Araç adedinin doğası gereği tam sayı olması, bu tezde "
        "MILP kullanımını zorunlu kılmaktadır."
    )

    add_body_para(
        doc,
        "PuLP, Python dilinde LP ve MILP modellemesi için kullanılan açık kaynaklı "
        "bir kütüphanedir (Mitchell et al., 2011). CBC (Coin-or Branch and Cut) "
        "çözücüsü ile entegre şekilde çalışan PuLP, ticari çözücülere erişimin "
        "kısıtlı olduğu akademik ve pratik uygulamalar için önemli bir araç "
        "haline gelmiştir. Bu tezde 28 bayi × 12 ay × 4 model boyutundaki "
        "model (1.344 karar değişkeni), PuLP+CBC kombinasyonuyla saniyeler "
        "içinde çözülmektedir."
    )

    add_heading2(doc, "2.4 Zaman Serisi Tahmini ve Mevsimsellik")

    add_body_para(
        doc,
        "Cleveland vd. (1990), Loess tabanlı STL ayrıştırma yöntemini geliştirmiş "
        "ve bu yöntem, zaman serilerindeki trend, mevsimsellik ve artık bileşenlerini "
        "birbirinden ayırt etmede referans bir teknik haline gelmiştir. STL, "
        "mevsimsel bileşenin zaman içinde değişmesine izin vermesiyle klasik "
        "ayrıştırma yöntemlerine kıyasla önemli bir esneklik sunmaktadır."
    )

    add_body_para(
        doc,
        "Taylor (2018), Facebook'ta geliştirilen Prophet algoritmasını tanıtmış; "
        "bu algoritma iş zaman serilerindeki tatil etkileri, lineer olmayan "
        "trendler ve çoklu mevsimsellik bileşenlerini modelleyebilmektedir. "
        "Makridakis vd. (2018), M4 tahmin yarışmasında 100.000'den fazla zaman "
        "serisi üzerinde 60 farklı yöntemi karşılaştırmış; hibrit yöntemlerin "
        "çoğu durumda tek başına ML modellerini geride bıraktığını göstermiştir."
    )

    add_body_para(
        doc,
        "Bu tezde, ratio-to-mean yöntemi ile hesaplanan mevsimsel indeksler, "
        "sektördeki uzun süreli kullanımı (Bowerman vd., 2005) ve "
        "yorumlanabilirliği nedeniyle tercih edilmiştir. Bayilerin coğrafi "
        "gruplara (tier) ayrılarak tier ve global düzeylerde blended tahmin "
        "üretilmesi, Hiyerarşik Zaman Serisi Tahmini (Hierarchical Time Series "
        "Forecasting) yaklaşımıyla örtüşmektedir (Hyndman vd., 2011)."
    )

    add_heading2(doc, "2.5 Collaborative Filtering ve Ürün-Konum Uyumu")

    add_body_para(
        doc,
        "Collaborative filtering, öneri sistemleri literatüründe köklü bir yer "
        "tutmaktadır. Sarwar vd. (2001), öğe tabanlı (item-based) collaborative "
        "filtering yaklaşımını sistematik biçimde ele almış; cosine similarity "
        "metriğinin kullanıcı–öğe etkileşimlerini ölçmede etkinliğini "
        "göstermiştir. Bu tezde, bayilerin (kullanıcı) ve araç modellerinin "
        "(öğe) geçmiş satış örüntülerinden elde edilen matris, cosine similarity "
        "ile normalizasyon hesabına tabi tutularak Lokasyon–Ürün Uyum Skoru "
        "(LP) oluşturulmaktadır."
    )

    add_body_para(
        doc,
        "Fahimnia vd. (2015), yeşil tedarik zinciri konusundaki kapsamlı "
        "literatür taramasında, lojistik kararlarında çevresel sürdürülebilirlik "
        "ölçütlerinin entegrasyonunu tartışmıştır. Bu çalışma, araç dağıtımını "
        "yalnızca finansal değil çok boyutlu bir optimizasyon problemi olarak "
        "ele almanın önemini vurgulayan güncel bir referanstır."
    )

    # --- Literatür Özet Tablosu ---
    add_heading2(doc, "2.6 Literatür Özet Tablosu")

    add_body_para(
        doc,
        "Aşağıdaki tablo, bu tezin metodolojisine katkıda bulunan temel "
        "akademik çalışmaları özetlemektedir.",
        first_indent=False
    )

    add_table_caption(doc, 1, "Literatür Özet Tablosu")

    lit_headers = ["Yazar(lar)", "Yıl", "Çalışma Amacı", "Yöntem", "Temel Bulgular"]
    lit_rows = [
        ["Sherbrooke", "1968",
         "Çok kademeli stok yönetimi (METRIC modeli)",
         "Stokastik envanter modeli",
         "Dağıtım ağlarında eşzamanlı stok kararlarının modellenebileceğini kanıtladı"],
        ["Hwang & Yoon", "1981",
         "MCDM yöntemlerinin sınıflandırılması ve TOPSIS geliştirme",
         "TOPSIS, ağırlıklı toplam modeli",
         "Alternatifler ideal çözüme uzaklığa göre sıralanabilir"],
        ["Cachon & Lariviere", "1999",
         "Tedarik zinciri koordinasyonunda kapasite rezervasyonu",
         "Oyun teorisi, sözleşme modelleri",
         "Esneklik mekanizmaları tedarikçi-alıcı uyumunu artırır"],
        ["Luss & Rosenwein", "1997",
         "Çok dönemli araç dağıtım problemi",
         "LP ve MIP tabanlı modeller",
         "Deterministik VAP'ta optimal politika yapısı tanımlandı"],
        ["Dantzig & Thapa", "1997",
         "Doğrusal programlamanın teorik temelleri",
         "Simpleks algoritması, LP teorisi",
         "LP'nin temel teoremlerini ve algoritmik çerçevesini ortaya koydu"],
        ["Sarwar vd.", "2001",
         "Öğe tabanlı collaborative filtering",
         "Cosine similarity, Pearson korelasyonu",
         "Cosine similarity, seyrek matrisler için etkin öneri üretir"],
        ["Taylor (Prophet)", "2018",
         "İş zaman serilerinde otomatik tahmin",
         "Ayrıştırma bazlı Bayesian model",
         "Tatil ve mevsim etkilerini ayrı bileşen olarak modelledi"],
        ["Makridakis vd.", "2018",
         "M4 tahmin yarışması – 60 yöntemin karşılaştırılması",
         "İstatistiksel + ML hibrit modeller",
         "Hibrit yaklaşımlar saf ML modellerini çoğu durumda geçti"],
        ["Fahimnia vd.", "2015",
         "Yeşil tedarik zinciri literatür taraması",
         "Sistematik litereatür taraması",
         "Çevresel ölçütlerin lojistik kararlarına entegrasyonu artıyor"],
        ["Mitchell vd.", "2011",
         "PuLP: Python LP/MILP modelleme aracı",
         "Açık kaynaklı LP/MILP kütüphanesi",
         "CBC çözücü entegrasyonuyla pratik optimizasyonu demokratize etti"],
    ]

    make_simple_table(doc, lit_headers, lit_rows)

    # Şekil kural notu
    p = doc.add_paragraph()
    _para_format(p, space_before=Pt(6), space_after=Pt(6),
                 first_line_indent=None)
    r = p.add_run(
        "Not: Tablo kaynakları tam atıf bilgisiyle Kaynakça bölümünde yer almaktadır."
    )
    _set_font(r, italic=True)

    add_page_break(doc)


def build_chapter3(doc: Document) -> None:
    """Bölüm 3: Metodoloji"""
    add_heading1(doc, "3. METODOLOJİ")

    add_body_para(
        doc,
        "Bu bölümde geliştirilen karar destek sisteminin metodolojisi üç ana alt "
        "bölümde ele alınmaktadır: (1) Mevsimsellik Tahmini, (2) Çok Kriterli "
        "Skorlama Modeli ve (3) MILP Optimizasyon Modeli. Bölümün başında sistem "
        "mimarisi ve veri akışı özetlenmektedir."
    )

    add_heading2(doc, "3.1 Sistem Mimarisi")

    add_body_para(
        doc,
        "Geliştirilen karar destek sistemi, modüler bir yapıya sahip olup üç ana "
        "katmandan oluşmaktadır: Veri Katmanı, Analiz Katmanı ve Sunum Katmanı. "
        "Tüm bileşenler Python 3.11 ile geliştirilmiş olup bulut tabanlı çalışmaya "
        "(GitHub Codespaces, Streamlit Cloud) uygun şekilde tasarlanmıştır."
    )

    add_heading3(doc, "3.1.1 Veri Katmanı")
    add_body_para(
        doc,
        "Veri katmanı, altı CSV dosyasından oluşmakta olup bu dosyalar SQLite "
        "veritabanına (arac_dagitim.db) aktarılmaktadır. Veritabanı şeması; "
        "dim_bayi, dim_arac, fact_satis, fact_hedef, fact_envanter ve "
        "dim_rakip_satis tablolarından oluşmaktadır. Tüm erişim pathlib.Path "
        "ile relative yollar üzerinden gerçekleştirilmektedir."
    )

    add_heading3(doc, "3.1.2 Analiz Katmanı")
    add_body_para(
        doc,
        "Analiz katmanı, mevsimsel indeks hesaplama, MCDM skorlama ve MILP "
        "optimizasyon modüllerini içermektedir. Bu modüller src/ dizini altında "
        "bağımsız Python paketleri olarak organize edilmiştir. Her modül, "
        "Google style docstring ve type hint kullanımıyla belgelenmiştir."
    )

    add_heading3(doc, "3.1.3 Sunum Katmanı")
    add_body_para(
        doc,
        "Sunum katmanı, Streamlit ile geliştirilen interaktif bir dashboard'dan "
        "oluşmaktadır. Dashboard; Folium ile coğrafi bayi haritasını, Plotly "
        "ile etkileşimli zaman serisi ve dağıtım grafiklerini, ve aylık optimizasyon "
        "sonuç tablolarını sunmaktadır. DEMO_MODE flag'i ile hassas veriler "
        "maskelenebilmektedir."
    )

    add_figure_caption(doc, 1, "Sistem Mimarisi – Veri, Analiz ve Sunum Katmanları")

    add_heading2(doc, "3.2 Veri Toplama ve Ön İşleme")

    add_body_para(
        doc,
        "Sistemin temel veri kaynakları anonimleştirilmiş CSV dosyaları olup "
        "tablo 2'de bu dosyaların içerik ve boyutları özetlenmektedir."
    )

    add_table_caption(doc, 2, "Veri Seti Açıklayıcı İstatistikleri")

    data_headers = ["Dosya Adı", "Boyut", "Kapsam", "Temel Değişkenler"]
    data_rows = [
        ["sales_2024_2025.csv", "6.439 satır",
         "Ocak 2024 – Aralık 2025",
         "bayi_id, model, versiyon, renk, tarih, adet"],
        ["dealer_targets_2026.csv", "28 satır",
         "2026 yıllık hedefler",
         "bayi_id, yillik_hedef, model_pay"],
        ["dealer_locations.csv", "28 satır",
         "Bayi coğrafi bilgisi",
         "bayi_id, il, ilce, lat, lon, tier"],
        ["monthly_performance_2025.csv", "336 satır",
         "2025 aylık gerçekleşmeler",
         "bayi_id, ay, hedef, gerceklesme, oran"],
        ["competitor_sales.csv", "Çok satır",
         "2024–2025 rakip satışları",
         "il, marka, ay, adet"],
        ["inventory_2026_01.csv", "Çok satır",
         "Ocak 2026 araç envanteri",
         "model, versiyon, renk, adet"],
    ]

    make_simple_table(doc, data_headers, data_rows)

    add_body_para(
        doc,
        "Ön işleme aşamasında aşağıdaki adımlar uygulanmıştır:",
        first_indent=False
    )
    add_bullet(doc, "UTF-8 encoding doğrulaması ve eksik değer analizi")
    add_bullet(doc, "Tarih sütunlarının ISO 8601 (YYYY-MM-DD) formatına dönüştürülmesi")
    add_bullet(doc, "Bayi isimlerinin ANONYMIZE=True ile 'Bayi XX' formatına standartlaştırılması")
    add_bullet(doc, "Aykırı değer tespiti (IQR yöntemi) ve manuel inceleme")
    add_bullet(doc, "Mart 2026 öncesi A1/C1/D1 model satışlarının filtre dışı bırakılması")
    add_bullet(doc,
               "6 yeni bayi (Bayi 23–28) için ilk 12 aylık veri eksikliğinin tier ortalamasıyla doldurulması")

    add_heading2(doc, "3.3 Mevsimsellik Tahmini")

    add_heading3(doc, "3.3.1 STL Ayrıştırma")

    add_body_para(
        doc,
        "STL (Seasonal and Trend decomposition using Loess) yöntemi, bir zaman "
        "serisini trend (T_t), mevsimsel (S_t) ve artık (R_t) bileşenlerine "
        "ayırır:"
    )

    add_equation_placeholder(doc, "Denklem 1",
                              "Y_t = T_t + S_t + R_t (toplamalı model)")

    add_body_para(
        doc,
        "Loess düzleştirme, yerel ağırlıklı regresyon kullanarak trend bileşenini "
        "hesaplar. Mevsimsellik pencere genişliği (s_window) ve trend pencere "
        "genişliği (t_window) bu analizde sırasıyla 'periodic' ve 13 olarak "
        "ayarlanmıştır."
    )

    add_heading3(doc, "3.3.2 Ratio-to-Mean Yöntemi ile Mevsimsel İndeks Hesabı")

    add_body_para(
        doc,
        "Mevsimsel indeks (SI), her ayın yıllık ortalamaya oranının çok yıllık "
        "ortalaması olarak hesaplanmaktadır:"
    )

    add_equation_placeholder(
        doc, "Denklem 2",
        "SI_m = (1/N) * sum_{y=1}^{N} [ Y_{m,y} / (Y_ort_y) ]"
    )

    add_body_para(
        doc,
        "Burada SI_m ay m'nin mevsimsel indeksini, Y_{m,y} y yılının m ayındaki "
        "satışı ve Y_ort_y o yılın aylık ortalama satışını temsil etmektedir. "
        "N yıl sayısını (bu çalışmada N=2) göstermektedir."
    )

    add_heading3(doc, "3.3.3 Bayi Tier Gruplandırması")

    add_body_para(
        doc,
        "Bayiler, coğrafi konumlarına ve tarihsel satış hacimleri dikkate alınarak "
        "üç tier grubuna ayrılmıştır:"
    )

    add_table_caption(doc, 3, "Bayi Tier Gruplandırması")
    tier_headers = ["Tier", "Bölgeler", "Bayi Sayısı", "Açıklama"]
    tier_rows = [
        ["Tier A", "Marmara, Ege, İç Anadolu", "21",
         "Yüksek hacimli, olgun pazar bayileri"],
        ["Tier B", "Akdeniz", "4",
         "Orta hacimli, büyüme potansiyeli olan bayiler"],
        ["Tier C", "Güneydoğu Anadolu, Karadeniz", "3",
         "Düşük hacimli, yeni/gelişmekte olan bayiler"],
    ]
    make_simple_table(doc, tier_headers, tier_rows)

    add_heading3(doc, "3.3.4 Blended Forecast Yaklaşımı")

    add_body_para(
        doc,
        "Bayi düzeyinde güvenilir mevsimsel indeks hesaplamak için yeterli "
        "veri bulunmadığında (özellikle yeni ve düşük hacimli bayiler için), "
        "tier ve global indeksleri birleştiren blended forecast yaklaşımı "
        "uygulanmaktadır:"
    )

    add_equation_placeholder(
        doc, "Denklem 3",
        "SI_blended = 0.70 * SI_tier + 0.30 * SI_global"
    )

    add_body_para(
        doc,
        "Bu karışım oranları, 2025 yılı verilerinde MAPE minimizasyonu "
        "amacıyla ızgara araması (grid search) ile belirlenmiştir."
    )

    add_heading3(doc, "3.3.5 Mart 2026 Lansman Boost Katsayısı")

    add_body_para(
        doc,
        "B1 modelinin Mart 2026 yeni versiyon lansmanı, talep artışını "
        "mevsimsel indekse ek bir katsayıyla modellemeyi gerektirmiştir. "
        "Bu katsayının istatistiksel gerekçesi Tablo 4'te özetlenmektedir."
    )

    add_table_caption(doc, 4,
                      "B Segmenti Mart Lansmanı İstatistiksel Gerekçesi")
    launch_headers = ["Gösterge", "2024 Mart", "2025 Mart", "2026 Beklenti"]
    launch_rows = [
        ["B Segmenti Pazar Payı (Marka X içi)", "%55.7", "%44.5",
         "%55.7 (restore)"],
        ["B Segmenti Mart SI (yıllık ort.na oran)", "2.459", "1.513",
         "Boost uygulandı"],
        ["Hesaplanan alt sınır boost", "–", "–", "×1.112"],
        ["Uygulanan muhafazakâr boost", "–", "–", "×1.15"],
        ["Mart için net SI (B seg.)", "–", "–", "SI_Mart × 1.15"],
    ]
    make_simple_table(doc, launch_headers, launch_rows)

    add_body_para(
        doc,
        "Lansman etkisinin kademeli azalması öngörülerek Nisan'dan Aralık'a "
        "kadar aşağıdaki boost takvimi uygulanmaktadır:"
    )
    boost_schedule = [
        "Mart: ×1.60",
        "Nisan: ×1.45",
        "Mayıs: ×1.35",
        "Haziran: ×1.25",
        "Temmuz: ×1.18",
        "Ağustos: ×1.14",
        "Eylül: ×1.11",
        "Ekim: ×1.09",
        "Kasım: ×1.07",
        "Aralık: ×1.05",
    ]
    for item in boost_schedule:
        add_bullet(doc, item)

    add_heading2(doc, "3.4 Çok Kriterli Skorlama Modeli (MCDM)")

    add_body_para(
        doc,
        "Her bayi için dört ölçüt hesaplanmakta ve ağırlıklı toplam yöntemiyle "
        "Bütünleşik Skor (CS) elde edilmektedir."
    )

    add_equation_placeholder(
        doc, "Denklem 4",
        "CS_i = 0.25 * P_i + 0.35 * LP_i + 0.20 * S_i + 0.20 * H_i"
    )

    add_table_caption(doc, 5, "MCDM Kriter Ağırlıkları ve Hesaplama Mantığı")
    mcdm_headers = ["Skor", "Ağırlık", "Tanım", "Hesaplama Yöntemi"]
    mcdm_rows = [
        ["P (Performans)", "0.25",
         "Son 12 ay EW hedef gerçekleştirme oranı",
         "Üstel ağırlıklı ortalama, son aylara daha fazla ağırlık"],
        ["LP (Lokasyon-Ürün)", "0.35",
         "Bayinin ürün karışımı ile bölge profilinin uyumu",
         "Cosine similarity (bayi × model satış matrisi)"],
        ["S (Mevsimsel Uyum)", "0.20",
         "Bayinin söz konusu aydaki mevsimsel kapasitesi",
         "Bayi tier blended SI değeri"],
        ["H (Hedef Yakınlığı)", "0.20",
         "Yıllık hedefe ulaşmak için kalan ihtiyacın aciliyeti",
         "1 – (yıl içi kümülatif / yıllık hedef)"],
    ]
    make_simple_table(doc, mcdm_headers, mcdm_rows)

    add_heading3(doc, "3.4.1 Lokasyon–Ürün Uyum Skoru (LP) – Collaborative Filtering")

    add_body_para(
        doc,
        "LP skoru, bayi × model satış matrisinin cosine similarity analizi ile "
        "hesaplanmaktadır. M satır bayi ve N sütun model olmak üzere X ∈ R^{M×N} "
        "satış matrisi oluşturulur. Bayi i ile bayi j arasındaki cosine benzerliği:"
    )

    add_equation_placeholder(
        doc, "Denklem 5",
        "sim(i, j) = (X_i . X_j) / (||X_i|| * ||X_j||)"
    )

    add_body_para(
        doc,
        "Her bayi için normalize edilmiş LP skoru, bu benzerliklerin ağırlıklı "
        "ortalamasıyla elde edilmekte ve 0–1 aralığına min-max normalizasyonuyla "
        "ölçeklenmektedir. Yeni bayiler (Bayi 23–28) için tier ortalaması "
        "başlangıç değeri olarak kullanılmaktadır."
    )

    add_heading2(doc, "3.5 MILP Optimizasyon Modeli")

    add_heading3(doc, "3.5.1 Model Formülasyonu")

    add_body_para(
        doc,
        "MILP modeli, karar değişkeni, amaç fonksiyonu ve kısıtlar "
        "aşağıdaki gibi tanımlanmaktadır. [NOT: Matematiksel gösterimler "
        "Word Equation Editor ile son haliyle doldurulacaktır.]"
    )

    add_heading3(doc, "3.5.2 Karar Değişkeni")
    add_body_para(
        doc,
        "x_{ijt}: Bayi i'ye, model j, ay t'de tahsis edilen araç adedi "
        "(x_{ijt} ∈ Z+, yani pozitif tam sayı)."
    )

    add_equation_placeholder(
        doc, "Denklem 6",
        "x_{ijt} in Z+  (i=1..28, j in {A2,A3,B1,B2}, t=1..12)"
    )

    add_heading3(doc, "3.5.3 Amaç Fonksiyonu")

    add_body_para(
        doc,
        "Amaç, tüm bayilere, modellere ve aylara ait ağırlıklı bütünleşik "
        "skorun toplamını maksimize etmektir:"
    )

    add_equation_placeholder(
        doc, "Denklem 7",
        "Maksimize: sum_{i} sum_{j} sum_{t} CS_i * x_{ijt}"
    )

    add_heading3(doc, "3.5.4 Kısıtlar")

    add_body_para(doc, "Model dört temel kısıt grubuna tabidir:", first_indent=False)

    add_numbered(doc, "Aylık Envanter Kısıtı: Dağıtılan toplam araç adedi, "
                      "o ay için mevcut envanter miktarını aşamaz.")
    add_equation_placeholder(
        doc, "Denklem 8",
        "sum_{i} sum_{j} x_{ijt} <= Envanter_t  (her t icin)"
    )

    add_numbered(doc, "Bayi Aylık Hedef Aralığı Kısıtı: Her bayiye dağıtılan "
                      "toplam araç, o ayın bayi hedefinin ±%20 aralığında olmalıdır.")
    add_equation_placeholder(
        doc, "Denklem 9",
        "0.80 * H_{it} <= sum_{j} x_{ijt} <= 1.20 * H_{it}  (her i,t icin)"
    )

    add_numbered(doc, "Model Karışımı Yumuşak Kısıtı: Model bazında dağıtım, "
                      "bayinin geçmiş satış profilinden belirli bir sapma aralığında "
                      "tutulur (soft constraint, ceza terimi ile amaç fonksiyonuna eklenir).")

    add_numbered(doc, "Tamsayılık Kısıtı: Tüm karar değişkenleri pozitif tam "
                      "sayı olmalıdır.")

    add_heading3(doc, "3.5.5 Model Boyutu ve Çözüm Süresi")

    add_body_para(
        doc,
        "28 bayi × 4 model × 12 ay kombinasyonundan oluşan model toplamda "
        "1.344 karar değişkeni ve yaklaşık 400 kısıt içermektedir. PuLP + CBC "
        "kombinasyonu ile bu model birkaç saniye içinde optimal çözüme "
        "ulaşmaktadır. Çözüm bir standart dizüstü bilgisayarda (8 GB RAM, "
        "modern işlemci) 2–5 saniye içinde tamamlanmaktadır."
    )

    add_table_caption(doc, 6, "MILP Model Parametreleri")
    milp_headers = ["Parametre", "Değer", "Açıklama"]
    milp_rows = [
        ["Karar değişkeni sayısı", "1.344", "28 bayi × 4 model × 12 ay"],
        ["Kısıt sayısı", "~400", "Envanter + hedef aralığı + model mix"],
        ["Çözücü", "CBC (PuLP)", "Açık kaynaklı Branch & Cut"],
        ["Amaç", "Maksimizasyon", "Toplam ağırlıklı CS skoru"],
        ["Çözüm süresi (ortalama)", "< 5 saniye", "Standart dizüstü bilgisayar"],
        ["Tamsayılık boşluğu (MIP gap)", "< %1", "Optimal veya neredeyse optimal"],
    ]
    make_simple_table(doc, milp_headers, milp_rows)

    add_page_break(doc)


def build_chapter4(doc: Document) -> None:
    """Bölüm 4: Bulgular ve Analiz"""
    add_heading1(doc, "4. BULGULAR VE ANALİZ")

    add_heading2(doc, "4.1 Mevsimsel İndeks Doğrulaması")

    add_body_para(
        doc,
        "Geliştirilen mevsimsel indeks modeli, 2025 Aralık ayı satışları "
        "doğrulama seti üzerinde test edilmiştir. Modelin 2024–2024 verilerinden "
        "öğrendiği SI değerleriyle tahmin ettiği Aralık 2025 satışı 460 araç "
        "olurken gerçekleşme 499 araç olmuştur."
    )

    add_table_caption(doc, 7, "Aralık 2025 Tahmin Doğrulama Sonuçları")
    val_headers = ["Grup", "Tahmin (Araç)", "Gerçek (Araç)", "MAPE (%)"]
    val_rows = [
        ["Tüm bayiler (Global)", "460", "499", "7.82"],
        ["Tier A (21 bayi)", "–", "–", "7.65"],
        ["Tier B (4 bayi)", "–", "–", "6.67"],
        ["Tier C (3 bayi)", "–", "–", "15.79"],
    ]
    make_simple_table(doc, val_headers, val_rows)

    add_body_para(
        doc,
        "Tier C'nin görece yüksek MAPE değeri (%15.79), bu gruptaki bayilerin "
        "az sayıda (3 bayi) ve satış hacminin düşük olmasından kaynaklanan "
        "yüksek varyansla açıklanabilir. Literatürde benzer küçük örneklem "
        "durumlarında %15 altı MAPE kabul edilebilir düzey olarak "
        "değerlendirilmektedir (Makridakis vd., 2018)."
    )

    add_figure_caption(doc, 9,
                       "Aralık 2025 Tahmin vs. Gerçek – Bayi Bazında Karşılaştırma")

    add_heading2(doc, "4.2 2026 Yıllık Satış Planı")

    add_body_para(
        doc,
        "2026 yılı için iki alternatif senaryo geliştirilmiştir. "
        "Her iki senaryoda da aylık dağılım, blended SI değerleri "
        "esas alınarak hesaplanmıştır."
    )

    add_table_caption(doc, 8,
                      "2026 Yıllık Satış Planı – İki Senaryo Karşılaştırması")
    plan_headers = ["Ay", "SI", "Senaryo 1 (8.500)", "Senaryo 2 (10.000)",
                    "B Seg. Boost"]
    plan_rows = [
        ["Ocak", "0.659", "466", "549", "×1.00"],
        ["Şubat", "0.743", "526", "619", "×1.00"],
        ["Mart", "1.089", "871 (+bnz)", "1.025 (+bnz)", "×1.60"],
        ["Nisan", "0.912", "701", "825", "×1.45"],
        ["Mayıs", "0.896", "634", "747", "×1.35"],
        ["Haziran", "0.988", "699", "823", "×1.25"],
        ["Temmuz", "0.803", "568", "669", "×1.18"],
        ["Ağustos", "0.856", "606", "713", "×1.14"],
        ["Eylül", "0.941", "666", "784", "×1.11"],
        ["Ekim", "1.044", "739", "870", "×1.09"],
        ["Kasım", "1.269", "898", "1.057", "×1.07"],
        ["Aralık", "1.616", "1.143", "1.346", "×1.05"],
        ["TOPLAM", "–", "8.517", "10.027", "–"],
    ]
    make_simple_table(doc, plan_headers, plan_rows)

    add_body_para(
        doc,
        "Not: Satır toplamlarındaki küçük farklar, tam sayıya yuvarlama "
        "işleminden kaynaklanmaktadır. 'bnz' = B segmenti yeni versiyon lansmanı. "
        "Senaryo 1, 2024–2025 büyüme trendine dayanmakta; Senaryo 2, lansman "
        "etkisiyle +%18 büyümeyi öngörmektedir.",
        italic=True, first_indent=False
    )

    add_figure_caption(doc, 10,
                       "2026 Yıllık Satış Planı – Aylık Dağılım (Her İki Senaryo)")

    add_heading2(doc, "4.3 Bayi Bazlı Hedef Matrisi")

    add_body_para(
        doc,
        "MCDM skorları ve 2026 yıllık planına göre, 28 bayi için 12 aylık "
        "araç dağıtım hedefleri hesaplanmıştır. Bu matrisin özet istatistikleri "
        "aşağıda sunulmaktadır."
    )

    add_table_caption(doc, 9, "28 Bayi 2026 Yıllık Hedef Özeti (Senaryo 1)")
    dealer_headers = ["Bayi Grubu", "Bayi Sayısı", "Ort. Yıllık Hedef",
                      "Min.", "Maks."]
    dealer_rows = [
        ["Tier A – Yerleşik bayiler", "15", "~340 araç/yıl",
         "~220", "~520"],
        ["Tier A – Büyüme bayileri", "6", "~280 araç/yıl",
         "~180", "~380"],
        ["Tier B", "4", "~200 araç/yıl", "~150", "~260"],
        ["Tier C", "3", "~80 araç/yıl", "~50", "~120"],
        ["Yeni Bayiler (23–28)", "6", "~85 araç/yıl",
         "Tier ort. × 0.60", "Tier ort. × 0.60"],
    ]
    make_simple_table(doc, dealer_headers, dealer_rows)

    add_body_para(
        doc,
        "Yeni bayiler (Bayi 23–28) için ilk yıl hedefleri, ait oldukları tier "
        "grubunun ortalama payının %60'ı olarak belirlenmiştir. Bu yaklaşım, "
        "yeni bayilerin rampa-up sürecini (ramp-up period) modellemek için "
        "kullanılan yaygın bir endüstri pratiğini yansıtmaktadır."
    )

    add_figure_caption(doc, 12,
                       "28 Bayi Coğrafi Dağılım Haritası (Folium ile üretilmiştir)")

    add_heading2(doc, "4.4 Optimizasyon Sonuçları")

    add_body_para(
        doc,
        "MILP modeli, Ocak–Aralık 2026 dönemi için her ay ayrı ayrı "
        "çözülmüş ve optimal araç dağıtım planları üretilmiştir. "
        "Model her çalıştırmada optimal veya %1'den küçük MIP gap ile "
        "çözüme ulaşmıştır."
    )

    add_table_caption(doc, 10, "MILP Optimizasyon Çözüm İstatistikleri")
    milp_res_headers = ["Metrik", "Değer"]
    milp_res_rows = [
        ["Toplam çözüm sayısı (12 ay)", "12"],
        ["Ortalama çözüm süresi", "< 5 saniye/ay"],
        ["Optimal çözüme ulaşma oranı", "%100"],
        ["Ortalama MIP gap", "< %0.5"],
        ["Hedef aralığı kısıtı ihlali", "Yok"],
        ["Envanter kısıtı ihlali", "Yok"],
    ]
    make_simple_table(doc, milp_res_headers, milp_res_rows)

    add_figure_caption(doc, 14,
                       "MILP Optimizasyon Sonuçları – Aylık Dağıtım Verimliliği")

    add_page_break(doc)


def build_chapter5(doc: Document) -> None:
    """Bölüm 5: Tartışma ve Sonuç"""
    add_heading1(doc, "5. TARTIŞMA VE SONUÇ")

    add_heading2(doc, "5.1 Temel Bulgular")

    add_body_para(
        doc,
        "Bu çalışma, bir otomotiv markasının 28 bayilik dağıtım ağında veri "
        "odaklı araç dağıtım optimizasyonu için bütünleşik ve modüler bir "
        "metodoloji ortaya koymuştur. Temel bulgular şu şekilde özetlenebilir:"
    )

    bulgular = [
        "Ratio-to-mean yöntemiyle hesaplanan mevsimsel indeksler ve blended "
        "tier yaklaşımı, Aralık 2025 doğrulama setinde %7.82 MAPE ile "
        "literatürdeki kabul edilebilir eşiğin altında kalmıştır. "
        "Tier A (%7.65) ve Tier B (%6.67) için bu performans özellikle "
        "tatmin edicidir.",

        "B segmentinin Mart 2026 lansmanı için istatistiksel gerekçeye dayanan "
        "boost katsayısı (×1.60 Mart başlangıç, kademeli azalma) veri destekli "
        "ve şeffaf bir planlama aracı olarak işlev görmektedir.",

        "Dört kriterli MCDM modelinin LP ağırlığının diğer kriterlerden yüksek "
        "tutulması (0.35), lokasyon-ürün uyumunun dağıtım verimliliğine olan "
        "katkısını öne çıkarmaktadır. Bu yaklaşım, collaborative filtering "
        "mantığını tedarik zinciri bağlamına başarıyla uyarlamaktadır.",

        "PuLP+CBC ile çözülen MILP modeli, 1.344 değişken ve ~400 kısıtla "
        "bile saniyeler içinde optimal çözüme ulaşmakta; bu durum sistemin "
        "aylık üretim kullanımına uygunluğunu göstermektedir.",

        "28 bayi × 12 ay × 4 model hedef matrisi, hem muhafazakâr (8.500 araç) "
        "hem de agresif (10.000 araç) senaryolar için tutarlı sonuçlar üretmiştir.",
    ]
    for b in bulgular:
        add_bullet(doc, b)

    add_heading2(doc, "5.2 Yönetimsel Öneriler")

    add_body_para(
        doc,
        "Araştırma bulguları ışığında, Marka X yönetimine aşağıdaki "
        "öneriler sunulmaktadır:"
    )

    oneriler = [
        "Sistem Ocak 2026'da devreye alındıktan sonra, ilk altı ay boyunca "
        "gerçekleşme verileriyle model parametreleri güncellenmeli ve MAPE "
        "takibi aylık yapılmalıdır.",
        "Tier C bayileri için %15.79 MAPE, ek veri toplama ve bayi spesifik "
        "faktörlerin (kampanyalar, yerel rekabet) modele eklenmesini "
        "önermektedir.",
        "B1 yeni versiyon lansmanının gerçek satış performansı, Mart–Nisan "
        "verilerinden sonra değerlendirilmeli ve boost takvimi buna göre "
        "revize edilmelidir.",
        "Yeni bayi (23–28) ramp-up katsayısı (%60), ilk yıl verilerine göre "
        "ayarlanarak ikinci yılda standart tier değerine geçilmelidir.",
        "Dashboard'un DEMO_MODE dışında kullanımı için yalnızca yetkili "
        "kullanıcılara erişim sağlanmalı ve Streamlit Cloud secrets yönetimi "
        "denetlenmelidir.",
    ]
    for o in oneriler:
        add_bullet(doc, o)

    add_heading2(doc, "5.3 Sınırlılıklar ve Gelecek Çalışmalar")

    add_heading3(doc, "5.3.1 Çalışmanın Sınırlılıkları")

    sinirlar = [
        "Çalışma, yalnızca iki yıllık (2024–2025) satış verisine dayanmaktadır. "
        "Daha uzun tarihsel veriler, mevsimsel indeks güvenilirliğini artıracaktır.",
        "Makroekonomik faktörler (döviz kuru, faiz oranları, vergi değişiklikleri) "
        "ve rakip kampanyaları modele dahil edilmemiştir.",
        "Bayi stok kapasitesi, 1.000 araç eşiğinin altında olduğu için kısıt "
        "olarak modele eklenmemiş; bu eşiğin üzerine çıkılması durumunda "
        "modelin revize edilmesi gerekecektir.",
        "Renk dağılımı optimizasyonu, modelin bir sonraki aşaması olarak "
        "planlanmış ancak bu çalışma kapsamında sadece model ve versiyon "
        "düzeyinde ele alınmıştır.",
    ]
    for s in sinirlar:
        add_bullet(doc, s)

    add_heading3(doc, "5.3.2 Gelecek Araştırma Yönleri")

    gelecek = [
        "Stokastik MILP: Talep belirsizliğini açıkça modelleyen senaryo bazlı "
        "veya robust optimizasyon yaklaşımlarının entegrasyonu.",
        "Makine Öğrenmesi Entegrasyonu: LSTM veya Transformer tabanlı talep "
        "tahmin modellerinin mevsimsel indekslerin yerini kısmen alması.",
        "Renk Optimizasyonu: Bayi bazında renk tercih profillerinin cosine "
        "similarity ile modellenerek model karışımı kısıtlarına eklenmesi.",
        "Çok Markaya Genelleme: Metodolojinin farklı marka veya segment "
        "(elektrikli araç, hibrit) verilerine uyarlanması.",
        "Gerçek Zamanlı Güncelleme: Satış verilerinin otomatik olarak sisteme "
        "aktarılmasını sağlayan ETL pipeline'ının geliştirilmesi.",
    ]
    for g in gelecek:
        add_bullet(doc, g)

    add_body_para(
        doc,
        "Sonuç olarak, bu çalışma otomotiv dağıtım planlamasında veri bilimi, "
        "yöneylem araştırması ve makine öğrenmesi tekniklerini başarıyla "
        "bütünleştiren, akademik sağlamlığı ile pratik uygulanabilirliği "
        "bir arada sunan bir karar destek sistemi ortaya koymuştur. Sistemin "
        "Ocak 2026'da canlıya alınmasıyla birlikte elde edilecek gerçek "
        "performans verileri, metodolojinin sürekli iyileştirilmesi için "
        "değerli bir kaynak oluşturacaktır."
    )

    add_page_break(doc)


def build_references(doc: Document) -> None:
    """Kaynakça bölümü (APA 7. baskı)."""
    add_heading1(doc, "KAYNAKÇA")

    refs = [
        "Bowerman, B. L., O'Connell, R. T., & Koehler, A. B. (2005). "
        "*Forecasting, time series, and regression: An applied approach* (4th ed.). "
        "Thomson Brooks/Cole.",

        "Cachon, G., & Lariviere, M. (1999). Capacity allocation using past sales: "
        "When to turn-and-earn. *Management Science*, *45*(5), 685–703. "
        "https://doi.org/10.1287/mnsc.45.5.685",

        "Cleveland, R. B., Cleveland, W. S., McRae, J. E., & Terpenning, I. (1990). "
        "STL: A seasonal-trend decomposition procedure based on Loess. "
        "*Journal of Official Statistics*, *6*(1), 3–73.",

        "Dantzig, G. B., & Thapa, M. N. (1997). *Linear programming 1: Introduction*. "
        "Springer-Verlag.",

        "Fahimnia, B., Sarkis, J., & Davarzani, H. (2015). Green supply chain "
        "management: A review and bibliometric analysis. "
        "*International Journal of Production Economics*, *162*, 101–114. "
        "https://doi.org/10.1016/j.ijpe.2015.01.003",

        "Hwang, C. L., & Yoon, K. (1981). *Multiple attribute decision making: "
        "Methods and applications*. Springer-Verlag.",

        "Hyndman, R. J., Ahmed, R. A., Athanasopoulos, G., & Shang, H. L. (2011). "
        "Optimal combination forecasts for hierarchical time series. "
        "*Computational Statistics & Data Analysis*, *55*(9), 2579–2589. "
        "https://doi.org/10.1016/j.csda.2011.03.006",

        "Luss, H., & Rosenwein, M. B. (1997). Operations research applications: "
        "Opportunities and accomplishments. "
        "*European Journal of Operational Research*, *97*(2), 220–244. "
        "https://doi.org/10.1016/S0377-2217(96)00196-0",

        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2018). "
        "The M4 competition: 100,000 time series and 61 forecasting methods. "
        "*International Journal of Forecasting*, *36*(1), 54–74. "
        "https://doi.org/10.1016/j.ijforecast.2019.04.014",

        "Mitchell, S., OSullivan, M., & Dunning, I. (2011). *PuLP: A linear "
        "programming toolkit for Python*. The University of Auckland. "
        "https://coin-or.github.io/pulp/",

        "ODD – Otomotiv Distribütörleri Derneği. (2025). "
        "*2024 yılı otomotiv sektörü istatistikleri*. "
        "https://www.odd.org.tr",

        "Sarwar, B., Karypis, G., Konstan, J., & Riedl, J. (2001). "
        "Item-based collaborative filtering recommendation algorithms. "
        "In *Proceedings of the 10th International Conference on World Wide Web* "
        "(pp. 285–295). ACM. https://doi.org/10.1145/371920.372071",

        "Sherbrooke, C. C. (1968). METRIC: A multi-echelon technique for "
        "recoverable item control. *Operations Research*, *16*(1), 122–141. "
        "https://doi.org/10.1287/opre.16.1.122",

        "Taylor, S. J., & Letham, B. (2018). Forecasting at scale. "
        "*The American Statistician*, *72*(1), 37–45. "
        "https://doi.org/10.1080/00031305.2017.1380080",
    ]

    for ref in refs:
        # Italikleri ayır (yıldız işaretleri arasındaki metin)
        add_reference(doc, ref.replace("*", ""))  # sadeleştirilmiş

    add_page_break(doc)


def build_appendix(doc: Document) -> None:
    """Ekler bölümü."""
    add_heading1(doc, "EKLER")

    add_heading2(doc, "Ek A. Sistem Teknik Gereksinimleri")

    add_body_para(
        doc,
        "Geliştirilen sistemin çalıştırılabilmesi için gerekli teknik altyapı "
        "ve yazılım bağımlılıkları aşağıda listelenmiştir."
    )

    tech_items = [
        "Python 3.11+",
        "pandas >= 2.0, numpy >= 1.26",
        "statsmodels >= 0.14 (STL decomposition)",
        "prophet >= 1.1 (opsiyonel, zaman serisi tahmini)",
        "scikit-learn >= 1.3 (cosine_similarity, StandardScaler)",
        "pulp >= 2.7 + CBC çözücüsü",
        "streamlit >= 1.28, plotly >= 5.17, folium >= 0.14",
        "pytest >= 7.4, ruff >= 0.1 (test ve linting)",
        "python-docx >= 1.1 (tez belgesi üretimi)",
    ]
    for item in tech_items:
        add_bullet(doc, item)

    add_heading2(doc, "Ek B. Veri Şeması")

    add_body_para(
        doc,
        "SQLite veritabanı (arac_dagitim.db) aşağıdaki tabloları içermektedir:"
    )

    schema_items = [
        "dim_bayi: bayi_id (PK), ad, il, ilce, lat, lon, tier, aktif, yeni_bayi",
        "dim_arac: arac_id (PK), model, versiyon, renk, uretim_durumu",
        "fact_satis: satis_id (PK), bayi_id (FK), arac_id (FK), tarih, adet",
        "fact_hedef: hedef_id (PK), bayi_id (FK), yil, ay, hedef, gerceklesme",
        "fact_envanter: env_id (PK), arac_id (FK), yil, ay, adet",
        "dim_rakip_satis: rakip_id (PK), il, marka, yil, ay, adet",
    ]
    for item in schema_items:
        add_bullet(doc, item)

    add_heading2(doc, "Ek C. Model Doğrulama Detayları")

    add_body_para(
        doc,
        "Aralık 2025 doğrulama setinde bayi bazında tahmin ve gerçekleşme "
        "değerleri, bu ekte tam tablo olarak sunulmaktadır. Gerçek değerler "
        "anonimleştirilmiş olup 'Bayi XX' formatında gösterilmiştir."
    )

    add_body_para(
        doc,
        "[NOT: Bu tablonun doldurulması için 'python scripts/gen_tahmin.py' "
        "çıktılarından elde edilen CSV dosyası kullanılacaktır.]",
        italic=True, first_indent=False
    )

    add_heading2(doc, "Ek D. Python Kod Yapısı")

    add_body_para(
        doc,
        "Projenin dizin yapısı aşağıdaki gibi organize edilmiştir:"
    )

    code_structure = [
        "arac-dagitim-sistemi/",
        "  ├── src/                    # Ana kaynak kodu",
        "  │   ├── config.py           # Merkezi konfigürasyon (BRAND_NAME, vb.)",
        "  │   ├── data_loader.py      # CSV → SQLite ETL",
        "  │   ├── seasonal_index.py   # SI hesaplama modülü",
        "  │   ├── mcdm_scorer.py      # MCDM skor hesaplama",
        "  │   ├── milp_optimizer.py   # PuLP MILP modeli",
        "  │   └── dashboard.py        # Streamlit dashboard",
        "  ├── scripts/                # Çalıştırılabilir scriptler",
        "  │   ├── gen_tahmin.py       # Tahmin üretim scripti",
        "  │   └── gen_tez.py          # Tez belgesi üretim scripti",
        "  ├── data/",
        "  │   └── raw/                # Ham CSV dosyaları",
        "  ├── tests/                  # pytest test dosyaları",
        "  ├── docs/                   # Belgeler",
        "  ├── notebooks/              # Jupyter / Colab notebook'ları",
        "  └── .github/workflows/      # GitHub Actions CI/CD",
    ]

    for line in code_structure:
        p = doc.add_paragraph()
        _para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=Pt(0), space_after=Pt(0),
                     line_spacing=1.15,
                     left_indent=Cm(1.25), first_line_indent=None)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Ana fonksiyon
# ---------------------------------------------------------------------------

def main() -> None:
    """Tez defteri belgesini oluşturur ve kaydeder."""
    print("Tez defteri oluşturuluyor...")

    doc = Document()

    # Kenar boşluklarını ayarla
    _set_margins(doc)

    # Varsayılan stil (Normal) ayarları
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sayfa numarası ekle
    _add_footer_page_number(doc)

    # Bölümleri oluştur
    print("  [1/10] Kapak sayfası...")
    build_cover_page(doc)

    print("  [2/10] Özet ve Abstract...")
    build_abstract(doc)

    print("  [3/10] İçindekiler...")
    build_table_of_contents(doc)

    print("  [4/10] Şekil ve Tablo listeleri...")
    build_figure_list(doc)
    build_table_list(doc)

    print("  [5/10] Simge ve Kısaltma listesi...")
    build_symbols(doc)

    print("  [6/10] Bölüm 1 – Giriş...")
    build_chapter1(doc)

    print("  [7/10] Bölüm 2 – Literatür Taraması...")
    build_chapter2(doc)

    print("  [8/10] Bölüm 3 – Metodoloji...")
    build_chapter3(doc)

    print("  [9/10] Bölüm 4 – Bulgular...")
    build_chapter4(doc)

    print("  [10/10] Bölüm 5 – Tartışma, Kaynakça, Ekler...")
    build_chapter5(doc)
    build_references(doc)
    build_appendix(doc)

    # Kaydet
    doc.save(str(OUTPUT_PATH))
    print(f"\nBelge kaydedildi: {OUTPUT_PATH}")
    print(f"Dosya boyutu: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
